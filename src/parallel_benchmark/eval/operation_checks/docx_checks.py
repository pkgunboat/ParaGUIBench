"""
Word 文档（.docx）属性检查原语。

每个函数接收文件路径和参数字典，返回标准化结果：
    {"pass": bool, "score": float 0.0~1.0, "reason": str}

依赖: python-docx
"""

import glob
import hashlib
import logging
import os
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, Cm, Emu

logger = logging.getLogger("eval.operation_checks.docx")


def _run_rgb(run) -> Optional[str]:
    """读取 Word run 的显式 RGB 颜色。

    输入:
        run: python-docx Run 对象。
    输出:
        六位大写 RGB 字符串；未显式设置或使用主题色时返回 None。
    """
    color = run.font.color
    if color is None or color.rgb is None:
        return None
    value = str(color.rgb).upper()
    return value[-6:] if len(value) >= 6 else value


def _is_red_run(run) -> bool:
    """判断 Word run 是否显式设为纯红色。

    输入:
        run: python-docx Run 对象。
    输出:
        RGB 为 FF0000 时返回 True，否则返回 False。
    """
    return _run_rgb(run) == "FF0000"


# ------------------------------------------------------------------
# 工具函数
# ------------------------------------------------------------------

def _load_document(file_path: str) -> Optional[Document]:
    """
    安全加载 docx 文件。

    输入:
        file_path: docx 文件路径
    输出:
        Document 对象；加载失败返回 None
    """
    try:
        return Document(file_path)
    except Exception as exc:
        logger.error("无法打开 docx 文件 %s: %s", file_path, exc)
        return None


def _ok(reason: str = "通过") -> Dict[str, Any]:
    """构造通过结果。"""
    return {"pass": True, "score": 1.0, "reason": reason}


def _fail(reason: str) -> Dict[str, Any]:
    """构造失败结果。"""
    return {"pass": False, "score": 0.0, "reason": reason}


def _partial(score: float, reason: str) -> Dict[str, Any]:
    """构造部分通过结果。严格阈值：仅当 score 等于 1.0 才算 pass。"""
    return {"pass": score >= 1.0 - 1e-9, "score": round(score, 4), "reason": reason}


def _config_error(reason: str) -> Dict[str, Any]:
    """评价器配置错误（缺参数等）：score=-1 哨兵，由上层冒泡为 evaluator_error。"""
    return {"pass": False, "score": -1.0, "status": "evaluator_error", "reason": reason}


# ------------------------------------------------------------------
# 检查函数
# ------------------------------------------------------------------

def check_max_consecutive_blank_lines(file_path: str, params: dict) -> dict:
    """
    检查文档中连续空段落数是否超过允许的最大值。

    输入:
        file_path: docx 文件路径
        params:
            max_allowed (int): 允许的最大连续空行数，默认 1
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    max_allowed = params.get("max_allowed", 1)
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    max_found = 0
    current_streak = 0
    violations = []

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "":
            current_streak += 1
            if current_streak > max_found:
                max_found = current_streak
            if current_streak > max_allowed:
                violations.append(i)
        else:
            current_streak = 0

    if max_found <= max_allowed:
        return _ok(f"最大连续空行 {max_found} ≤ {max_allowed}")

    return _fail(
        f"发现连续空行 {max_found} 行（允许 {max_allowed}），"
        f"共 {len(violations)} 处违规"
    )


def check_font_name(file_path: str, params: dict) -> dict:
    """
    检查文档中所有 run 的字体是否为指定字体名称。

    输入:
        file_path: docx 文件路径
        params:
            font_name (str): 期望的字体名称，如 "Times New Roman"
            threshold (float): 符合比例阈值，默认 0.9（90% 以上的 run 字体正确即通过）
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    expected_font = params.get("font_name", "")
    threshold = params.get("threshold", 0.9)

    if not expected_font:
        return _config_error("参数缺少 font_name")

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    total_runs = 0
    matched_runs = 0

    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            total_runs += 1
            run_font = run.font.name
            # 有些 run 继承段落/样式字体，font.name 可能为 None
            if run_font and run_font.lower() == expected_font.lower():
                matched_runs += 1
            elif run_font is None:
                # 尝试从段落样式获取字体
                style_font = para.style.font.name if para.style and para.style.font else None
                if style_font and style_font.lower() == expected_font.lower():
                    matched_runs += 1

    if total_runs == 0:
        return _ok("文档无可检查的文本 run")

    ratio = matched_runs / total_runs
    if ratio >= threshold:
        return _ok(f"字体匹配率 {ratio:.1%}（{matched_runs}/{total_runs}）")

    return _partial(
        ratio,
        f"字体匹配率 {ratio:.1%}（{matched_runs}/{total_runs}），期望 ≥ {threshold:.0%}"
    )


def check_line_spacing(file_path: str, params: dict) -> dict:
    """
    检查文档段落的行距是否为指定值。

    输入:
        file_path: docx 文件路径
        params:
            spacing (float): 期望的行距倍数，如 2.0 表示双倍行距
            threshold (float): 符合比例阈值，默认 0.9
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    expected_spacing = params.get("spacing")
    threshold = params.get("threshold", 0.9)

    if expected_spacing is None:
        return _config_error("参数缺少 spacing")

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    total_paras = 0
    matched_paras = 0

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        total_paras += 1

        pf = para.paragraph_format
        if pf.line_spacing is None:
            continue

        # line_spacing 可能是 float（倍数）或 Pt 值（固定磅值）
        if pf.line_spacing_rule in (
            WD_LINE_SPACING.MULTIPLE,
            WD_LINE_SPACING.DOUBLE,
            WD_LINE_SPACING.ONE_POINT_FIVE,
        ):
            actual = float(pf.line_spacing)
        elif isinstance(pf.line_spacing, (int, float)):
            actual = float(pf.line_spacing)
        else:
            # Emu/Pt 固定值，转换为近似倍数（基准 12pt）
            try:
                actual = float(pf.line_spacing) / Pt(12)
            except Exception:
                continue

        if abs(actual - expected_spacing) < 0.05:
            matched_paras += 1

    if total_paras == 0:
        return _ok("文档无可检查的非空段落")

    ratio = matched_paras / total_paras
    if ratio >= threshold:
        return _ok(f"行距匹配率 {ratio:.1%}（{matched_paras}/{total_paras}）")

    return _partial(
        ratio,
        f"行距匹配率 {ratio:.1%}（{matched_paras}/{total_paras}），"
        f"期望 {expected_spacing} 倍行距，阈值 ≥ {threshold:.0%}"
    )


def check_heading_hierarchy(file_path: str, params: dict) -> dict:
    """
    检查文档的标题层级是否符合预期规则。

    输入:
        file_path: docx 文件路径
        params:
            rules (list[dict]): 标题规则列表，每项格式：
                {
                    "pattern": str,          # 正则表达式匹配标题文本
                    "expected_style": str,   # 期望的样式名称，如 "Heading 1"
                }
            threshold (float): 符合比例阈值，默认 0.8
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    rules = params.get("rules", [])
    threshold = params.get("threshold", 0.8)

    if not rules:
        return _config_error("参数缺少 rules")

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    total_checks = 0
    matched_checks = 0
    mismatches = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        for rule in rules:
            pattern = rule.get("pattern", "")
            expected_style = rule.get("expected_style", "")
            if not pattern or not expected_style:
                continue

            if re.search(pattern, text):
                total_checks += 1
                actual_style = para.style.name if para.style else ""
                if actual_style == expected_style:
                    matched_checks += 1
                else:
                    mismatches.append({
                        "text": text[:50],
                        "expected": expected_style,
                        "actual": actual_style,
                    })

    if total_checks == 0:
        return _fail("未匹配到需要检查的标题")

    ratio = matched_checks / total_checks
    if ratio >= threshold:
        return _ok(
            f"标题层级匹配率 {ratio:.1%}（{matched_checks}/{total_checks}）"
        )

    mismatch_summary = "; ".join(
        f"'{m['text']}' 期望 {m['expected']} 实际 {m['actual']}"
        for m in mismatches[:5]
    )
    return _partial(
        ratio,
        f"标题层级匹配率 {ratio:.1%}，不匹配: {mismatch_summary}"
    )


def check_has_toc(file_path: str, params: dict) -> dict:
    """
    检查文档是否包含目录（Table of Contents）。

    通过检测 TOC 域代码或特定样式（TOC Heading / TOC 1-9）来判断。

    输入:
        file_path: docx 文件路径
        params: {}（无额外参数）
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    # 方法 1：检查 XML 中是否存在 TOC 域代码
    from lxml import etree
    body = doc.element.body
    # w:fldChar + w:instrText 中包含 "TOC" 字样
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    instr_texts = body.findall(".//w:instrText", namespaces)
    for instr in instr_texts:
        if instr.text and "TOC" in instr.text.upper():
            return _ok("检测到 TOC 域代码")

    # 方法 2：必须存在至少两条非空 TOC 1..9 目录条目。
    # "TOC Heading" 只是目录标题，单独出现不能证明已生成目录。
    toc_entry_pattern = re.compile(r"^(?:toc\s*[1-9]|目录\s*[1-9])$", re.IGNORECASE)
    toc_entries = [
        para for para in doc.paragraphs
        if para.text.strip()
        and para.style
        and para.style.name
        and toc_entry_pattern.match(para.style.name.strip())
    ]
    if len(toc_entries) >= 2:
        return _ok(f"检测到 {len(toc_entries)} 条 TOC 目录条目")

    # 方法 3：检查 SDT（结构化文档标签）中的 TOC
    sdt_elements = body.findall(".//w:sdt", namespaces)
    for sdt in sdt_elements:
        sdt_pr = sdt.find("w:sdtPr", namespaces)
        if sdt_pr is not None:
            doc_part = sdt_pr.find("w:docPartGallery", namespaces)
            if doc_part is not None:
                val = doc_part.get(f"{{{namespaces['w']}}}val", "")
                if "toc" in val.lower() or "table of contents" in val.lower():
                    entry_texts = [
                        "".join(node.itertext()).strip()
                        for node in sdt.findall(".//w:p", namespaces)
                    ]
                    if len([text for text in entry_texts if text]) >= 2:
                        return _ok("检测到含目录条目的 SDT 结构")

    return _fail("未检测到目录（TOC）")


def check_first_line_indent(file_path: str, params: dict) -> dict:
    """
    检查文档正文段落是否有首行缩进。

    输入:
        file_path: docx 文件路径
        params:
            min_indent_cm (float): 最小首行缩进值（厘米），默认 0.5
            threshold (float): 符合比例阈值，默认 0.8
            skip_styles (list[str]): 跳过检查的样式名称列表（如标题样式）
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    min_indent_cm = params.get("min_indent_cm", 0.5)
    threshold = params.get("threshold", 0.8)
    skip_styles = set(params.get("skip_styles", [
        "Heading 1", "Heading 2", "Heading 3", "Heading 4",
        "Title", "Subtitle", "TOC Heading",
    ]))

    min_indent_emu = Cm(min_indent_cm)

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    total_paras = 0
    indented_paras = 0

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        if para.style and para.style.name in skip_styles:
            continue

        total_paras += 1
        pf = para.paragraph_format
        first_indent = pf.first_line_indent

        if first_indent is not None and first_indent >= min_indent_emu:
            indented_paras += 1

    if total_paras == 0:
        return _ok("文档无需检查首行缩进的正文段落")

    ratio = indented_paras / total_paras
    if ratio >= threshold:
        return _ok(f"首行缩进率 {ratio:.1%}（{indented_paras}/{total_paras}）")

    return _partial(
        ratio,
        f"首行缩进率 {ratio:.1%}（{indented_paras}/{total_paras}），"
        f"期望 ≥ {threshold:.0%}"
    )


# ------------------------------------------------------------------
# 任务专用检查函数
# ------------------------------------------------------------------

def check_batchword002_tab_indent(file_path: str, params: dict) -> dict:
    """
    BatchoperationWord-002 专用：检查每个正文段落是否有标准首行缩进。

    判断方式（满足任一即视为有 tab 缩进）：
      1. 段落文本以 '\\t' 开头
      2. 段落中第一个 run 的文本以 '\\t' 开头
      3. 段落格式 first_line_indent 大于等于 min_indent_cm

    输入:
        file_path: docx 文件路径
        params:
            min_indent_cm (float): 标准首行缩进的最小厘米数，默认 0.5。
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    skip_styles = {
        "Heading 1", "Heading 2", "Heading 3", "Heading 4",
        "Title", "Subtitle", "TOC Heading",
    }
    min_indent = Cm(float(params.get("min_indent_cm", 0.5)))

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    total_paras = 0
    tab_paras = 0

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        if para.style and para.style.name in skip_styles:
            continue

        total_paras += 1

        # 检查段落文本或首个 run 是否以 tab 开头
        if para.text.startswith("\t"):
            tab_paras += 1
        elif para.runs and para.runs[0].text.startswith("\t"):
            tab_paras += 1
        elif (
            para.paragraph_format.first_line_indent is not None
            and para.paragraph_format.first_line_indent >= min_indent
        ):
            tab_paras += 1

    if total_paras == 0:
        return _ok("文档无需检查的正文段落")

    ratio = tab_paras / total_paras
    if ratio >= 0.8:
        return _ok(f"Tab 缩进率 {ratio:.1%}（{tab_paras}/{total_paras}）")

    return _partial(
        ratio,
        f"Tab 缩进率 {ratio:.1%}（{tab_paras}/{total_paras}）"
    )


# ------------------------------------------------------------------
# 扩展检查函数
# ------------------------------------------------------------------

def check_heading_style_exists(file_path: str, params: dict) -> dict:
    """
    检查文档是否存在指定样式的标题。

    输入:
        file_path: docx 文件路径
        params:
            style_name (str): 期望的标题样式名称，如 "Heading 1"
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    style_name = params.get("style_name", "")

    if not style_name:
        return _config_error("参数缺少 style_name")

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    found = False
    count = 0
    for para in doc.paragraphs:
        if para.style and para.style.name == style_name:
            found = True
            count += 1

    if found:
        return _ok(f"存在 {count} 个 '{style_name}' 样式段落")

    return _fail(f"未找到 '{style_name}' 样式的段落")


def check_has_table(file_path: str, params: dict) -> dict:
    """
    检查文档是否存在表格。

    输入:
        file_path: docx 文件路径
        params:
            min_tables (int): 最少表格数量，默认 1
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    min_tables = params.get("min_tables", 1)

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    table_count = len(doc.tables)

    if table_count >= min_tables:
        return _ok(f"文档包含 {table_count} 个表格")

    return _fail(f"文档仅有 {table_count} 个表格（期望至少 {min_tables} 个）")


def _normalize_table_value(value: Any) -> str:
    """归一化 Word 表格值，用于跨格式比较。

    输入:
        value: 表格单元格文本或任务配置中的期望值。
    输出:
        去空白、数字分组逗号并转小写后的字符串。
    """
    text = str(value).strip().casefold()
    text = re.sub(r"(?<=\d),(?=\d{3}\b)", "", text)
    return re.sub(r"\s+", "", text)


def check_table_contains_expected_values(file_path: str, params: dict) -> dict:
    """检查文档表格的结构与任务数据内容。

    输入:
        file_path: docx 文件路径。
        params:
            expected_values_by_file (dict[str, list]): 按文件名配置的必须出现值。
            min_rows (int): 单个表格最少行数，默认 1。
            min_columns (int): 单个表格最少列数，默认 1。
    输出:
        包含任务期望数据的表格存在时返回标准评价结果。
    """
    expected_by_file = params.get("expected_values_by_file", {})
    min_rows = int(params.get("min_rows", 1))
    min_columns = int(params.get("min_columns", 1))
    filename = os.path.basename(file_path)
    expected_values = expected_by_file.get(filename)
    if not expected_values:
        return _config_error(f"未配置 {filename} 的期望表格值")

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    qualifying_tables = [
        table for table in doc.tables
        if len(table.rows) >= min_rows and len(table.columns) >= min_columns
    ]
    if not qualifying_tables:
        return _fail(f"未找到至少 {min_rows}行×{min_columns}列的表格")

    table_values = {
        _normalize_table_value(cell.text)
        for table in qualifying_tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    }
    expected_normalized = [
        _normalize_table_value(value) for value in expected_values
    ]
    missing = [
        original for original, normalized in zip(expected_values, expected_normalized)
        if normalized not in table_values
    ]
    if not missing:
        return _ok(
            f"表格结构正确且包含全部 {len(expected_values)} 个期望值"
        )
    ratio = (len(expected_values) - len(missing)) / len(expected_values)
    return _partial(
        ratio,
        f"表格缺少 {len(missing)}/{len(expected_values)} 个期望值: "
        f"{', '.join(map(str, missing[:5]))}",
    )


def check_vowels_colored_red(file_path: str, params: dict) -> dict:
    """
    检查文档中元音字母（aeiouAEIOU）是否被标记为红色。

    输入:
        file_path: docx 文件路径
        params:
            threshold (float): 符合比例阈值，默认 0.8
            max_non_vowel_red_ratio (float): 辅音字母被误标红的最大比例，
                默认 0.05。
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    threshold = params.get("threshold", 0.8)
    max_non_vowel_red_ratio = params.get("max_non_vowel_red_ratio", 0.05)
    vowels = set("aeiouAEIOU")

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    total_vowels = 0
    red_vowels = 0
    total_non_vowels = 0
    red_non_vowels = 0

    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            if not text:
                continue
            for char in text:
                if char in vowels:
                    total_vowels += 1
                    if _is_red_run(run):
                        red_vowels += 1
                elif char.isalpha():
                    total_non_vowels += 1
                    if _is_red_run(run):
                        red_non_vowels += 1

    if total_vowels == 0:
        return _ok("文档无元音字母可检查")

    recall = red_vowels / total_vowels
    false_positive_ratio = (
        red_non_vowels / total_non_vowels if total_non_vowels else 0.0
    )
    specificity = 1.0 - false_positive_ratio
    score = min(recall, specificity)
    if recall >= threshold and false_positive_ratio <= max_non_vowel_red_ratio:
        return _ok(
            f"元音红色标记率 {recall:.1%}，"
            f"非元音误标红率 {false_positive_ratio:.1%}"
        )

    return _partial(
        score,
        f"元音红色率 {recall:.1%}，"
        f"非元音误标红率 {false_positive_ratio:.1%}"
    )


def check_uppercase_words_have_parentheses(file_path: str, params: dict) -> dict:
    """
    检查文档中纯大写的单词（如 MAC）后是否都有括号。

    输入:
        file_path: docx 文件路径
        params:
            threshold (float): 符合比例阈值，默认 0.9
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    threshold = params.get("threshold", 0.9)

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    import re
    uppercase_word_pattern = re.compile(r'\b[A-Z]{2,}\b')

    total_uppercase = 0
    with_parentheses = 0
    details = []

    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + " "

    matches = uppercase_word_pattern.findall(full_text)
    total_uppercase = len(matches)

    for para in doc.paragraphs:
        para_text = para.text
        for match in uppercase_word_pattern.finditer(para_text):
            word = match.group()
            end_pos = match.end()
            if re.match(r"\s*\(", para_text[end_pos:]):
                with_parentheses += 1
            else:
                details.append(word)

    if total_uppercase == 0:
        return _ok("文档无纯大写单词可检查")

    ratio = with_parentheses / total_uppercase
    if ratio >= threshold:
        return _ok(f"大写单词括号率 {ratio:.1%}（{with_parentheses}/{total_uppercase}）")

    return _partial(ratio, f"大写词括号率 {ratio:.1%}（{with_parentheses}/{total_uppercase}），无括号: {', '.join(details[:3])}")


def check_highlighted_words_capitalized(file_path: str, params: dict) -> dict:
    """
    检查黄色高亮的词是否有大写字母开头。

    输入:
        file_path: docx 文件路径
        params:
            highlight_color (str): 高亮颜色，默认 "FFFF00"（黄色）
            threshold (float): 符合比例阈值，默认 0.8
            require_bold (bool): 是否同时要求高亮文本加粗，默认 False。
            allowed_leading_articles (list[str]): 地名可接受的小写冠词，
                默认 ["the", "a", "an"]。
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    highlight_color = params.get("highlight_color", "FFFF00").upper()
    threshold = params.get("threshold", 0.8)
    require_bold = params.get("require_bold", False)
    allowed_articles = {
        article.lower() for article in params.get(
            "allowed_leading_articles", ["the", "a", "an"]
        )
    }

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    total_highlighted = 0
    capitalized_count = 0

    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            if not text:
                continue
            if run.font.highlight_color:
                hl = str(run.font.highlight_color).upper()
                if highlight_color in hl or "YELLOW" in hl or run.font.highlight_color == 7:
                    total_highlighted += 1
                    words = re.findall(r"[A-Za-z][A-Za-z'\-]*", text.strip())
                    if words and words[0].lower() in allowed_articles:
                        words = words[1:]
                    capitalized = bool(words and words[0][0].isupper())
                    correctly_bold = not require_bold or run.font.bold is True
                    if capitalized and correctly_bold:
                        capitalized_count += 1

    if total_highlighted == 0:
        return _fail("未找到黄色高亮文本")

    ratio = capitalized_count / total_highlighted
    if ratio >= threshold:
        return _ok(f"高亮词大写开头率 {ratio:.1%}（{capitalized_count}/{total_highlighted}）")

    return _partial(ratio, f"大写开头率 {ratio:.1%}（{capitalized_count}/{total_highlighted}）")


def check_misspelled_words_highlighted(file_path: str, params: dict) -> dict:
    """
    检查特定拼写错误的词是否被黄色高亮标记。

    输入:
        file_path: docx 文件路径
        params:
            expected_highlights (dict): 期望被高亮的词及其所在文档的映射
                例如: {"intrenational": "travel", "conmference": "travel"}
            highlight_color (str): 高亮颜色，默认 "FFFF00"
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    expected_highlights = params.get("expected_highlights", {})
    highlight_color = params.get("highlight_color", "FFFF00").upper()

    if not expected_highlights:
        return _config_error("参数缺少 expected_highlights")

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    total = len(expected_highlights)
    matched = 0
    details = []

    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text.lower()
            if not text:
                continue
            for misspelled, doc_key in expected_highlights.items():
                if misspelled.lower() in text:
                    if run.font.highlight_color:
                        hl = str(run.font.highlight_color).upper()
                        if highlight_color in hl or "YELLOW" in hl or run.font.highlight_color == 7:
                            matched += 1
                            details.append(f"'{misspelled}' 在 {doc_key} 中已高亮")
                        else:
                            details.append(f"'{misspelled}' 未高亮")
                    else:
                        details.append(f"'{misspelled}' 未高亮")

    if matched == total:
        return _ok(f"全部 {total} 个错误词都已黄色高亮")

    ratio = matched / total
    return _partial(ratio, f"{matched}/{total} 错误词已高亮: {', '.join(details[:5])}")


def check_heading_colors_different(file_path: str, params: dict) -> dict:
    """
    检查文档中不同标题的颜色是否互不相同。

    输入:
        file_path: docx 文件路径
        params:
            heading_styles (list[str]): 需要检查的标题样式列表
            threshold (float): 符合比例阈值，默认 0.8
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    heading_styles = params.get("heading_styles", ["Heading 1", "Heading 2", "Heading 3"])
    threshold = params.get("threshold", 0.8)

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    heading_colors = []

    for para in doc.paragraphs:
        if para.style and para.style.name in heading_styles:
            for run in para.runs:
                if run.text.strip():
                    if run.font.color and run.font.color.rgb:
                        color = str(run.font.color.rgb).upper()
                        heading_colors.append((para.style.name, run.text[:20], color))
                    else:
                        heading_colors.append((para.style.name, run.text[:20], "无颜色"))

    if len(heading_colors) < 2:
        return _fail("可检查的着色标题数量不足")

    colors_only = [c[2] for c in heading_colors]
    unique_colors = set(colors_only)

    ratio = len(unique_colors) / len(colors_only) if colors_only else 0.0

    if len(unique_colors) >= len(colors_only) * threshold:
        return _ok(f"全部 {len(unique_colors)} 个标题颜色都不同")

    return _partial(ratio, f"{len(unique_colors)} 种颜色 / {len(colors_only)} 个标题")


def check_heading_palette_and_references(result_dir: str, params: dict) -> dict:
    """检查多文档标题配色及互引文档的颜色约束。

    此检查以每个文档第一个非空段落为主标题，因为任务 fixture
    的标题是 Normal/Normal (Web) 加粗段落，而非 Heading 1..3 样式。

    输入:
        result_dir: Agent 产出根目录。
        params:
            expected_files (list[str]): 参与着色的 docx 文件名。
            palette (list[str]): 允许的六位 RGB 颜色列表。
    输出:
        所有标题均着色且互相引用的文档颜色不同时返回通过。
    """
    expected_files = params.get("expected_files", [])
    palette = {
        str(color).upper().lstrip("#")[-6:]
        for color in params.get(
            "palette", ["FF0000", "FFFF00", "0000FF", "00B050", "00FF00", "008000"]
        )
    }
    if not expected_files:
        return _config_error("参数缺少 expected_files")

    documents = {}
    for filename in expected_files:
        matches = sorted(
            path for path in glob.glob(os.path.join(result_dir, "**", filename), recursive=True)
            if os.path.isfile(path)
        )
        if len(matches) != 1:
            return _fail(f"{filename} 期望唯一文件，实际找到 {len(matches)} 个")
        document = _load_document(matches[0])
        if document is None:
            return _fail(f"无法打开文件: {filename}")
        documents[filename] = document

    colors = {}
    for filename, document in documents.items():
        title = next((paragraph for paragraph in document.paragraphs if paragraph.text.strip()), None)
        if title is None:
            return _fail(f"{filename} 无非空主标题")
        title_colors = set()
        for run in title.runs:
            if not run.text.strip():
                continue
            color = _run_rgb(run)
            if color is None and title.style is not None:
                style_color = title.style.font.color
                if style_color is not None and style_color.rgb is not None:
                    color = str(style_color.rgb).upper()[-6:]
            if color is None:
                return _fail(f"{filename} 主标题存在未显式着色文本")
            title_colors.add(color)
        if len(title_colors) != 1:
            return _fail(f"{filename} 主标题使用了 {len(title_colors)} 种颜色")
        color = next(iter(title_colors))
        if color not in palette:
            return _fail(f"{filename} 主标题颜色 {color} 不在允许色板中")
        colors[filename] = color

    references = []
    for source_name, document in documents.items():
        source_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        for target_name in expected_files:
            if target_name == source_name:
                continue
            target_stem = os.path.splitext(target_name)[0]
            if re.search(rf"\b{re.escape(target_stem)}\b", source_text, re.IGNORECASE):
                references.append((source_name, target_name))
                if colors[source_name] == colors[target_name]:
                    return _fail(
                        f"互引文档 {source_name} 与 {target_name} 颜色相同"
                    )
    if not references:
        return _fail("未在文档中识别到互引关系")
    return _ok(
        f"{len(colors)} 个文档主标题使用允许色板，"
        f"{len(references)} 条互引关系颜色均不同"
    )


def check_image_name_matches_doc(file_path: str, params: dict) -> dict:
    """
    检查文档中插入的图片是否为同名源图，并可校验宽度。

    Word/LibreOffice 会将嵌入媒体重命名为 image1.ext，因此不再比较
    DOCX 包内部件名，而是比较嵌入二进制与文档旁同名源图的 SHA-256。

    输入:
        file_path: docx 文件路径
        params:
            expected_width_cm (float|None): 期望图片宽度；默认不校验。
            width_tolerance_cm (float): 宽度绝对容差，默认 0.15cm。
            source_extensions (list[str]): 可接受的源图扩展名。
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    expected_width_cm = params.get("expected_width_cm")
    width_tolerance_cm = float(params.get("width_tolerance_cm", 0.15))
    source_extensions = params.get(
        "source_extensions", [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff"]
    )

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    doc_stem = os.path.splitext(os.path.basename(file_path))[0]
    source_candidates = []
    for extension in source_extensions:
        source_candidates.extend(
            glob.glob(os.path.join(os.path.dirname(file_path), doc_stem + extension))
        )
        source_candidates.extend(
            glob.glob(os.path.join(os.path.dirname(file_path), doc_stem + extension.upper()))
        )
    # 大小写不敏感的文件系统（macOS/Windows）上，小写与大写扩展名两次 glob 会
    # 返回同一文件的两种拼写，按路径字符串去重会误判为「多个同名源图」；
    # 改以 inode 标识去重。
    unique_by_inode = {}
    for candidate in sorted(set(source_candidates)):
        try:
            stat = os.stat(candidate)
        except OSError:
            continue
        unique_by_inode.setdefault((stat.st_dev, stat.st_ino), candidate)
    source_candidates = list(unique_by_inode.values())
    if len(source_candidates) != 1:
        return _fail(
            f"期望唯一同名源图，实际找到 {len(source_candidates)} 个"
        )

    with open(source_candidates[0], "rb") as stream:
        source_digest = hashlib.sha256(stream.read()).hexdigest()

    image_parts = []
    for relationship in doc.part.rels.values():
        if relationship.reltype.endswith("/image"):
            image_parts.append(relationship.target_part)

    if not image_parts:
        return _fail("文档中无内联图片")

    embedded_digests = {
        hashlib.sha256(part.blob).hexdigest() for part in image_parts
    }
    if source_digest not in embedded_digests:
        return _fail("嵌入图片与文档同名源图内容不一致")

    if expected_width_cm is not None:
        namespaces = {
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        }
        expected_emu = float(Cm(float(expected_width_cm)))
        tolerance_emu = float(Cm(width_tolerance_cm))
        extents = doc.element.body.findall(".//wp:extent", namespaces)
        widths = [float(extent.get("cx", "0")) for extent in extents]
        if not any(abs(width - expected_emu) <= tolerance_emu for width in widths):
            return _fail(
                f"未找到宽度约 {float(expected_width_cm):g}cm 的匹配图片"
            )

    return _ok(f"已嵌入同名源图 {os.path.basename(source_candidates[0])}")


def check_docx_word_count(file_path: str, params: dict) -> dict:
    """
    检查Word文档的字数是否大于指定值。

    输入:
        file_path: docx 文件路径
        params:
            min_words (int): 最少字数要求，默认 100
            exclude_headings (bool): 是否排除标题/Title 样式段落的词数，
                默认 False（保持历史行为）。开放写作任务建议开启，
                避免模板自带的标题词数白送基础分。
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    min_words = params.get("min_words", 100)
    exclude_headings = params.get("exclude_headings", False)

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    word_count = 0
    for para in doc.paragraphs:
        style_name = (para.style.name or "") if para.style is not None else ""
        if exclude_headings and (
            style_name.startswith("Heading") or style_name == "Title"
        ):
            continue
        words = para.text.split()
        word_count += len(words)

    scope = "正文" if exclude_headings else "文档"
    if word_count >= min_words:
        return _ok(f"{scope}字数 {word_count} >= {min_words}")

    ratio = word_count / min_words
    return _partial(ratio, f"{scope}字数 {word_count} < {min_words}")


def check_headings_have_body(file_path: str, params: dict) -> dict:
    """
    检查文档中每个标题（Heading 样式）之后是否有实质正文。

    用于"补全各小节正文"类开放写作任务：仅靠总词数无法约束内容写在
    哪里——agent 把全部文字堆在一个小节下、或写在标题外的任意位置都
    能通过词数检查。本检查按小节统计：每个 Heading 与下一个 Heading
    （或文档结尾）之间的非标题段落计为该小节正文，
    得分 = 正文词数达标的小节数 / 小节总数。

    输入:
        file_path: docx 文件路径
        params:
            min_words_per_heading (int): 每个小节正文最少词数，默认 15
            heading_style_prefix (str): 标题样式前缀，默认 "Heading"
                （Title 样式视为文档题目，不计入小节）
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    min_words = params.get("min_words_per_heading", 15)
    heading_prefix = params.get("heading_style_prefix", "Heading")

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    # 按段落顺序切分小节：遇到 Heading 开新节，其后的非标题段落归入该节
    sections = []  # [(heading_text, body_word_count)]
    current = None  # 当前小节的 [标题文本, 正文词数]
    for para in doc.paragraphs:
        style_name = (para.style.name or "") if para.style is not None else ""
        if style_name.startswith(heading_prefix):
            if current is not None:
                sections.append(tuple(current))
            current = [para.text.strip(), 0]
        elif style_name != "Title" and current is not None:
            current[1] += len(para.text.split())
    if current is not None:
        sections.append(tuple(current))

    if not sections:
        return _fail(f"文档中没有 '{heading_prefix}' 样式的标题（可能被删除或样式丢失）")

    satisfied = [s for s in sections if s[1] >= min_words]
    ratio = len(satisfied) / len(sections)
    lacking = [f"'{h[:30]}'({w}词)" for h, w in sections if w < min_words]

    if ratio >= 1.0:
        return _ok(f"全部 {len(sections)} 个小节均有 >= {min_words} 词正文")
    return _partial(
        ratio,
        f"{len(satisfied)}/{len(sections)} 个小节正文达标，"
        f"不足的小节: {', '.join(lacking[:5])}",
    )


def check_docx_has_hyperlink(file_path: str, params: dict) -> dict:
    """检查 Word 文档的超链接数量、提示段落覆盖率与目标。

    输入:
        file_path: docx 文件路径。
        params:
            threshold (float): 兼容旧规则的链接数/段落数阈值，默认 0.5。
            cue_pattern (str|None): 需要覆盖的提示段落正则；默认识别
                "see ... .pptx Page N" 形式。
            require_cue_coverage (bool): 是否要求每个提示段落都有链接，
                默认 False。
            validate_cited_target (bool): 是否要求链接目标文件名与段落
                引用的 pptx 名称一致，默认 False。
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    threshold = params.get("threshold", 0.5)
    require_cue_coverage = params.get("require_cue_coverage", False)
    validate_cited_target = params.get("validate_cited_target", False)
    cue_pattern = re.compile(
        params.get(
            "cue_pattern",
            r"(?P<target>[A-Za-z0-9_.\-]+\.pptx)\s+Page\s+(?P<page>\d+)",
        ),
        re.IGNORECASE,
    )

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    body = doc.element.body
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    hyperlinks = body.findall(".//w:hyperlink", namespaces)
    link_count = len(hyperlinks)
    if link_count == 0:
        return _fail("文档中不存在超链接")

    if require_cue_coverage:
        cue_count = 0
        covered_count = 0
        failures = []
        for paragraph in body.findall(".//w:p", namespaces):
            paragraph_text = "".join(
                node.text or "" for node in paragraph.findall(".//w:t", namespaces)
            )
            citation = cue_pattern.search(paragraph_text)
            if citation is None:
                continue
            cue_count += 1
            paragraph_links = paragraph.findall(".//w:hyperlink", namespaces)
            targets = []
            for hyperlink in paragraph_links:
                relationship_id = hyperlink.get(f"{{{namespaces['r']}}}id")
                if relationship_id and relationship_id in doc.part.rels:
                    targets.append(str(doc.part.rels[relationship_id].target_ref))
            cited_name = os.path.basename(citation.groupdict().get("target", "").strip())
            target_matches = any(
                os.path.basename(target.split("#", 1)[0]).casefold() == cited_name.casefold()
                for target in targets
            )
            if paragraph_links and (not validate_cited_target or target_matches):
                covered_count += 1
            else:
                failures.append(cited_name or paragraph_text[:40])
        if cue_count == 0:
            return _fail("未识别到需要插入链接的 PPT 页面提示")
        if covered_count == cue_count:
            return _ok(f"全部 {cue_count} 个 PPT 页面提示均有正确文件链接")
        return _partial(
            covered_count / cue_count,
            f"仅覆盖 {covered_count}/{cue_count} 个提示段落，"
            f"缺失或目标错误: {', '.join(failures[:3])}",
        )

    total_paras = len([paragraph for paragraph in doc.paragraphs if paragraph.text.strip()])
    if total_paras == 0:
        return _ok(f"文档无段落，但包含 {link_count} 个超链接")
    ratio = min(1.0, link_count / total_paras)
    if ratio >= threshold:
        return _ok(f"文档包含 {link_count} 个超链接")
    return _partial(ratio, f"仅 {link_count} 个超链接")


# ---- 以下三项自 2026-07 审计修复线（docs/eval_audit/ 修复包）补回 ----
# 线 C 的 Word 检查未覆盖高亮加粗、图片宽度与超链接目标三类断言，
# 分别为 BatchOperationWord-003 / -010 与 CombinationDocs-006 所引用。


def _run_has_highlight(run, highlight_color: str = "FFFF00") -> bool:
    color = run.font.highlight_color
    if not color:
        return False
    color_text = str(color).upper()
    color_value = getattr(color, "value", None)
    return (
        highlight_color.upper() in color_text
        or "YELLOW" in color_text
        or color_value == 7
        or color == 7
    )


def check_highlighted_words_bold(file_path: str, params: dict) -> dict:
    """
    检查被黄色高亮的文本是否同时加粗。

    用于“highlight ... in yellow bold text”类任务；只检查已经高亮的
    文本是否满足 bold，避免把未被识别的 gold 实体硬编码到 evaluator。
    """
    highlight_color = params.get("highlight_color", "FFFF00")
    threshold = params.get("threshold", 0.8)

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    total_highlighted = 0
    bold_highlighted = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if not (run.text or "").strip():
                continue
            if not _run_has_highlight(run, highlight_color):
                continue
            total_highlighted += 1
            style_bold = para.style.font.bold if para.style and para.style.font else None
            if run.font.bold is True or (run.font.bold is None and style_bold is True):
                bold_highlighted += 1

    if total_highlighted == 0:
        return _fail("未找到黄色高亮文本")

    ratio = bold_highlighted / total_highlighted
    if ratio >= threshold:
        return _ok(f"高亮文本加粗率 {ratio:.1%}（{bold_highlighted}/{total_highlighted}）")
    return _partial(ratio, f"高亮文本加粗率 {ratio:.1%}（{bold_highlighted}/{total_highlighted}）")


def check_docx_image_width(file_path: str, params: dict) -> dict:
    """
    检查文档中内联图片宽度是否接近期望厘米数。
    """
    width_cm = float(params.get("width_cm", 5.0))
    tolerance_cm = float(params.get("tolerance_cm", 0.15))
    threshold = params.get("threshold", 1.0)
    expected_width = Cm(width_cm)
    tolerance = Cm(tolerance_cm)

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    total = len(doc.inline_shapes)
    if total == 0:
        return _fail("文档中无内联图片")

    matched = 0
    details = []
    for idx, shape in enumerate(doc.inline_shapes, 1):
        actual = shape.width
        if abs(actual - expected_width) <= tolerance:
            matched += 1
        else:
            actual_cm = actual / Cm(1)
            details.append(f"图{idx}: {actual_cm:.2f}cm")

    ratio = matched / total
    if ratio >= threshold:
        return _ok(f"图片宽度匹配率 {ratio:.1%}（{matched}/{total}），期望 {width_cm:g}cm")
    return _partial(ratio, f"图片宽度匹配率 {ratio:.1%}，不匹配: {', '.join(details[:5])}")


def check_docx_hyperlink_targets(file_path: str, params: dict) -> dict:
    """
    检查 Word 超链接目标是否指向指定文件类型。
    """
    expected_extensions = [
        ext.lower() if str(ext).startswith(".") else f".{str(ext).lower()}"
        for ext in params.get("expected_extensions", [".pptx", ".ppt"])
    ]
    min_links = int(params.get("min_links", 1))

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    targets = []
    for hyperlink in doc.element.body.findall(".//w:hyperlink", namespaces):
        rid = hyperlink.get(f"{{{namespaces['r']}}}id")
        if rid and rid in doc.part.rels:
            rel = doc.part.rels[rid]
            targets.append(str(rel.target_ref))

    if not targets:
        return _fail("文档中不存在可解析目标的超链接")

    matched = [
        target for target in targets
        if any(ext in target.lower() for ext in expected_extensions)
    ]
    if len(matched) >= min_links:
        return _ok(f"{len(matched)}/{len(targets)} 个超链接目标指向 {expected_extensions}")
    return _partial(
        len(matched) / max(min_links, len(targets)),
        f"仅 {len(matched)} 个超链接目标指向 {expected_extensions}: {targets[:5]}",
    )


# 线 C 以 check_heading_palette_and_references 覆盖了标题配色与参考文献两项，
# 但审计修复线的调色板独立断言仍被 tests/test_operation_coverage_regressions.py
# 引用，一并保留以免回退该断言。


def check_heading_colors_in_palette(file_path: str, params: dict) -> dict:
    """
    检查标题颜色是否限定在指定调色板内。
    """
    heading_styles = params.get("heading_styles", ["Heading 1", "Heading 2", "Heading 3"])
    allowed_colors = {str(c).upper().replace("#", "") for c in params.get(
        "allowed_colors", ["FF0000", "FFFF00", "0000FF", "00FF00"]
    )}
    threshold = params.get("threshold", 1.0)
    tolerance = float(params.get("rgb_tolerance", 60.0))

    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")

    def _parse_rgb(rgb: str):
        if not rgb:
            return None
        if len(rgb) == 8:
            rgb = rgb[2:]
        if len(rgb) != 6:
            return None
        try:
            return (int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16))
        except ValueError:
            return None

    def _color_family(rgb: str) -> Optional[str]:
        parsed = _parse_rgb(rgb)
        if parsed is None:
            return None
        import colorsys
        r, g, b = parsed
        h, s, v = colorsys.rgb_to_hsv(r / 255, g / 255, b / 255)
        hue = h * 360
        if (hue <= 15 or hue >= 345) and s >= 0.55 and v >= 0.5:
            return "red"
        if 45 <= hue <= 70 and s >= 0.45 and v >= 0.5:
            return "yellow"
        if 200 <= hue <= 260 and s >= 0.45 and v >= 0.35:
            return "blue"
        if 80 <= hue <= 165 and s >= 0.35 and v >= 0.3:
            return "green"
        return None

    def _rgb_distance(rgb1: str, rgb2: str) -> float:
        parsed1 = _parse_rgb(rgb1)
        parsed2 = _parse_rgb(rgb2)
        if parsed1 is None or parsed2 is None:
            return float("inf")
        return sum((a - b) ** 2 for a, b in zip(parsed1, parsed2)) ** 0.5

    def _matches_allowed_color(rgb: str) -> bool:
        normalized_rgb = rgb[2:] if len(rgb) == 8 else rgb
        if normalized_rgb in allowed_colors:
            return True
        if any(_rgb_distance(normalized_rgb, allowed) <= tolerance for allowed in allowed_colors):
            return True
        family = _color_family(normalized_rgb)
        return family is not None and family in allowed_families

    allowed_families = {
        family for color in allowed_colors
        if (family := _color_family(color))
    }

    total = 0
    matched = 0
    details = []
    for para in doc.paragraphs:
        if not (para.style and para.style.name in heading_styles):
            continue
        for run in para.runs:
            if not run.text.strip():
                continue
            total += 1
            rgb = str(run.font.color.rgb).upper() if run.font.color and run.font.color.rgb else ""
            normalized_rgb = rgb[2:] if len(rgb) == 8 else rgb
            if _matches_allowed_color(rgb):
                matched += 1
            else:
                details.append(f"{run.text[:20]}={normalized_rgb or '无颜色'}")

    if total == 0:
        return _fail("未找到可检查颜色的标题文本")
    ratio = matched / total
    if ratio >= threshold:
        return _ok(f"标题颜色调色板匹配率 {ratio:.1%}（{matched}/{total}）")
    return _partial(ratio, f"标题颜色调色板匹配率 {ratio:.1%}，不匹配: {', '.join(details[:5])}")



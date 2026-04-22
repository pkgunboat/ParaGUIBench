"""
Operation 文件操作任务 Pipeline — 多线程并行版本。

基于 run_QA_pipeline_parallel.py，Stage 1/2 完全复用，Stage 3 评估逻辑
从 QA 的文本匹配改为文件内容比对（.docx / .xlsx / .pptx）。

用法:
    # 顺序执行（默认）
    python run_self_operation_pipeline_parallel.py

    # 2 个任务并行，每任务 3 个 VM
    python run_self_operation_pipeline_parallel.py -p 2 -n 3

    # 指定 GT 缓存目录
    python run_self_operation_pipeline_parallel.py --gt-cache-dir /tmp/my_gt_cache
"""

from __future__ import annotations

import argparse
import atexit
import json
import logging
import os
import queue
import shutil
import subprocess
import sys
import tempfile
import threading
import time
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any, Dict, List, Tuple

import requests

# ============================================================
# 路径设置
# ============================================================

current_dir = os.path.dirname(os.path.abspath(__file__))
examples_dir = os.path.dirname(current_dir)         # ubuntu_env/examples/
ubuntu_env_dir = os.path.dirname(examples_dir)       # ubuntu_env/
parallel_benchmark_dir = os.path.join(ubuntu_env_dir, "parallel_benchmark")

if parallel_benchmark_dir not in sys.path:
    sys.path.insert(0, parallel_benchmark_dir)
if ubuntu_env_dir not in sys.path:
    sys.path.insert(0, ubuntu_env_dir)
if examples_dir not in sys.path:
    sys.path.insert(0, examples_dir)

# ============================================================
# 从原始 pipeline 导入可复用函数（不修改原文件）
# ============================================================

from run_QA_pipeline import (  # noqa: E402
    ensure_conda_env,
    load_task_config,
    load_evaluator,
    parse_prepare_script_path,
    extract_execution_summary,
    TASKS_LIST_DIR,
)

# ============================================================
# 从 QA parallel pipeline 导入可复用函数
# ============================================================

from run_QA_pipeline_parallel import (  # noqa: E402
    execute_on_vm_with_ip,
    wait_for_vm_ready_with_ip,
    _download_task_files_on_vm_with_ip,
    get_ssh_credentials,
    run_ssh_command,
    setup_logging,
    get_task_logger,
    rebuild_containers_parallel,
    cleanup_group_containers,
    stage2_execute_agent_parallel,
)

from parallel_agents.plan_agent_thought_action import (  # noqa: E402
    calculate_cost,
)

# ============================================================
# 从 Docker 并行管理器导入
# ============================================================

from desktop_env.providers.docker.parallel_manager import (  # noqa: E402
    ContainerSetConfig,
    MemoryGuard,
    allocate_ports_for_group,
    scan_remote_docker_ports,
)

# ============================================================
# 评估器 metrics 导入
# ============================================================

import openpyxl  # noqa: E402

from desktop_env.evaluators.metrics.comprehensive import (  # noqa: E402
    compare_file_comprehensive,
)

from desktop_env.controllers.python import PythonController  # noqa: E402
from parallel_agents_as_tools.seed18_gui_agent_as_tool import Seed18GUIAgentTool  # noqa: E402
from parallel_agents_as_tools.claude_gui_agent_as_tool import ClaudeGUIAgentTool  # noqa: E402
from parallel_agents_as_tools.kimi_gui_agent_as_tool import KimiGUIAgentTool  # noqa: E402

# ============================================================
# 常量
# ============================================================

OUTPUT_JSON_PATH = os.path.join(
    ubuntu_env_dir, "logs", "run_self_operation_pipeline_parallel.json"
)

# 全局追踪：记录所有已启动的容器组（用于 atexit 清理）
_active_groups: Dict[int, ContainerSetConfig] = {}
_active_groups_lock = threading.Lock()

# ============================================================
# 活跃端口注册表（用于 GlobalScreensaverHeartbeat 动态端口列表）
# ============================================================

_active_ports: Dict[int, List[int]] = {}  # group_id -> [server_port, ...]
_active_ports_lock = threading.Lock()


def register_group_ports(group_id: int, server_ports: List[int]) -> None:
    """
    注册某组的 VM server 端口到全局活跃端口表。

    输入:
        group_id: 容器组编号
        server_ports: 该组所有 VM 的 server 端口列表
    """
    with _active_ports_lock:
        _active_ports[group_id] = list(server_ports)


def unregister_group_ports(group_id: int) -> None:
    """
    从全局活跃端口表中注销某组的端口。

    输入:
        group_id: 容器组编号
    """
    with _active_ports_lock:
        _active_ports.pop(group_id, None)


def get_all_active_ports() -> List[int]:
    """
    获取所有活跃组的 server 端口（扁平化列表）。

    输出:
        所有活跃 VM 的 server 端口列表
    """
    with _active_ports_lock:
        ports = []
        for port_list in _active_ports.values():
            ports.extend(port_list)
        return ports


# ============================================================
# GlobalScreensaverHeartbeat — 支持动态端口列表
# ============================================================

class GlobalScreensaverHeartbeat:
    """
    全局防黑屏心跳守护线程，支持动态端口列表。

    与单次禁用屏保的区别：
    - 端口列表不在初始化时固定，而是每次心跳时从 get_all_active_ports() 动态获取
    - 适用于多任务并行场景：不同组的 VM 在不同时间启动和关闭

    输入：
        vm_ip: VM 宿主 IP
        interval_sec: 心跳间隔（秒），默认 180（3 分钟）
    """

    def __init__(self, vm_ip: str, interval_sec: int = 180):
        self.vm_ip = vm_ip
        self.interval_sec = interval_sec
        self._stop_event = threading.Event()
        self._thread = None

    def _heartbeat_loop(self) -> None:
        """
        心跳循环主体，在后台线程中运行。
        每隔 interval_sec 秒向所有活跃 VM 发送屏保重置命令。
        """
        heartbeat_script = (
            "import subprocess, os\n"
            "env = os.environ.copy()\n"
            "env['DISPLAY'] = ':0'\n"
            "env['DBUS_SESSION_BUS_ADDRESS'] = 'unix:path=/run/user/1000/bus'\n"
            "try:\n"
            "    subprocess.run(['dbus-send', '--session',\n"
            "        '--dest=org.gnome.ScreenSaver', '--type=method_call',\n"
            "        '/org/gnome/ScreenSaver',\n"
            "        'org.gnome.ScreenSaver.SetActive', 'boolean:false'],\n"
            "        env=env, capture_output=True, timeout=5)\n"
            "    subprocess.run(['xset', 's', 'reset'],\n"
            "        env=env, capture_output=True, timeout=5)\n"
            "except Exception:\n"
            "    pass\n"
            "print('heartbeat_ok')\n"
        )

        log = logging.getLogger("pipeline.heartbeat")

        while not self._stop_event.is_set():
            if self._stop_event.wait(timeout=self.interval_sec):
                break

            # 动态获取当前活跃的端口列表
            ports = get_all_active_ports()
            if not ports:
                continue

            log.debug("心跳: 向 %d 个 VM 发送屏保重置", len(ports))
            for port in ports:
                try:
                    url = f"http://{self.vm_ip}:{port}/execute"
                    payload = json.dumps({
                        "command": ["python", "-c", heartbeat_script],
                        "shell": False,
                    })
                    requests.post(
                        url,
                        headers={"Content-Type": "application/json"},
                        data=payload,
                        timeout=10,
                    )
                except Exception:
                    pass  # 静默忽略，不影响主流程

    def start(self) -> None:
        """启动心跳守护线程。"""
        if self._thread is not None and self._thread.is_alive():
            return
        self._stop_event.clear()
        self._thread = threading.Thread(
            target=self._heartbeat_loop,
            name="global-screensaver-heartbeat",
            daemon=True,
        )
        self._thread.start()
        logging.getLogger("pipeline.heartbeat").info(
            "GlobalScreensaverHeartbeat 已启动（间隔 %ds）", self.interval_sec
        )

    def stop(self) -> None:
        """停止心跳守护线程。"""
        self._stop_event.set()
        if self._thread is not None:
            self._thread.join(timeout=5)
            self._thread = None
        logging.getLogger("pipeline.heartbeat").info("GlobalScreensaverHeartbeat 已停止")


# ============================================================
# 批量禁用屏保（Stage 1 使用）
# ============================================================

def disable_screensaver_parallel(
    vm_ip: str,
    vm_ports: List[int],
    log: logging.Logger,
) -> None:
    """
    在指定端口的所有 VM 中禁用屏保和锁屏（防黑屏第一层：预防）。

    输入:
        vm_ip: VM 宿主 IP
        vm_ports: VM server 端口列表
        log: logger
    """
    log.info("禁用所有 VM 的屏保和锁屏...")
    disable_script = (
        "import subprocess, os\n"
        "env = os.environ.copy()\n"
        "env['DISPLAY'] = ':0'\n"
        "env['DBUS_SESSION_BUS_ADDRESS'] = 'unix:path=/run/user/1000/bus'\n"
        "cmds = [\n"
        "    ['gsettings', 'set', 'org.gnome.desktop.session', 'idle-delay', '0'],\n"
        "    ['gsettings', 'set', 'org.gnome.desktop.screensaver', 'lock-enabled', 'false'],\n"
        "    ['gsettings', 'set', 'org.gnome.desktop.screensaver', 'idle-activation-enabled', 'false'],\n"
        "]\n"
        "for cmd in cmds:\n"
        "    try:\n"
        "        subprocess.run(cmd, env=env, capture_output=True, timeout=5)\n"
        "    except Exception:\n"
        "        pass\n"
        "try:\n"
        "    subprocess.run(['xset', 's', 'off'], env=env, capture_output=True, timeout=5)\n"
        "    subprocess.run(['xset', '-dpms'], env=env, capture_output=True, timeout=5)\n"
        "    subprocess.run(['xset', 's', 'noblank'], env=env, capture_output=True, timeout=5)\n"
        "except Exception:\n"
        "    pass\n"
        "print('screensaver_disabled')\n"
    )

    for port in vm_ports:
        try:
            url = f"http://{vm_ip}:{port}/execute"
            payload = json.dumps({
                "command": ["python", "-c", disable_script],
                "shell": False,
            })
            resp = requests.post(
                url,
                headers={"Content-Type": "application/json"},
                data=payload,
                timeout=15,
            )
            if resp.status_code == 200:
                output = resp.json().get("output", "")
                if "screensaver_disabled" in output:
                    log.info("  VM %d 屏保已禁用", port)
                else:
                    log.warning("  VM %d 屏保禁用返回异常: %s", port, output[:100])
            else:
                log.warning("  VM %d 屏保禁用失败 (HTTP %d)", port, resp.status_code)
        except Exception as exc:
            log.warning("  VM %d 屏保禁用失败: %s", port, exc)


# HuggingFace 数据集基础 API URL
HF_API_BASE = "https://huggingface.co/api/datasets/leeLegendary/Parallel_benchmark"
HF_DOWNLOAD_BASE = "https://huggingface.co/datasets/leeLegendary/Parallel_benchmark/resolve/main"


# ============================================================
# 任务级自定义文件比对函数
# ============================================================


def _is_red_family(argb: str) -> bool:
    """
    判断 aRGB 颜色是否属于红色系（R>180, G<100, B<100）。

    输入:
        argb: 8 位 aRGB 字符串（如 "FFC9211E"、"FFFF0000"）

    输出:
        bool
    """
    if not argb or len(argb) < 6:
        return False
    try:
        hex_rgb = argb[-6:]  # 取后 6 位（去掉 alpha）
        r = int(hex_rgb[0:2], 16)
        g = int(hex_rgb[2:4], 16)
        b = int(hex_rgb[4:6], 16)
        return r > 180 and g < 100 and b < 100
    except (ValueError, IndexError):
        return False


def _extract_red_indicator(cell) -> bool:
    """
    检查单元格是否被标为红色（同时检查字体颜色和填充颜色）。

    Agent 可能通过两种方式标红：
      1. 字体颜色（font.color）— 红色字体
      2. 填充前景色（fill.fgColor / fill.start_color）— 红色背景

    仅检查 RGB 类型的颜色值（theme 类型的 .rgb 属性会返回错误字符串而非真实颜色）。

    输入:
        cell: openpyxl Cell 对象

    输出:
        bool，是否被标为红色
    """
    # 检查字体颜色
    try:
        fc = cell.font.color
        if fc and getattr(fc, 'type', None) == 'rgb' and fc.rgb:
            if _is_red_family(str(fc.rgb)):
                return True
    except (AttributeError, TypeError):
        pass

    # 检查填充前景色（fgColor 和 start_color 两个属性）
    try:
        fill = cell.fill
        if fill:
            for color_attr in ('fgColor', 'start_color'):
                color = getattr(fill, color_attr, None)
                if color and getattr(color, 'type', None) == 'rgb' and color.rgb:
                    if _is_red_family(str(color.rgb)):
                        return True
    except (AttributeError, TypeError):
        pass

    return False


def compare_xlsx_negative_profit_red(gt_path: str, result_path: str) -> float:
    """
    excel-006 专用评估：检查负利润行是否被标为红色。

    评估逻辑:
      1. 从 GT 确定哪些行是负利润行（利润值 < 0）
      2. 在 Agent 结果中检查这些行是否被标红（字体颜色或填充颜色）
      3. 同时检查非负利润行是否未被标红（防止全部标红拿分）

    输入:
        gt_path: GT xlsx 文件路径（用于确定负利润行）
        result_path: Agent 结果 xlsx 文件路径

    输出:
        float，0.0~1.0
    """
    _log = logging.getLogger("pipeline.excel006")

    try:
        gt_wb = openpyxl.load_workbook(gt_path)
        result_wb = openpyxl.load_workbook(result_path)
    except Exception as exc:
        _log.error("[excel-006] 无法打开文件: %s", exc)
        return 0.0

    gt_ws = gt_wb.active
    result_ws = result_wb.active

    # 找利润列（表头行，搜索前 10 行以兼容不同表头位置）
    profit_col = None
    header_row = None
    for row in range(1, min(11, gt_ws.max_row + 1)):
        for col in range(1, gt_ws.max_column + 1):
            header = str(gt_ws.cell(row, col).value or "")
            if "Profit" in header or "利润" in header:
                profit_col = col
                header_row = row
                break
        if profit_col:
            break

    if not profit_col:
        _log.warning("[excel-006] 未找到利润列")
        return 0.0

    data_start_row = header_row + 1

    # 遍历数据行
    correct = 0
    total = 0
    for row in range(data_start_row, gt_ws.max_row + 1):
        profit = gt_ws.cell(row, profit_col).value
        if profit is None or not isinstance(profit, (int, float)):
            continue
        # 跳过汇总行
        first_cell_val = str(gt_ws.cell(row, 1).value or "").upper()
        if first_cell_val in ("TOTAL", "合计", "年度总计", "总计"):
            continue

        total += 1
        is_negative = profit < 0

        # 检查结果文件中该行所有列（字体颜色 + 填充颜色），任一列为红色即算标红
        is_red = False
        for col in range(1, result_ws.max_column + 1):
            if _extract_red_indicator(result_ws.cell(row, col)):
                is_red = True
                break

        if is_negative and is_red:
            correct += 1      # 负利润行被标红 ✓
        elif not is_negative and not is_red:
            correct += 1      # 非负利润行未标红 ✓

    gt_wb.close()
    result_wb.close()

    if total == 0:
        _log.warning("[excel-006] 没有有效数据行")
        return 0.0

    score = correct / total
    _log.info("[excel-006] 红色高亮评估: %d/%d correct, score=%.2f", correct, total, score)
    return score


def compare_pptx_first_slide_transition(gt_path: str, result_path: str) -> float:
    """
    ppt-001 专用评估：检查第一张幻灯片是否设置了 dissolve 过渡效果。

    ppt-001 的 GT 未包含 transition（制作疏漏），导致默认评估器跳过 transition 检查。
    本函数直接检查 Agent 结果文件：
      - 第一张幻灯片必须有 'dissolve' transition → 得 1.0
      - 否则 → 得 0.0

    输入:
        gt_path: GT pptx 文件路径（本评估器不使用 GT 内容）
        result_path: Agent 结果 pptx 文件路径

    输出:
        float，0.0 或 1.0
    """
    from desktop_env.evaluators.metrics.slides import extract_all_transitions

    _log = logging.getLogger("pipeline.ppt001")
    try:
        result_transitions = extract_all_transitions(result_path)
    except Exception as exc:
        _log.error("[ppt-001] 提取 transition 失败: %s", exc)
        return 0.0

    if not result_transitions:
        _log.info("[ppt-001] 结果文件无幻灯片")
        return 0.0

    first_transition = result_transitions[0]
    if first_transition and first_transition.lower() == "dissolve":
        _log.info("[ppt-001] ✓ 第一张幻灯片有 dissolve transition")
        return 1.0

    _log.info("[ppt-001] ✗ 第一张幻灯片 transition=%s（期望 dissolve）",
              first_transition)
    return 0.0


def compare_docx_style_focused(gt_path: str, result_path: str) -> float:
    """
    Word-001 专用评估：侧重样式修改的 docx 比对。

    Word-001 任务要求修复标题层级（Normal Web → Heading 1/2），文本内容不变。
    默认权重 text=0.4 会给未修改的文件很高的基线分（0.6+）导致假阳性。
    本函数使用侧重样式的权重分配，确保未修改文件得分低于 pass 阈值。

    权重:
      - text（0.1）：段落文本（任务不要求改文本，低权重）
      - table（0.1）：表格数据（任务不涉及表格）
      - style（0.45）：段落样式名称（核心检查项，高权重）
      - run_format（0.35）：run 级别格式

    输入:
        gt_path: GT docx 文件路径
        result_path: 待评估 docx 文件路径

    输出:
        float，0.0~1.0
    """
    from desktop_env.evaluators.metrics.comprehensive import compare_docx_comprehensive

    weights = {"text": 0.1, "table": 0.1, "style": 0.45, "run_format": 0.35}
    return compare_docx_comprehensive(gt_path, result_path, weights=weights)


def compare_docx_vowels_red(gt_path: str, result_path: str) -> float:
    """
    Word-007 专用评估：检查所有元音字母是否被标为红色。

    评估逻辑:
      1. 用 compare_docx_files 检查文本内容是否一致（基线验证）
      2. 遍历结果文件所有 run，逐字符提取颜色信息
      3. 对每个字母字符分类检查：
         - 元音字符(a,e,i,o,u) → 应为红色（R>180, G<100, B<100）
         - 非元音字母 → 应不为红色（防止全部标红拿分）
      4. 得分 = 0.3 * text_score + 0.7 * classification_accuracy

    输入:
        gt_path: GT docx 文件路径
        result_path: 待评估 docx 文件路径

    输出:
        float，0.0~1.0
    """
    from docx import Document as DocxDocument
    from desktop_env.evaluators.metrics.docs import compare_docx_files

    _log = logging.getLogger("pipeline.word007")
    VOWELS = set("aeiouAEIOU")

    # ── 第 1 步：文本内容一致性 ──
    try:
        text_score = float(compare_docx_files(result_path, gt_path))
    except Exception as exc:
        _log.error("[word-007] compare_docx_files 异常: %s", exc)
        text_score = 0.0

    # ── 第 2 步：逐字符颜色分类检查 ──
    try:
        doc = DocxDocument(result_path)
    except Exception as exc:
        _log.error("[word-007] 无法打开结果文件: %s", exc)
        return 0.0

    total = 0
    correct = 0

    for para in doc.paragraphs:
        for run in para.runs:
            # 获取该 run 的字体颜色
            run_is_red = False
            try:
                color_obj = run.font.color
                if color_obj and color_obj.rgb:
                    argb = str(color_obj.rgb)
                    run_is_red = _is_red_family(argb)
            except (AttributeError, TypeError):
                pass

            # 逐字符检查（同一个 run 内所有字符颜色相同）
            for ch in run.text:
                if not ch.isalpha():
                    continue  # 跳过空白、标点、数字
                total += 1
                is_vowel = ch in VOWELS
                if is_vowel and run_is_red:
                    correct += 1   # 元音被标红 ✓
                elif not is_vowel and not run_is_red:
                    correct += 1   # 非元音未标红 ✓

    if total == 0:
        _log.warning("[word-007] 未找到字母字符")
        return 0.3 * text_score

    classification_acc = correct / total
    score = 0.3 * text_score + 0.7 * classification_acc
    _log.info("[word-007] text=%.2f, class_acc=%.2f (%d/%d), final=%.2f",
              text_score, classification_acc, correct, total, score)
    return score


def compare_docx_font_times_new_roman(gt_path: str, result_path: str) -> float:
    """
    Word-008 专用评估：检查字体是否全部改为 Times New Roman。

    评估逻辑:
      1. 用 compare_docx_files 检查文本内容是否一致
      2. 遍历结果文件所有非空 run，检查 run.font.name
      3. 统计 font_name == "Times New Roman" 的 run 比例
      4. 得分 = 0.3 * text_score + 0.7 * font_match_ratio

    输入:
        gt_path: GT docx 文件路径
        result_path: 待评估 docx 文件路径

    输出:
        float，0.0~1.0
    """
    from docx import Document as DocxDocument
    from desktop_env.evaluators.metrics.docs import compare_docx_files

    _log = logging.getLogger("pipeline.word008")

    # ── 第 1 步：文本内容一致性 ──
    try:
        text_score = float(compare_docx_files(result_path, gt_path))
    except Exception as exc:
        _log.error("[word-008] compare_docx_files 异常: %s", exc)
        text_score = 0.0

    # ── 第 2 步：检查字体名称 ──
    try:
        doc = DocxDocument(result_path)
    except Exception as exc:
        _log.error("[word-008] 无法打开结果文件: %s", exc)
        return 0.0

    total_runs = 0
    tnr_runs = 0

    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue  # 跳过空 run
            total_runs += 1
            font_name = run.font.name
            if font_name == "Times New Roman":
                tnr_runs += 1

    if total_runs == 0:
        _log.warning("[word-008] 未找到非空 run")
        return 0.3 * text_score

    font_ratio = tnr_runs / total_runs
    score = 0.3 * text_score + 0.7 * font_ratio
    _log.info("[word-008] text=%.2f, font_ratio=%.2f (%d/%d), final=%.2f",
              text_score, font_ratio, tnr_runs, total_runs, score)
    return score


def compare_docx_double_spacing(gt_path: str, result_path: str) -> float:
    """
    Word-009 专用评估：检查文档是否改为双倍行距。

    评估逻辑:
      1. 用 compare_docx_files 检查文本内容是否一致
      2. 用 compare_line_spacing 检查行距是否与 GT 一致
      3. 得分 = 0.3 * text_score + 0.7 * line_spacing_score

    输入:
        gt_path: GT docx 文件路径
        result_path: 待评估 docx 文件路径

    输出:
        float，0.0~1.0
    """
    from desktop_env.evaluators.metrics.docs import (
        compare_docx_files,
        compare_line_spacing,
    )

    _log = logging.getLogger("pipeline.word009")

    # ── 第 1 步：文本内容一致性 ──
    try:
        text_score = float(compare_docx_files(result_path, gt_path))
    except Exception as exc:
        _log.error("[word-009] compare_docx_files 异常: %s", exc)
        text_score = 0.0

    # ── 第 2 步：行距一致性（compare_line_spacing 返回 0 或 1） ──
    try:
        line_spacing_score = float(compare_line_spacing(result_path, gt_path))
    except Exception as exc:
        _log.error("[word-009] compare_line_spacing 异常: %s", exc)
        line_spacing_score = 0.0

    score = 0.3 * text_score + 0.7 * line_spacing_score
    _log.info("[word-009] text=%.2f, line_spacing=%.2f, final=%.2f",
              text_score, line_spacing_score, score)
    return score


def compare_xlsx_hide_na_focused(gt_path: str, result_path: str) -> float:
    """
    Excel-008 专用评估：隐藏包含 N/A 的行，侧重行属性检查。

    默认 compare_xlsx_comprehensive 中 row_props 仅 0.1 权重，
    导致未修改文件（数据相同、样式相同但行未隐藏）得分 0.9，产生假阳性。
    本函数将 row_props 权重提升至 0.6，确保隐藏行操作是评分关键。

    权重:
      - data（0.2）：单元格数据值
      - style（0.2）：单元格样式
      - row_props（0.6）：行属性（隐藏行，核心检查项）

    输入:
        gt_path: GT xlsx 文件路径
        result_path: 待评估 xlsx 文件路径

    输出:
        float，0.0~1.0
    """
    from desktop_env.evaluators.metrics.comprehensive import compare_xlsx_comprehensive

    return compare_xlsx_comprehensive(
        gt_path, result_path,
        dimension_weights={"data": 0.2, "style": 0.2, "row_props": 0.6},
    )


# key = task_id, value = callable(gt_path, result_path) -> float
CUSTOM_FILE_EVALUATORS: Dict[str, Any] = {
    "Operation-FileOperate-Batchoperationexcel-006": compare_xlsx_negative_profit_red,
    "Operation-FileOperate-Batchoperationppt-001": compare_pptx_first_slide_transition,
    "Operation-FileOperate-BatchoperationWord-001": compare_docx_style_focused,
    # 第三梯队自定义评估器
    "Operation-FileOperate-BatchoperationWord-007": compare_docx_vowels_red,
    "Operation-FileOperate-BatchoperationWord-008": compare_docx_font_times_new_roman,
    "Operation-FileOperate-BatchoperationWord-009": compare_docx_double_spacing,
    "Operation-FileOperate-Batchoperationexcel-008": compare_xlsx_hide_na_focused,
}


# ============================================================
# 任务扫描
# ============================================================

def load_task_list(task_list_path: str) -> set:
    """
    从任务列表文件加载允许运行的 task_id 集合。

    文件格式：每行一个 task_id，空行和 # 开头的行被忽略。

    输入:
        task_list_path: 任务列表文件路径

    输出:
        set[str]，task_id 集合；文件不存在或为空时返回空集合
    """
    if not os.path.isfile(task_list_path):
        logging.getLogger("pipeline").warning("任务列表文件不存在: %s", task_list_path)
        return set()

    allowed = set()
    with open(task_list_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#"):
                allowed.add(line)
    return allowed


def scan_operation_tasks(
    tasks_dir: str,
    allowed_task_ids: set = None,
) -> List[Tuple[str, str, Dict]]:
    """
    扫描 tasks_list 目录，筛选 Operation 文件操作任务。

    筛选条件:
      - task_id 以 "Operation-FileOperate-" 开头
      - task_tag == "FileOperate"
      - task_type != "QA"
      - task_id 不包含 "coding"（已有专属 pipeline）
      - evaluator_path 不指向 .json（OSWorld 脚本任务，走另一套流程）
      - task_uid 非空
      - 若提供 allowed_task_ids，则 task_id 必须在该集合内

    输入:
        tasks_dir: 任务 JSON 文件所在目录
        allowed_task_ids: 可选的 task_id 白名单；为 None 时不做额外过滤

    输出:
        [(task_uid, task_path, task_config), ...]，按 task_id 排序
    """
    results = []

    if not os.path.isdir(tasks_dir):
        logging.getLogger("pipeline").error("任务目录不存在: %s", tasks_dir)
        return results

    for fname in os.listdir(tasks_dir):
        if not fname.endswith(".json"):
            continue

        task_path = os.path.join(tasks_dir, fname)
        try:
            config = load_task_config(task_path)
        except Exception:
            continue

        task_id = config.get("task_id", "")
        task_uid = config.get("task_uid", "")
        task_tag = config.get("task_tag", "")
        task_type = config.get("task_type", "")
        evaluator_path = config.get("evaluator_path", "")

        # 基本过滤
        if not task_id.startswith("Operation-FileOperate-"):
            continue
        if task_tag != "FileOperate":
            continue
        if task_type == "QA":
            continue
        if not task_uid:
            continue

        # 排除 coding 任务（已有专属 pipeline）
        if "coding" in task_id.lower():
            continue

        # 排除 evaluator_path 指向 .json 的 OSWorld 脚本任务
        # 但在白名单模式下允许通过（由白名单显式控制）
        if evaluator_path and evaluator_path.endswith(".json"):
            if allowed_task_ids is None:
                continue

        # 白名单过滤：仅运行指定的任务
        if allowed_task_ids is not None and task_id not in allowed_task_ids:
            continue

        results.append((task_uid, task_path, config))

    # 按 task_id 排序以保证可复现顺序
    results.sort(key=lambda x: x[2].get("task_id", ""))
    return results


# ============================================================
# 扁平化 shared 目录（支持指定 vm_ip 的版本）
# ============================================================

def _flatten_shared_dir_with_ip(vm_ip: str, vm_port: int, subdir: str, log: logging.Logger) -> bool:
    """
    将 subdir 下的文件扁平化到 /home/user/shared 根目录，并删除嵌套父目录。
    是 _flatten_shared_dir 的 vm_ip 参数化版本。

    输入:
        vm_ip: VM 的 IP 地址
        vm_port: VM 端口
        subdir: 子目录路径（如 "benchmark_dataset/xxx-uuid"）
        log: logger

    输出:
        bool（是否整理成功）
    """
    if not subdir:
        log.warning("subdir 为空，跳过整理")
        return True

    parent_dir = os.path.dirname(subdir)
    if not parent_dir:
        log.warning("subdir 无父目录，跳过删除父目录")
        return True

    src_root = f"/home/user/shared/{subdir}"
    parent_root = f"/home/user/shared/{parent_dir}"
    cmd = (
        "bash -c "
        f"\"find '{src_root}' -type f -print0 | "
        "xargs -0 -I{} mv -f {} /home/user/shared/ && "
        f"rm -rf '{parent_root}'\""
    )
    result = execute_on_vm_with_ip(vm_ip, vm_port, cmd)
    if result.get("status") != "success":
        log.warning("共享目录整理失败: %s", result.get("error", "Unknown"))
        return False
    log.info("已将嵌套文件移动到 shared 根目录")
    return True


# ============================================================
# 覆写 init_vm_parallel：下载后调用 flatten
# ============================================================

def init_vm_with_flatten(
    vm_port: int,
    vnc_port: int,
    prepare_url: str,
    shared_host_dir: str,
    vm_ip: str,
    is_first_vm: bool,
    rebuilt: bool,
    log: logging.Logger,
) -> bool:
    """
    VM 初始化（在 QA parallel 版基础上，增加下载后扁平化步骤）。

    与 run_QA_pipeline_parallel.init_vm_parallel 的区别:
      - 步骤 [5/6] 下载完数据后，额外调用 _flatten_shared_dir_with_ip()
        将嵌套文件移到 /home/user/shared/ 根目录

    输入:
        vm_port: VM server 端口
        vnc_port: VNC 端口
        prepare_url: 任务数据 URL
        shared_host_dir: 宿主机共享目录路径（该组专用）
        vm_ip: 宿主机 IP
        is_first_vm: 是否为该组的第一个 VM（仅第一个 VM 下载数据）
        rebuilt: 是否为重建后首次初始化
        log: logger

    输出:
        bool
    """
    log.info("初始化 VM (port %d, VNC http://%s:%d/)", vm_port, vm_ip, vnc_port)

    wait_time = 120 if rebuilt else 30
    if not wait_for_vm_ready_with_ip(vm_ip, vm_port, max_wait=wait_time):
        log.error("VM %d 无法响应", vm_port)
        return False

    # [1/6] 检查并安装 sshfs（含重试机制）
    log.info("[1/6] 检查 sshfs...")
    result = execute_on_vm_with_ip(vm_ip, vm_port, "which sshfs")
    if result.get("status") != "success":
        _ = execute_on_vm_with_ip(
            vm_ip, vm_port,
            'bash -c "echo password | sudo -S systemctl stop packagekit || true; '
            'echo password | sudo -S systemctl disable packagekit || true"',
        )
        # apt update: 超时 150 秒，最多重试 3 次
        apt_ok = False
        for attempt in range(3):
            result = execute_on_vm_with_ip(
                vm_ip, vm_port,
                'bash -c "echo password | sudo -S apt update -qq"',
                timeout=150,
            )
            if result.get("status") == "success":
                apt_ok = True
                break
            log.warning(
                "apt update 失败 (第 %d/3 次): %s",
                attempt + 1, result.get("error", "Unknown"),
            )
            if attempt < 2:
                import time as _time
                _time.sleep(5)
        if not apt_ok:
            log.error("apt update 重试 3 次均失败，中断初始化")
            return False
        # apt install sshfs: 同样增加超时
        result = execute_on_vm_with_ip(
            vm_ip, vm_port,
            'bash -c "echo password | sudo -S DEBIAN_FRONTEND=noninteractive '
            'apt install -y -qq sshfs"',
            timeout=150,
        )
        if result.get("status") != "success":
            log.error("安装 sshfs 失败: %s", result.get("error", "Unknown"))
            return False
        result = execute_on_vm_with_ip(vm_ip, vm_port, "which sshfs")
        if result.get("status") != "success":
            log.error("sshfs 验证失败")
            return False
        log.info("  sshfs 安装完成")
    else:
        log.info("  sshfs 已安装")

    # [2/6] 准备 shared 目录
    log.info("[2/6] 准备 shared 目录...")
    cmd = (
        'bash -c "echo password | sudo -S fusermount3 -u /home/user/shared 2>/dev/null; '
        'mkdir -p /home/user/shared"'
    )
    result = execute_on_vm_with_ip(vm_ip, vm_port, cmd)
    if result.get("status") != "success":
        log.error("准备 shared 目录失败: %s", result.get("error", "Unknown"))
        return False

    # [3/6] 挂载 shared（使用参数化路径）
    log.info("[3/6] 挂载 shared (%s)...", shared_host_dir)
    _creds = get_ssh_credentials(vm_ip)
    cmd = (
        f"bash -c \"echo '{_creds['ssh_password']}' | sshfs {_creds['ssh_host']}:{shared_host_dir} "
        "/home/user/shared -o password_stdin -o StrictHostKeyChecking=no\""
    )
    result = execute_on_vm_with_ip(vm_ip, vm_port, cmd)
    if result.get("status") != "success":
        log.error("挂载 shared 失败: %s", result.get("error", "Unknown"))
        return False

    # [4/6] 验证挂载
    log.info("[4/6] 验证 shared 挂载...")
    result = execute_on_vm_with_ip(vm_ip, vm_port, "ls /home/user/shared")
    if result.get("status") != "success":
        log.error("shared 挂载验证失败: %s", result.get("error", "Unknown"))
        return False

    # [5/6] 仅第一个 VM 下载任务数据 + 扁平化
    if is_first_vm:
        if prepare_url:
            log.info("[5/6] 下载任务数据到 shared...")
            if not _download_task_files_on_vm_with_ip(
                vm_ip, vm_port, prepare_url,
                host_shared_dir=shared_host_dir,
            ):
                return False

            # 关键修复：下载后扁平化目录结构
            # 文件下载到 /home/user/shared/benchmark_dataset/{task_uid}/...
            # 需要移到 /home/user/shared/ 根目录，否则 Agent 找不到文件
            try:
                _, _, subdir = parse_prepare_script_path(prepare_url)
                if subdir:
                    log.info("扁平化 shared 目录: %s", subdir)
                    _flatten_shared_dir_with_ip(vm_ip, vm_port, subdir, log)
            except Exception as exc:
                log.warning("扁平化目录失败（非致命）: %s", exc)
        else:
            log.info("[5/6] 任务未提供 prepare_script_path，跳过下载")
    else:
        log.info("[5/6] 跳过下载（使用 shared 中的文件）")

    # [6/6] 启动 Chrome 并打开 Bing
    log.info("[6/6] 启动 Chrome 并打开 Bing...")
    cmd = (
        'bash -c "nohup python3 -c \\"import subprocess, time, os; '
        'env = os.environ.copy(); '
        "env['DISPLAY'] = ':0'; "
        "subprocess.Popen(['google-chrome', '--no-first-run', '--no-default-browser-check', "
        "'https://www.bing.com'], env=env); "
        'time.sleep(2)\\" '
        '> /tmp/bootstrap_chrome.log 2>&1 &"'
    )
    _ = execute_on_vm_with_ip(vm_ip, vm_port, cmd)
    log.info("  Chrome 已启动并打开 Bing（异步执行）")

    log.info("VM %d 初始化成功", vm_port)
    return True


# ============================================================
# 覆写 stage1：使用 init_vm_with_flatten 替代 init_vm_parallel
# ============================================================

def stage1_initialize_with_flatten(
    task_config: Dict[str, Any],
    config: ContainerSetConfig,
    log: logging.Logger,
) -> bool:
    """
    Stage 1: 环境初始化（与 QA parallel 版本相同，但使用带 flatten 的 init_vm）。

    输入:
        task_config: 任务配置
        config: 容器组配置
        log: logger

    输出:
        bool
    """
    log.info("STAGE 1: 环境初始化 (组 %d)", config.group_id)

    prepare_url = task_config.get("prepare_script_path", "")
    if not prepare_url:
        log.warning("任务配置缺少 prepare_script_path，将跳过下载步骤但继续执行任务")

    # 通过 SSH 在宿主机上创建共享目录
    log.info("通过 SSH 在宿主机上创建共享目录: %s", config.shared_host_dir)
    _creds = get_ssh_credentials(config.vm_ip)
    try:
        result = run_ssh_command(
            _creds["ssh_password"], _creds["ssh_opts"], _creds["ssh_host"],
            f"mkdir -p {config.shared_host_dir} && chmod 777 {config.shared_host_dir}",
            timeout=30,
        )
        if result.returncode == 0:
            log.info("共享目录创建成功")
        else:
            log.warning("共享目录创建失败: %s", result.stderr)
    except Exception as exc:
        log.warning("创建共享目录异常: %s", exc)

    # 重建容器
    rebuilt = rebuild_containers_parallel(config, log)
    if not rebuilt:
        log.error("容器重建失败，终止初始化")
        return False

    # 初始化各 VM（使用带 flatten 的版本）
    vm_pairs = config.get_vm_pairs()
    success_count = 0
    for idx, (vm_port, vnc_port) in enumerate(vm_pairs):
        if init_vm_with_flatten(
            vm_port=vm_port,
            vnc_port=vnc_port,
            prepare_url=prepare_url,
            shared_host_dir=config.shared_host_dir,
            vm_ip=config.vm_ip,
            is_first_vm=(idx == 0),
            rebuilt=True,
            log=log,
        ):
            success_count += 1
        else:
            log.warning("VM %d 初始化失败，继续下一个...", vm_port)

    # 禁用屏保（防黑屏第一层）
    if success_count > 0:
        disable_screensaver_parallel(config.vm_ip, config.get_server_ports(), log)

    log.info("初始化完成: %d/%d 个 VM 成功", success_count, len(vm_pairs))
    return success_count == len(vm_pairs)


# ============================================================
# GT 下载
# ============================================================

def download_gt_from_hf(task_uid: str, local_gt_base: str, log: logging.Logger) -> str:
    """
    从 HuggingFace answer_files/{task_uid}/ 下载 GT 文件到 Mac 本地。
    有缓存机制：如果本地目录已存在且非空，跳过下载。

    注意: HuggingFace 上 GT 目录按 task_uid（UUID）组织，不是 task_id。

    输入:
        task_uid: 任务 UID（如 "a1510a05-9fca-46ba-b95d-451dd5779194"）
        local_gt_base: 本地 GT 缓存根目录
        log: logger

    输出:
        GT 目录路径（成功）或空字符串（失败）
    """
    gt_dir = os.path.join(local_gt_base, task_uid)

    # 缓存检查：目录存在且包含文件则跳过下载
    if os.path.isdir(gt_dir) and os.listdir(gt_dir):
        log.info("GT 缓存命中: %s (%d 个文件)", task_uid, len(os.listdir(gt_dir)))
        return gt_dir

    os.makedirs(gt_dir, exist_ok=True)

    # 获取文件列表
    api_url = (
        f"{HF_API_BASE}/tree/main/answer_files/{task_uid}?recursive=1"
    )
    try:
        resp = requests.get(api_url, timeout=30)
        resp.raise_for_status()
        file_entries = resp.json()
    except Exception as exc:
        log.error("获取 GT 文件列表失败: %s (URL: %s)", exc, api_url)
        return ""

    if not file_entries:
        log.warning("GT 文件列表为空: %s", task_uid)
        return ""

    # 逐个下载
    downloaded = 0
    for entry in file_entries:
        # HF API 返回的 entry 格式: {"type": "file", "path": "answer_files/{task_uid}/xxx.docx", ...}
        if entry.get("type") != "file":
            continue

        remote_path = entry.get("path", "")
        if not remote_path:
            continue

        # 计算本地文件名：去掉 answer_files/{task_uid}/ 前缀
        prefix = f"answer_files/{task_uid}/"
        if remote_path.startswith(prefix):
            rel_name = remote_path[len(prefix):]
        else:
            rel_name = os.path.basename(remote_path)

        local_path = os.path.join(gt_dir, rel_name)
        local_dir = os.path.dirname(local_path)
        os.makedirs(local_dir, exist_ok=True)

        download_url = f"{HF_DOWNLOAD_BASE}/{remote_path}"
        try:
            resp = requests.get(download_url, timeout=60, stream=True)
            resp.raise_for_status()
            with open(local_path, "wb") as f:
                for chunk in resp.iter_content(chunk_size=8192):
                    f.write(chunk)
            downloaded += 1
        except Exception as exc:
            log.error("下载 GT 文件失败: %s → %s", download_url, exc)
            return ""

    log.info("GT 下载完成: %s (%d 个文件)", task_uid, downloaded)
    return gt_dir if downloaded > 0 else ""


# ============================================================
# Agent 结果下载
# ============================================================

def download_agent_result_from_host(
    shared_host_dir: str,
    local_result_dir: str,
    vm_ip: str,
    log: logging.Logger,
) -> str:
    """
    通过 SSH + tar 管道从宿主机 shared 目录下载 Agent 结果文件到 Mac 本地。

    输入:
        shared_host_dir: 宿主机 shared 目录路径（如 /home/agentlab/shared/group_0）
        local_result_dir: Mac 本地结果目录（将创建或清空）
        vm_ip: 宿主机 IP
        log: logger

    输出:
        本地结果目录路径（成功）或空字符串（失败）
    """
    os.makedirs(local_result_dir, exist_ok=True)

    _creds = get_ssh_credentials(vm_ip)
    ssh_env = os.environ.copy()
    ssh_env["SSHPASS"] = _creds["ssh_password"]

    # 在宿主机上 tar 打包 → 管道传输到本地解压
    remote_cmd = f"tar czf - -C {shared_host_dir} ."
    ssh_cmd = (
        ["sshpass", "-e", "ssh"]
        + _creds["ssh_opts"]
        + [_creds["ssh_host"], remote_cmd]
    )
    tar_extract_cmd = ["tar", "xzf", "-", "-C", local_result_dir]

    try:
        ssh_proc = subprocess.Popen(
            ssh_cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, env=ssh_env,
        )
        tar_proc = subprocess.Popen(
            tar_extract_cmd, stdin=ssh_proc.stdout, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        )
        if ssh_proc.stdout:
            ssh_proc.stdout.close()

        _, tar_err = tar_proc.communicate(timeout=120)
        ssh_proc.wait(timeout=10)

        if tar_proc.returncode != 0:
            log.error("tar 解压失败: %s", tar_err.decode(errors="replace"))
            return ""

        # 统计下载的文件数
        file_count = sum(len(files) for _, _, files in os.walk(local_result_dir))
        log.info("结果文件下载完成: %d 个文件 → %s", file_count, local_result_dir)
        return local_result_dir

    except Exception as exc:
        log.error("下载结果文件失败: %s", exc)
        return ""


# ============================================================
# 文件比对评估
# ============================================================


def _build_filename_map(directory: str) -> Dict[str, str]:
    """
    递归扫描目录，建立 filename -> 完整路径 的映射。
    如果存在同名文件，保留路径最短（最浅层）的。

    输入:
        directory: 目录路径

    输出:
        {filename: full_path} 映射
    """
    mapping: Dict[str, str] = {}
    for root, _, files in os.walk(directory):
        for fname in files:
            fpath = os.path.join(root, fname)
            if fname not in mapping:
                mapping[fname] = fpath
            else:
                # 保留路径更短的（更浅层的文件优先）
                if len(fpath) < len(mapping[fname]):
                    mapping[fname] = fpath
    return mapping


def _compare_single_file(gt_path: str, result_path: str, log: logging.Logger) -> float:
    """
    综合比对单个文件（docx/xlsx/pptx），委托给 comprehensive 模块。

    输入:
        gt_path: GT 文件路径
        result_path: Agent 结果文件路径
        log: logger

    输出:
        float，比对得分（0.0 ~ 1.0）
    """
    try:
        score = compare_file_comprehensive(gt_path, result_path)
        return score
    except Exception as exc:
        log.error("文件比对异常: %s vs %s → %s", gt_path, result_path, exc)
        return 0.0


def match_and_compare_files(
    gt_dir: str,
    result_dir: str,
    _task_config: Dict[str, Any],
    log: logging.Logger,
) -> Dict[str, Any]:
    """
    核心评估逻辑：将 GT 文件与 Agent 结果文件按文件名匹配并逐个比对。

    匹配策略:
      1. 精确匹配（优先）
      2. 大小写不敏感匹配（fallback）

    输入:
        gt_dir: GT 文件目录
        result_dir: Agent 结果文件目录
        _task_config: 任务配置（查找 CUSTOM_FILE_EVALUATORS 中的自定义比对函数）
        log: logger

    输出:
        {
            "score": float,           # 0.0~1.0
            "pass": bool,             # score >= 0.5
            "file_scores": {filename: score},
            "missing_files": [...],
            "gt_file_count": int,
            "result_file_count": int,
            "reason": str,
        }
    """
    task_id = _task_config.get("task_id", "")
    custom_compare_fn = CUSTOM_FILE_EVALUATORS.get(task_id)

    gt_map = _build_filename_map(gt_dir)
    result_map = _build_filename_map(result_dir)

    # 过滤掉非文档文件（只比对 .docx/.xlsx/.pptx）
    doc_extensions = {".docx", ".xlsx", ".pptx"}
    gt_map = {k: v for k, v in gt_map.items()
              if os.path.splitext(k)[1].lower() in doc_extensions}

    if not gt_map:
        log.warning("GT 目录中没有可比对的文档文件")
        return {
            "score": 0.0,
            "pass": False,
            "file_scores": {},
            "missing_files": [],
            "gt_file_count": 0,
            "result_file_count": len(result_map),
            "reason": "GT 目录中没有文档文件",
        }

    # 建立大小写不敏感的查找索引
    result_map_lower = {k.lower(): v for k, v in result_map.items()}

    file_scores: Dict[str, float] = {}
    missing_files: List[str] = []

    for gt_name, gt_path in gt_map.items():
        # 策略 1：精确匹配
        if gt_name in result_map:
            result_path = result_map[gt_name]
        # 策略 2：大小写不敏感
        elif gt_name.lower() in result_map_lower:
            result_path = result_map_lower[gt_name.lower()]
        else:
            log.warning("GT 文件 '%s' 在结果目录中找不到匹配", gt_name)
            missing_files.append(gt_name)
            file_scores[gt_name] = 0.0
            continue

        log.info("比对: %s", gt_name)
        if custom_compare_fn:
            log.info("使用自定义比对: %s", task_id)
            score = custom_compare_fn(gt_path, result_path)
        else:
            score = _compare_single_file(gt_path, result_path, log)
        file_scores[gt_name] = score
        log.info("  得分: %.2f", score)

    # 总分 = 各文件得分的平均值
    total_score = sum(file_scores.values()) / len(file_scores) if file_scores else 0.0

    return {
        "score": total_score,
        "pass": total_score >= 0.5,
        "file_scores": file_scores,
        "missing_files": missing_files,
        "gt_file_count": len(gt_map),
        "result_file_count": len(result_map),
        "reason": (
            f"比对 {len(file_scores)} 个文件，"
            f"缺失 {len(missing_files)} 个，"
            f"平均得分 {total_score:.2f}"
        ),
    }


# ============================================================
# Stage 3: Operation 任务评估
# ============================================================

def stage3_evaluate_operation(
    task_config: Dict[str, Any],
    agent_result: Dict[str, Any],
    task_path: str,
    config: ContainerSetConfig,
    gt_cache_dir: str,
    log: logging.Logger,
    save_result_dir: str = "",
) -> Dict[str, Any]:
    """
    Stage 3: Operation 任务评估。

    流程:
      1. 如果任务有自定义 evaluator_path (.py) → 动态加载并调用
      2. 否则走默认文件比对：
         a. 下载 GT 到 Mac 本地（有缓存）
         b. 下载 Agent 结果到 Mac 本地
         c. 文件比对评估
         d. 保存结果文件（如指定 save_result_dir）+ 清理临时文件

    输入:
        task_config: 任务配置
        agent_result: Stage 2 的执行结果
        task_path: 任务 JSON 文件路径
        config: 容器组配置（用于获取 shared_host_dir 等）
        gt_cache_dir: GT 文件缓存目录
        log: logger
        save_result_dir: 结果文件持久化目录（可选，为空则不保存）

    输出:
        评估结果字典（包含 saved_result_path 字段，指示结果文件保存位置）
    """
    task_id = task_config.get("task_id", "unknown")
    evaluator_path = task_config.get("evaluator_path", "")

    # 路径 0：规则化评估（task_config 中包含 eval_rules 时优先使用）
    if task_config.get("eval_rules"):
        log.info("使用规则化评估器（eval_rules）: %s", task_id)
        try:
            from eval.operation_evaluator import evaluate as rule_evaluate

            # 下载 Agent 结果到本地（复用现有逻辑）
            local_result_dir = os.path.join(
                tempfile.gettempdir(),
                "operation_result",
                f"group_{config.group_id}",
                task_id,
            )
            if os.path.exists(local_result_dir):
                shutil.rmtree(local_result_dir, ignore_errors=True)

            result_dir = download_agent_result_from_host(
                config.shared_host_dir, local_result_dir, config.vm_ip, log,
            )
            if not result_dir:
                return {"score": 0.0, "pass": False, "reason": "Agent 结果下载失败"}

            # 保存结果文件
            saved_path = ""
            if save_result_dir:
                try:
                    task_save_dir = os.path.join(save_result_dir, task_id)
                    if os.path.exists(task_save_dir):
                        shutil.rmtree(task_save_dir, ignore_errors=True)
                    shutil.copytree(local_result_dir, task_save_dir)
                    saved_path = task_save_dir
                except Exception as exc:
                    log.warning("保存结果文件失败: %s", exc)

            eval_result = rule_evaluate(result_dir, task_config)

            # 清理临时文件
            try:
                shutil.rmtree(local_result_dir, ignore_errors=True)
            except Exception:
                pass

            eval_result["saved_result_path"] = saved_path
            return eval_result
        except Exception as exc:
            log.error("规则化评估器执行失败: %s", exc, exc_info=True)
            return {"score": 0.0, "pass": False, "reason": f"规则化评估异常: {exc}"}

    # 路径 1：自定义 evaluator（.py 脚本）
    if evaluator_path and evaluator_path.endswith(".py"):
        log.info("使用自定义 evaluator: %s", evaluator_path)
        try:
            # load_evaluator 内部会拼接 parallel_benchmark_dir，直接传相对路径
            evaluator_module = load_evaluator(evaluator_path)
            # 自定义 evaluator 通常接受 (task_path, final_answer) 或类似签名
            execution_record = (
                agent_result.get("execution_record", {})
                if isinstance(agent_result, dict) else {}
            )
            summary = extract_execution_summary(execution_record)
            final_answer = summary.get("model_output_answer", "")
            eval_result = evaluator_module.evaluate(task_path, final_answer)
            return eval_result
        except Exception as exc:
            log.error("自定义 evaluator 执行失败: %s", exc)
            return {"score": 0.0, "pass": False, "reason": f"evaluator 异常: {exc}"}

    # 路径 2：OSWorld JSON 评测配置
    if evaluator_path and evaluator_path.endswith(".json"):
        log.info("使用 OSWorld JSON 评测配置: %s", evaluator_path)
        try:
            from eval.osworld_evaluator import evaluate_osworld_task
            json_path = os.path.join(parallel_benchmark_dir, evaluator_path)
            vm_pairs = config.get_vm_pairs()
            vm_port = vm_pairs[0][0]  # 使用第一个 VM
            return evaluate_osworld_task(
                evaluator_json_path=json_path,
                vm_ip=config.vm_ip,
                vm_port=vm_port,
                shared_host_dir=config.shared_host_dir,
                log=log,
            )
        except Exception as exc:
            log.error("OSWorld 评测执行失败: %s", exc, exc_info=True)
            return {"score": 0.0, "pass": False, "reason": f"OSWorld 评测异常: {exc}"}

    # 路径 3：默认文件比对
    log.info("使用默认文件比对评估")

    # Step A: 下载 GT（有缓存，多组任务共享同一份 GT）
    # 注意: HuggingFace 上 GT 按 task_uid 组织，不是 task_id
    task_uid = task_config.get("task_uid", "")
    if not task_uid:
        return {"score": 0.0, "pass": False, "reason": f"任务缺少 task_uid: {task_id}"}
    gt_dir = download_gt_from_hf(task_uid, gt_cache_dir, log)

    # Step B: 下载 Agent 结果到 Mac 本地
    # 使用 group_id + task_id 双重隔离，防止并行冲突
    local_result_dir = os.path.join(
        tempfile.gettempdir(),
        "operation_result",
        f"group_{config.group_id}",
        task_id,
    )
    # 清理旧结果（如果存在）
    if os.path.exists(local_result_dir):
        shutil.rmtree(local_result_dir, ignore_errors=True)

    result_dir = download_agent_result_from_host(
        config.shared_host_dir, local_result_dir, config.vm_ip, log,
    )

    # GT 不可用时：保存 Agent 结果供人工评价，返回 unknown
    if not gt_dir:
        log.warning("GT 不可用，跳过自动评估，保存 Agent 结果供人工评价")
        saved_path = ""
        if save_result_dir and result_dir:
            try:
                task_save_dir = os.path.join(save_result_dir, task_id)
                if os.path.exists(task_save_dir):
                    shutil.rmtree(task_save_dir, ignore_errors=True)
                shutil.copytree(local_result_dir, task_save_dir)
                saved_path = task_save_dir
                log.info("Agent 结果已保存: %s", task_save_dir)
            except Exception as exc:
                log.warning("保存结果失败: %s", exc)
        try:
            shutil.rmtree(local_result_dir, ignore_errors=True)
        except Exception:
            pass
        return {
            "score": -1, "pass": None, "status": "unknown",
            "reason": f"GT 不可用 ({task_id}, uid={task_uid})，已保存结果供人工评价",
            "saved_result_path": saved_path,
        }

    if not result_dir:
        return {
            "score": 0.0,
            "pass": False,
            "reason": "Agent 结果下载失败",
        }

    # Step C-1: 先保存结果文件（在评估之前，确保即使评估崩溃也能保留 Agent 产物）
    saved_path = ""
    if save_result_dir:
        try:
            task_save_dir = os.path.join(save_result_dir, task_id)
            if os.path.exists(task_save_dir):
                shutil.rmtree(task_save_dir, ignore_errors=True)
            shutil.copytree(local_result_dir, task_save_dir)
            saved_path = task_save_dir
            log.info("Agent 结果文件已保存到: %s", task_save_dir)
        except Exception as exc:
            log.warning("保存结果文件失败: %s", exc)

    # Step C-2: 文件比对评估
    try:
        eval_result = match_and_compare_files(gt_dir, result_dir, task_config, log)
    except Exception as exc:
        log.error("文件比对评估异常: %s", exc)
        eval_result = {"score": 0.0, "pass": False, "reason": f"比对异常: {exc}"}

    # Step D: 清理临时文件
    try:
        shutil.rmtree(local_result_dir, ignore_errors=True)
    except Exception:
        pass

    eval_result["saved_result_path"] = saved_path
    return eval_result


# ============================================================
# Stage 2: 纯 GUI Agent 执行（gui_only 模式）
# ============================================================

def stage2_execute_gui_only(
    task_config: Dict[str, Any],
    task_uid: str,
    config: "ContainerSetConfig",
    log: logging.Logger,
    gui_agent: str = "seed18",
    max_rounds: int = 200,
    gui_timeout: int = 3600,
    output_dir: str = "",
) -> Tuple[Dict[str, Any], PythonController]:
    """
    纯 GUI Agent 模式的 Stage 2：单个 GUI Agent 在单台 VM 上完成完整任务。
    不经过 Plan Agent 任务分解，直接调用 GUI Agent。

    输入:
        task_config: 任务配置
        task_uid: 任务 UID
        config: 容器组配置（仅使用第一个 VM）
        log: logger
        gui_agent: GUI Agent 类型（seed18 / claude / kimi）
        max_rounds: 最大执行轮次
        gui_timeout: 超时时间（秒）
        output_dir: 执行记录输出目录（为空则使用 ubuntu_env/logs/）

    输出:
        (result, controller_vm1) — result 格式与 stage2_execute_agent_parallel 兼容
    """
    log.info("STAGE 2 [gui_only]: 单个 GUI Agent 独立执行任务")

    task_instruction = task_config.get("instruction", "")
    if not task_instruction:
        raise ValueError("任务配置缺少 instruction")

    log.info("任务描述: %s", task_instruction[:200])
    log.info("GUI Agent: %s | 最大轮次: %d | 超时: %ds", gui_agent, max_rounds, gui_timeout)

    # 仅使用第一个 VM
    vm_ports = config.get_server_ports()
    first_port = vm_ports[0]
    controller_vm1 = PythonController(vm_ip=config.vm_ip, server_port=first_port)

    try:
        screenshot = controller_vm1.get_screenshot()
        log.info("VM1 (port %d) connected - Screenshot: %d bytes",
                 first_port, len(screenshot) if screenshot else 0)
    except Exception as e:
        log.warning("VM1 (port %d) connection warning: %s", first_port, e)

    # 根据 gui_agent 参数创建对应的 Tool 实例
    if gui_agent == "claude":
        gui_tool = ClaudeGUIAgentTool(controller=controller_vm1)
    elif gui_agent == "kimi":
        gui_tool = KimiGUIAgentTool(controller=controller_vm1)
    elif gui_agent == "seed18":
        gui_tool = Seed18GUIAgentTool(controller=controller_vm1, prompt_mode="gui_only")
    elif gui_agent == "gpt54":
        from parallel_agents_as_tools.gpt54_gui_agent_as_tool import GPT54GUIAgentTool
        gui_tool = GPT54GUIAgentTool(controller=controller_vm1, prompt_mode="gui_only")
    elif gui_agent == "gpt54_fc":
        from parallel_agents_as_tools.gpt_gui_agent_as_tool import GPTGUIAgentTool
        gui_tool = GPTGUIAgentTool(
            controller=controller_vm1,
            model_name="gpt-5.4-mini",
            api_config_key="pincc",
        )
    else:
        log.warning("未知的 gui_agent: %s，fallback 到 seed18", gui_agent)
        gui_tool = Seed18GUIAgentTool(controller=controller_vm1, prompt_mode="gui_only")

    start_time = time.time()
    gui_result = gui_tool.execute(
        task=task_instruction,
        max_rounds=max_rounds,
        timeout=gui_timeout,
    )
    elapsed_time = time.time() - start_time
    log.info("纯 GUI Agent 执行完成，耗时: %.2fs", elapsed_time)

    # 格式转换：GUI Agent result → Pipeline 统一格式
    final_answer = gui_result.get("result", "")
    gui_status = gui_result.get("status", "failure")
    gui_model = gui_result.get("model_name", gui_agent)
    gui_token = gui_result.get("gui_token_usage", {})
    gui_steps = gui_result.get("steps", [])
    gui_rounds_timing = gui_result.get("rounds_timing", [])

    # 构建与 extract_execution_summary 兼容的 execution_record
    execution_record = {
        "plan_agent": {
            "model_name": "",
            "rounds": [],
            "summary": {"total_rounds": 0},
        },
        "devices": [{
            "device_id": f"{config.vm_ip}:{first_port}",
            "agents": [{
                "model_name": gui_model,
                "summary": {
                    "total_rounds": len(gui_steps),
                    "final_status": gui_status,
                },
            }],
        }],
        "summary": {
            "final_answer": final_answer,
            "status": gui_status,
            "total_rounds": len(gui_steps),
            "mode": "gui_only",
        },
        "steps": gui_steps,
        "rounds_timing": gui_rounds_timing,
    }

    # 构建与 Plan Agent 兼容的 token_usage
    token_usage = {
        "plan_agent": {},
        "gui_agent": gui_token,
        "plan_agent_model": "",
        "gui_agent_model": gui_model,
    }

    result = {
        "execution_record": execution_record,
        "token_usage": token_usage,
    }

    # 保存执行记录
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    _record_dir = output_dir if output_dir else os.path.join(ubuntu_env_dir, "logs")
    os.makedirs(_record_dir, exist_ok=True)
    record_path = os.path.join(
        _record_dir, f"operation_gui_only_{task_uid}_{timestamp}.json"
    )
    try:
        with open(record_path, "w", encoding="utf-8") as f:
            json.dump({
                "task_uid": task_uid,
                "instruction": task_instruction,
                "gui_result": gui_result,
                "elapsed_time": elapsed_time,
            }, f, ensure_ascii=False, indent=2, default=str)
        log.info("执行记录已保存: %s", record_path)
    except Exception as exc:
        log.warning("保存执行记录失败: %s", exc)

    return result, controller_vm1


# ============================================================
# 单任务完整流程
# ============================================================

def run_single_task(
    task_uid: str,
    task_path: str,
    task_config: Dict[str, Any],
    available_groups: queue.Queue,
    args: argparse.Namespace,
    memory_guard: MemoryGuard,
    output_results: Dict[str, Any] = None,
    results_lock: threading.Lock = None,
    output_json_path: str = "",
) -> Dict[str, Any]:
    """
    单个任务的完整执行流程（在 Worker 线程中运行）。

    通过 available_groups 队列获取可用的 group_id，保证同一 group_id
    不会被两个线程同时使用，避免容器名称冲突。任务完成后归还 group_id。

    流程:
        0. available_groups.get() — 获取可用 group_id（阻塞直到有空闲组）
        1. memory_guard.acquire() — 申请内存额度
        2. allocate_ports_for_group() — 动态分配端口
        3. stage1_initialize_with_flatten() — 重建容器 + 初始化 VM + 扁平化
        4. stage2_execute_agent_parallel() — Agent 执行
        5. stage3_evaluate_operation() — Operation 文件比对评估
        6. 清理容器 + memory_guard.release() + 归还 group_id

    输入:
        task_uid: 任务 UID
        task_path: 任务 JSON 路径
        task_config: 任务配置字典
        available_groups: 可用 group_id 队列（线程安全）
        args: 命令行参数
        memory_guard: 内存管理器

    输出:
        task_result 字典
    """
    # 0. 从队列获取可用 group_id
    group_id = available_groups.get()

    # 构建容器组配置
    config = ContainerSetConfig(
        group_id=group_id,
        num_vms=args.vms_per_task,
        vm_memory=args.vm_memory,
        vm_cpu_cores=args.vm_cpu_cores,
        shared_host_dir=f"{args.shared_base_dir}/group_{group_id}",
        vm_ip=args.vm_ip,
        docker_image=args.docker_image,
        qcow2_path=args.qcow2_path,
    )

    task_id = task_config.get("task_id", "")
    log = get_task_logger(group_id, task_uid)
    log.info("获得组 %d，开始执行任务 %s (%s)", group_id, task_uid[:8], task_id)

    instruction = task_config.get("instruction", "")

    task_result: Dict[str, Any] = {
        "task_uid": task_uid,
        "task_id": task_id,
        "instruction": instruction,
        "model_output_answer": "",
        "plan_agent_model": "",
        "gui_agent_model": "",
        "plan_agent_total_rounds": 0,
        "evaluator_output": None,
        "token_usage": None,
        "plan_agent_last_round_output": "",
        "plan_agent_last_round_messages": [],
        "interrupted": False,
        "interrupt_reason": "",
        "group_id": group_id,
    }

    # 1. 申请内存额度
    if not memory_guard.acquire(config.num_vms):
        task_result["interrupted"] = True
        task_result["interrupt_reason"] = "memory_guard_timeout"
        log.error("内存申请超时，跳过任务")
        available_groups.put(group_id)
        return task_result

    try:
        # 2. 动态分配端口（含远程端口扫描，自动避开已占用端口）
        log.info("为组 %d 分配端口（扫描远程已用端口）...", group_id)
        _creds_port = get_ssh_credentials(config.vm_ip)
        remote_ports = scan_remote_docker_ports(
            ssh_password=_creds_port["ssh_password"],
            ssh_opts=_creds_port["ssh_opts"],
            ssh_host=_creds_port["ssh_host"],
            conda_activate=_creds_port["conda_activate"],
        )
        config.containers = allocate_ports_for_group(
            config.num_vms, group_id, extra_used_ports=remote_ports,
        )

        # 注册到全局追踪
        with _active_groups_lock:
            _active_groups[group_id] = config

        # 注册端口到心跳服务
        register_group_ports(group_id, config.get_server_ports())

        # 3. Stage 1: 环境初始化（使用带 flatten 的版本）
        if not stage1_initialize_with_flatten(task_config, config, log):
            task_result["interrupted"] = True
            task_result["interrupt_reason"] = "stage1_initialize_failed"
            log.error("环境初始化失败，跳过当前任务")
            return task_result

        # 4. Stage 2: Agent 执行
        agent_mode = getattr(args, "agent_mode", "plan")
        try:
            if agent_mode == "gui_only":
                result, _ = stage2_execute_gui_only(
                    task_config, task_uid, config, log,
                    gui_agent=getattr(args, "gui_agent", "seed18"),
                    max_rounds=getattr(args, "gui_max_rounds", 200),
                    gui_timeout=getattr(args, "gui_timeout", 3600),
                    output_dir=os.environ.get("ABLATION_RECORD_DIR", ""),
                )
            else:
                result, _ = stage2_execute_agent_parallel(task_config, task_uid, config, log)
        except Exception as exc:
            task_result["interrupted"] = True
            task_result["interrupt_reason"] = f"stage2_execute_exception: {exc}"
            log.error("Agent 执行失败: %s", exc)
            return task_result

        # Plan 模式执行记录保存（与 gui_only 模式保持一致的输出逻辑）
        if agent_mode != "gui_only":
            _record_dir = os.environ.get("ABLATION_RECORD_DIR", "") or os.path.join(ubuntu_env_dir, "logs")
            os.makedirs(_record_dir, exist_ok=True)
            _ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            _record_path = os.path.join(_record_dir, f"operation_plan_{task_uid}_{_ts}.json")
            try:
                with open(_record_path, "w", encoding="utf-8") as _f:
                    json.dump({
                        "task_uid": task_uid,
                        "instruction": task_config.get("instruction", ""),
                        "agent_mode": "plan",
                        "result": result,
                    }, _f, ensure_ascii=False, indent=2, default=str)
                log.info("Plan 模式执行记录已保存: %s", _record_path)
            except Exception as _exc:
                log.warning("保存 Plan 模式执行记录失败: %s", _exc)

        # 保存 Plan Agent 执行状态（避免 API 错误等被静默吞掉）
        if isinstance(result, dict):
            if not result.get("success", True):
                task_result["plan_agent_error"] = result.get("error", "unknown_error")
            if result.get("status"):
                task_result["plan_agent_status"] = result["status"]

        # 提取执行摘要
        execution_record = (
            result.get("execution_record", {}) if isinstance(result, dict) else {}
        )
        if execution_record:
            summary_info = extract_execution_summary(execution_record)
            task_result.update(summary_info)
        else:
            task_result["interrupted"] = True
            task_result["interrupt_reason"] = "missing_execution_record"

        # 提取 token 消耗
        raw_token = result.get("token_usage") if isinstance(result, dict) else None
        if raw_token:
            plan_usage = raw_token.get("plan_agent", {})
            gui_usage = raw_token.get("gui_agent", {})
            plan_model = raw_token.get("plan_agent_model", "")
            gui_model = raw_token.get("gui_agent_model", "unknown")
            plan_cost = calculate_cost(plan_usage, plan_model)
            gui_cost = calculate_cost(gui_usage, gui_model)
            task_result["token_usage"] = {
                "plan_agent": {
                    **plan_usage,
                    "model": plan_model,
                    "cost_usd": plan_cost["total_cost"],
                },
                "gui_agent": {
                    **gui_usage,
                    "model": gui_model,
                    "cost_usd": gui_cost["total_cost"],
                },
                "total_cost_usd": plan_cost["total_cost"] + gui_cost["total_cost"],
            }

        # ---- 中间保存：Stage 2 结果先落盘，防止 Stage 3 崩溃丢失执行记录 ----
        if output_results is not None and results_lock is not None and output_json_path:
            try:
                with results_lock:
                    output_results[task_uid] = dict(task_result)
                    with open(output_json_path, "w", encoding="utf-8") as _f:
                        json.dump(output_results, _f, ensure_ascii=False, indent=2)
                log.info("Stage 2 结果已中间保存: %s", task_uid[:8])
            except Exception as _save_exc:
                log.warning("[中间保存] 写入失败: %s", _save_exc)

        # 5. Stage 3: Operation 文件比对评估
        try:
            eval_result = stage3_evaluate_operation(
                task_config=task_config,
                agent_result=result,
                task_path=task_path,
                config=config,
                gt_cache_dir=args.gt_cache_dir,
                log=log,
                save_result_dir=getattr(args, "save_result_dir", ""),
            )
            task_result["evaluator_output"] = eval_result
        except Exception as exc:
            task_result["interrupted"] = True
            task_result["interrupt_reason"] = f"stage3_evaluate_exception: {exc}"
            task_result["evaluator_output"] = {
                "pass": False, "score": 0.0,
                "error": f"evaluator_exception: {exc}",
            }
            log.error("评估失败: %s", exc)

        log.info("任务 %s 执行完成", task_uid[:8])
        return task_result

    finally:
        # 6. 清理容器 + 释放内存 + 归还 group_id
        unregister_group_ports(group_id)
        cleanup_group_containers(config, log)
        memory_guard.release(config.num_vms)

        with _active_groups_lock:
            _active_groups.pop(group_id, None)

        available_groups.put(group_id)
        log.info("组 %d 已释放", group_id)


# ============================================================
# atexit 清理：确保异常退出时也能清理容器
# ============================================================

def _atexit_cleanup() -> None:
    """程序退出时清理所有活跃的容器组。"""
    log = logging.getLogger("pipeline.cleanup")
    with _active_groups_lock:
        groups = dict(_active_groups)

    if not groups:
        return

    log.info("程序退出，清理 %d 个活跃容器组...", len(groups))
    for group_id, config in groups.items():
        try:
            cleanup_group_containers(config, log)
        except Exception as exc:
            log.warning("清理组 %d 失败: %s", group_id, exc)


atexit.register(_atexit_cleanup)


# ============================================================
# 参数解析
# ============================================================

def parse_args() -> argparse.Namespace:
    """
    解析命令行参数。

    输出:
        argparse.Namespace
    """
    parser = argparse.ArgumentParser(
        description="Operation 文件操作任务 Pipeline — 多线程并行版本",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例:\n"
            "  # 顺序执行（默认）\n"
            "  python run_self_operation_pipeline_parallel.py\n\n"
            "  # 2 个任务并行，每任务 3 个 VM\n"
            "  python run_self_operation_pipeline_parallel.py -p 2 -n 3\n\n"
            "  # 4 个任务并行，指定 GT 缓存目录\n"
            "  python run_self_operation_pipeline_parallel.py -p 4 --gt-cache-dir /tmp/my_gt\n"
        ),
    )
    parser.add_argument(
        "-p", "--max-parallel-tasks",
        type=int, default=1,
        help="最大并发任务数（默认 1 = 顺序执行）",
    )
    parser.add_argument(
        "-n", "--vms-per-task",
        type=int, default=5,
        help="每个任务启动的 VM 数量（默认 5，可设为 1-5）",
    )
    parser.add_argument(
        "--vm-memory",
        type=str, default="1G",
        help='每个 QEMU VM 内存（默认 "1G"）',
    )
    parser.add_argument(
        "--vm-cpu-cores",
        type=str, default="1",
        help='每个 VM CPU 核数（默认 "1"）',
    )
    parser.add_argument(
        "--memory-limit-gb",
        type=float, default=48.0,
        help="容器区可用总内存上限 GiB（默认 48.0）",
    )
    parser.add_argument(
        "--vm-ip",
        type=str, default="10.1.110.143",
        help="Docker 宿主机 IP（默认 10.1.110.143）",
    )
    parser.add_argument(
        "--shared-base-dir",
        type=str, default="/home/agentlab/shared",
        help="共享目录根路径（默认 /home/agentlab/shared）",
    )
    parser.add_argument(
        "--qcow2-path",
        type=str,
        default="/home/agentlab/code/parallel-efficient-benchmark/ubuntu_env/docker_vm_data/Ubuntu.qcow2",
        help="VM 磁盘镜像路径",
    )
    parser.add_argument(
        "--docker-image",
        type=str, default="happysixd/osworld-docker-sshfs",
        help="Docker 镜像名",
    )
    parser.add_argument(
        "--gt-cache-dir",
        type=str, default=os.path.join(current_dir, "gt_cache"),
        help="GT 文件缓存目录（默认脚本同目录下的 gt_cache/）",
    )
    parser.add_argument(
        "--task-list",
        type=str, default=os.path.join(current_dir, "task_list.txt"),
        help="任务列表文件路径，每行一个 task_id（默认同目录下 task_list.txt）",
    )
    parser.add_argument(
        "--task-ids",
        nargs="*",
        help="直接指定 task_id 列表（优先于 --task-list）",
    )
    parser.add_argument(
        "--no-task-list",
        action="store_true",
        help="忽略任务列表文件，运行所有匹配的任务（动态扫描模式）",
    )
    parser.add_argument(
        "--save-result-dir",
        type=str, default="",
        help="Agent 结果文件持久化目录（按 task_id 子目录保存，用于事后复现和判断）。"
             "不指定则不保存。",
    )
    parser.add_argument(
        "--skip-completed-dir",
        type=str, default="",
        help="跳过已完成任务：指定历史结果目录路径，该目录下每个子目录名视为已完成的 task_id。"
             "支持逗号分隔多个目录。",
    )
    parser.add_argument(
        "--output-json-path",
        type=str, default="",
        help="自定义输出 JSON 路径（默认 logs/run_self_operation_pipeline_parallel.json）",
    )
    # ---------- gui_only 模式参数 ----------
    parser.add_argument(
        "--agent-mode",
        type=str, default="plan",
        choices=["plan", "gui_only"],
        help="Agent 模式：plan（Plan Agent 分解 + 多 VM）或 gui_only（单 GUI Agent + 单 VM）",
    )
    parser.add_argument(
        "--gui-agent",
        type=str, default="seed18",
        choices=["seed18", "claude", "kimi"],
        help="gui_only 模式下使用的 GUI Agent 类型（默认 seed18）",
    )
    parser.add_argument(
        "--gui-max-rounds",
        type=int, default=200,
        help="gui_only 模式下 GUI Agent 最大执行轮次（默认 200）",
    )
    parser.add_argument(
        "--gui-timeout",
        type=int, default=3600,
        help="gui_only 模式下 GUI Agent 超时时间（秒，默认 3600）",
    )
    return parser.parse_args()


# ============================================================
# 主流程
# ============================================================

def main() -> None:
    """
    主流程：多线程并行 Operation 文件操作任务调度器。

    1. 解析命令行参数
    2. 扫描 Operation 文件操作任务列表
    3. 创建 MemoryGuard
    4. 使用 ThreadPoolExecutor 并行提交任务
    5. 收集结果并写入 JSON
    """
    args = parse_args()

    # 消融实验环境变量覆盖（run_ablation.py 通过 subprocess 环境变量传递）
    _ablation_agent_mode = os.environ.get("ABLATION_AGENT_MODE", "")
    _ablation_gui_agent = os.environ.get("ABLATION_GUI_AGENT", "")
    if _ablation_agent_mode:
        args.agent_mode = _ablation_agent_mode
    if _ablation_gui_agent:
        args.gui_agent = _ablation_gui_agent

    # gui_only 模式下强制 vms_per_task=1
    if args.agent_mode == "gui_only":
        args.vms_per_task = 1

    setup_logging(args.max_parallel_tasks)
    log = logging.getLogger("pipeline.main")

    # 打印消融覆盖信息
    if _ablation_agent_mode or _ablation_gui_agent:
        log.info("消融环境变量覆盖: agent_mode=%s, gui_agent=%s",
                 _ablation_agent_mode or "(未覆盖)", _ablation_gui_agent or "(未覆盖)")

    # conda 环境检查
    required_env = os.environ.get("REQUIRED_CONDA_ENV", "")
    strict_check = os.environ.get("REQUIRED_CONDA_ENV_STRICT", "0") == "1"
    ensure_conda_env(required_env, strict=strict_check)

    log.info("=" * 80)
    log.info("Operation 文件操作任务 Pipeline — 多线程并行版本")
    log.info("  模式: %s | 并发数: %d | VM/任务: %d | VM 内存: %s | CPU: %s | 内存上限: %.1f GiB",
             args.agent_mode, args.max_parallel_tasks, args.vms_per_task,
             args.vm_memory, args.vm_cpu_cores, args.memory_limit_gb)
    if args.agent_mode == "gui_only":
        log.info("  GUI Agent: %s | 最大轮次: %d | 超时: %ds",
                 args.gui_agent, args.gui_max_rounds, args.gui_timeout)
    log.info("  GT 缓存目录: %s", args.gt_cache_dir)
    if args.save_result_dir:
        os.makedirs(args.save_result_dir, exist_ok=True)
        log.info("  结果文件保存目录: %s", args.save_result_dir)
    else:
        log.info("  结果文件保存: 未启用（使用 --save-result-dir 开启）")
    log.info("=" * 80)

    # 确定任务白名单（优先级：--task-ids > --task-list > 动态扫描全部）
    allowed_task_ids = None
    if args.task_ids:
        allowed_task_ids = set(args.task_ids)
        log.info("使用命令行指定的 %d 个任务", len(allowed_task_ids))
    elif not args.no_task_list:
        allowed_task_ids = load_task_list(args.task_list)
        if allowed_task_ids:
            log.info("从任务列表 %s 加载了 %d 个任务", args.task_list, len(allowed_task_ids))
        else:
            log.warning("任务列表为空或不存在: %s，将扫描所有匹配的任务", args.task_list)
            allowed_task_ids = None
    else:
        log.info("--no-task-list 模式：将扫描所有匹配的任务")

    # 扫描任务（按白名单过滤）
    task_items = scan_operation_tasks(TASKS_LIST_DIR, allowed_task_ids=allowed_task_ids)
    log.info("共检测到 Operation 文件操作任务数量: %d", len(task_items))

    # 跳过已完成的任务（通过历史结果目录中的子目录名判断）
    if args.skip_completed_dir:
        completed_task_ids: set = set()
        for one_dir in args.skip_completed_dir.split(","):
            one_dir = one_dir.strip()
            if not one_dir or not os.path.isdir(one_dir):
                continue
            for name in os.listdir(one_dir):
                if os.path.isdir(os.path.join(one_dir, name)):
                    completed_task_ids.add(name)
        if completed_task_ids:
            before_count = len(task_items)
            task_items = [
                (uid, path, cfg) for uid, path, cfg in task_items
                if cfg.get("task_id", "") not in completed_task_ids
            ]
            skipped = before_count - len(task_items)
            log.info("跳过已完成任务: %d 个（来自 %s）", skipped, args.skip_completed_dir)

    if not task_items:
        log.warning("未找到 Operation 文件操作任务，退出")
        return

    # 打印任务列表
    for i, (uid, _, cfg) in enumerate(task_items):
        log.info("  [%d] %s (UID: %s)", i + 1, cfg.get("task_id", ""), uid[:8])

    # 创建内存管理器
    memory_guard = MemoryGuard(args.memory_limit_gb, args.vm_memory)

    # 创建 group_id 池
    available_groups: queue.Queue = queue.Queue()
    for g in range(args.max_parallel_tasks):
        available_groups.put(g)
    log.info("已初始化 %d 个容器组槽位", args.max_parallel_tasks)

    # 启动全局防黑屏心跳守护线程
    heartbeat = GlobalScreensaverHeartbeat(vm_ip=args.vm_ip, interval_sec=180)
    heartbeat.start()

    # 结果收集
    output_results: Dict[str, Any] = {}
    results_lock = threading.Lock()
    output_json_path = os.path.abspath(
        args.output_json_path if args.output_json_path else OUTPUT_JSON_PATH
    )
    os.makedirs(os.path.dirname(output_json_path), exist_ok=True)

    # 并行调度
    completed_count = 0
    total_count = len(task_items)

    with ThreadPoolExecutor(
        max_workers=args.max_parallel_tasks,
        thread_name_prefix="Worker",
    ) as executor:
        futures = {}

        for i, (task_uid, task_path, task_config) in enumerate(task_items):
            log.info("提交任务 %d/%d | %s (UID: %s)",
                     i + 1, total_count,
                     task_config.get("task_id", ""), task_uid[:8])

            fut = executor.submit(
                run_single_task,
                task_uid, task_path, task_config,
                available_groups, args, memory_guard,
                output_results, results_lock, output_json_path,
            )
            futures[fut] = (task_uid, i + 1)

        # 收集结果
        for fut in as_completed(futures):
            task_uid, _index = futures[fut]
            try:
                task_result = fut.result()
            except Exception as exc:
                log.error("任务 %s 异常: %s", task_uid[:8], exc)
                task_result = {
                    "task_uid": task_uid,
                    "interrupted": True,
                    "interrupt_reason": f"uncaught_exception: {exc}",
                }

            with results_lock:
                output_results[task_uid] = task_result
                completed_count += 1

            # 实时持久化中间结果
            evaluator_output = task_result.get("evaluator_output")
            if task_result.get("interrupted"):
                status = "INTERRUPTED"
            elif evaluator_output and evaluator_output.get("pass"):
                status = "PASS"
            else:
                status = "FAIL"

            task_id = task_result.get("task_id", "")
            score_str = ""
            if evaluator_output and "score" in evaluator_output:
                score_str = f" (score: {evaluator_output['score']:.2f})"

            log.info(
                "任务完成 %d/%d | %s | 状态: %s%s",
                completed_count, total_count, task_id, status, score_str,
            )

            # 每完成一个任务就写一次中间结果
            try:
                with results_lock:
                    with open(output_json_path, "w", encoding="utf-8") as f:
                        json.dump(output_results, f, ensure_ascii=False, indent=2)
            except Exception as exc:
                log.warning("写入中间结果失败: %s", exc)

    # 停止心跳
    heartbeat.stop()

    # 写入最终结果
    with open(output_json_path, "w", encoding="utf-8") as f:
        json.dump(output_results, f, ensure_ascii=False, indent=2)

    log.info("=" * 80)
    log.info("全部任务执行完成 (%d/%d)", completed_count, total_count)
    log.info("输出结果文件: %s", output_json_path)
    log.info("=" * 80)

    # 统计汇总
    passed = sum(
        1 for r in output_results.values()
        if r.get("evaluator_output") and r.get("evaluator_output", {}).get("pass")
    )
    interrupted = sum(1 for r in output_results.values() if r.get("interrupted"))
    failed = total_count - passed - interrupted

    log.info("统计: 通过 %d | 失败 %d | 中断 %d | 总计 %d",
             passed, failed, interrupted, total_count)

    # 输出每个任务的详细得分
    log.info("-" * 60)
    for uid, res in sorted(output_results.items(), key=lambda x: x[1].get("task_id", "")):
        tid = res.get("task_id", uid[:8])
        ev = res.get("evaluator_output")
        if res.get("interrupted"):
            log.info("  %s: INTERRUPTED (%s)", tid, res.get("interrupt_reason", ""))
        elif ev:
            log.info("  %s: %s (score=%.2f)", tid,
                     "PASS" if ev.get("pass") else "FAIL", ev.get("score", 0))
        else:
            log.info("  %s: NO_EVAL", tid)


if __name__ == "__main__":
    main()

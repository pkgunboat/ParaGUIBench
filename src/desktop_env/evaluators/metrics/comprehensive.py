"""
综合文件比对评估模块。

对 docx / xlsx / pptx 分别提供多维度综合评分函数，同时覆盖内容和格式。
原子级比对函数分布在 docs.py、table.py、slides.py 中，本模块负责编排组合。

典型用法:
    from desktop_env.evaluators.metrics.comprehensive import (
        compare_docx_comprehensive,
        compare_xlsx_comprehensive,
        compare_pptx_comprehensive,
        compare_file_comprehensive,
    )
    score = compare_file_comprehensive(gt_path, result_path)
"""

import logging
import os
from typing import Any, Dict, List

import openpyxl
from docx import Document as DocxDocument

from desktop_env.evaluators.metrics.docs import (
    compare_docx_files,
    compare_docx_tables,
    compare_docx_paragraph_styles,
    compare_docx_run_formatting,
)
from desktop_env.evaluators.metrics.table import compare_table
from desktop_env.evaluators.metrics.slides import (
    compare_pptx_files,
    extract_all_transitions,
    compare_pptx_transitions,
)

logger = logging.getLogger("desktopenv.metric.comprehensive")


# ============================================================
# DOCX 综合评分
# ============================================================

def compare_docx_comprehensive(
    gt_path: str,
    result_path: str,
    weights: Dict[str, float] = None,
) -> float:
    """
    综合比较两个 docx 文件，从内容和格式两个维度评分。

    四个维度及默认权重:
      - text（0.4）：段落纯文本比较（compare_docx_files）
      - table（0.2）：表格数据比较（compare_docx_tables），无表格时得 1.0
      - style（0.2）：段落样式名称比较（compare_docx_paragraph_styles）
      - run_format（0.2）：run 级别格式比较（compare_docx_run_formatting）

    输入:
        gt_path: GT docx 文件路径
        result_path: 待评估 docx 文件路径
        weights: 可选的权重字典，键为 text/table/style/run_format

    输出:
        float，综合得分 0.0~1.0
    """
    if weights is None:
        weights = {"text": 0.4, "table": 0.2, "style": 0.2, "run_format": 0.2}

    # 维度 1: 段落文本
    text_score = float(compare_docx_files(result_path, gt_path))
    logger.debug("[docx] 段落文本得分: %.2f", text_score)

    # 加载文档对象（后续维度共用）
    try:
        doc_result = DocxDocument(result_path)
        doc_gt = DocxDocument(gt_path)
    except Exception as exc:
        logger.error("[docx] 无法打开文档: %s", exc)
        return 0.0

    # 维度 2: 表格内容
    if not doc_gt.tables:
        table_score = 1.0
    else:
        table_score = float(compare_docx_tables(result_path, gt_path))
    logger.debug("[docx] 表格内容得分: %.2f", table_score)

    # 维度 3: 段落样式
    style_score = compare_docx_paragraph_styles(doc_result, doc_gt)
    logger.debug("[docx] 段落样式得分: %.2f", style_score)

    # 维度 4: Run 级别格式
    run_fmt_score = compare_docx_run_formatting(doc_result, doc_gt)
    logger.debug("[docx] Run 格式得分: %.2f", run_fmt_score)

    score = (text_score * weights.get("text", 0.4)
             + table_score * weights.get("table", 0.2)
             + style_score * weights.get("style", 0.2)
             + run_fmt_score * weights.get("run_format", 0.2))
    logger.debug("[docx] 综合得分: %.2f", score)
    return score


# ============================================================
# XLSX 综合评分
# ============================================================

def compare_xlsx_comprehensive(
    gt_path: str,
    result_path: str,
    extra_style_props: List[str] = None,
    extra_row_props: List[str] = None,
    dimension_weights: Dict[str, float] = None,
) -> float:
    """
    综合比较两个 xlsx 文件，从数据值、样式、行属性三个维度加权评分。

    为每个 sheet 生成三组 rules，分别传递给 compare_table 独立评分，
    最后按维度权重加权求和。这样某个维度失败不会导致整体归零。

    三个维度及默认权重:
      - data（0.6）：单元格数据值比对
      - style（0.3）：单元格样式比对（字体加粗/颜色/斜体/下划线/背景色/数字格式/对齐）
      - row_props（0.1）：行属性比对（隐藏行等）

    输入:
        gt_path: GT xlsx 文件路径
        result_path: 待评估 xlsx 文件路径
        extra_style_props: 额外需要检查的样式属性列表（追加到默认列表）
        extra_row_props: 额外需要检查的行属性列表（追加到默认列表）
        dimension_weights: 可选的维度权重字典，键为 data/style/row_props

    输出:
        float，综合得分 0.0~1.0
    """
    if dimension_weights is None:
        dimension_weights = {"data": 0.6, "style": 0.3, "row_props": 0.1}

    try:
        gt_wb = openpyxl.load_workbook(gt_path, read_only=True)
        result_wb = openpyxl.load_workbook(result_path, read_only=True)
    except Exception as exc:
        logger.error("[xlsx] 无法打开文件: %s", exc)
        return 0.0

    gt_sheets = gt_wb.sheetnames
    result_sheets = result_wb.sheetnames
    gt_wb.close()
    result_wb.close()

    # 默认样式属性
    style_props = [
        "font_bold", "font_color", "font_italic",
        "font_underline", "bgcolor", "fgcolor",
        "number_format", "horizontal_alignment",
    ]
    if extra_style_props:
        style_props.extend(extra_style_props)

    # 默认行属性
    row_props = ["hidden"]
    if extra_row_props:
        row_props.extend(extra_row_props)

    # 按维度分别构建规则列表
    data_rules: List[Dict[str, Any]] = []
    style_rules: List[Dict[str, Any]] = []
    row_prop_rules: List[Dict[str, Any]] = []

    for i, sheet_name in enumerate(gt_sheets):
        if sheet_name in result_sheets:
            result_idx = result_sheets.index(sheet_name)
        elif i < len(result_sheets):
            result_idx = i
        else:
            logger.warning("[xlsx] GT sheet '%s' 在结果中找不到对应项", sheet_name)
            continue

        # 数据值规则
        data_rules.append({
            "type": "sheet_data",
            "sheet_idx0": result_idx,
            "sheet_idx1": f"EI{i}",
        })

        # 样式规则
        style_rules.append({
            "type": "style",
            "sheet_idx0": result_idx,
            "sheet_idx1": f"EI{i}",
            "props": style_props,
        })

        # 行属性规则
        row_prop_rules.append({
            "type": "row_props",
            "sheet_idx0": result_idx,
            "sheet_idx1": f"EI{i}",
            "props": row_props,
        })

    if not data_rules:
        logger.warning("[xlsx] 未能构建任何比对规则")
        return 0.0

    # 各维度独立评估（避免 AND 耦合：某个维度失败不影响其他维度）
    data_score = float(compare_table(
        result=result_path, expected=gt_path, rules=data_rules))
    style_score = float(compare_table(
        result=result_path, expected=gt_path, rules=style_rules))
    row_score = float(compare_table(
        result=result_path, expected=gt_path, rules=row_prop_rules))

    logger.debug("[xlsx] 数据值得分: %.2f, 样式得分: %.2f, 行属性得分: %.2f",
                 data_score, style_score, row_score)

    score = (data_score * dimension_weights.get("data", 0.6)
             + style_score * dimension_weights.get("style", 0.3)
             + row_score * dimension_weights.get("row_props", 0.1))
    logger.debug("[xlsx] 综合得分: %.2f", score)
    return score


# ============================================================
# PPTX 综合评分
# ============================================================

def compare_pptx_comprehensive(
    gt_path: str,
    result_path: str,
    content_weight: float = 0.7,
    transition_weight: float = 0.3,
) -> float:
    """
    综合比较两个 pptx 文件，包含内容格式和过渡动画两个维度。

    维度及默认权重:
      - 内容+格式（0.7）：复用 compare_pptx_files（检查文本/字体/颜色/对齐/背景/备注等）
      - 过渡动画（0.3）：逐张比较 transition 类型
        如果 GT 没有任何 transition，跳过此维度（权重全部给内容格式）

    输入:
        gt_path: GT pptx 文件路径
        result_path: 待评估 pptx 文件路径
        content_weight: 内容格式维度权重（默认 0.7）
        transition_weight: 过渡动画维度权重（默认 0.3）

    输出:
        float，综合得分 0.0~1.0
    """
    content_format_score = float(compare_pptx_files(result_path, gt_path))
    logger.debug("[pptx] 内容格式得分: %.2f", content_format_score)

    gt_transitions = extract_all_transitions(gt_path)

    # GT 没有任何 transition → 不检查此维度
    if all(t is None for t in gt_transitions):
        logger.debug("[pptx] GT 无 transition，跳过动画维度")
        return content_format_score

    transition_score = compare_pptx_transitions(gt_path, result_path)
    logger.debug("[pptx] 过渡动画得分: %.2f", transition_score)

    score = content_format_score * content_weight + transition_score * transition_weight
    logger.debug("[pptx] 综合得分: %.2f", score)
    return score


# ============================================================
# 统一入口
# ============================================================

def compare_file_comprehensive(
    gt_path: str,
    result_path: str,
    **kwargs,
) -> float:
    """
    根据文件扩展名自动选择综合比对函数。

    支持 .docx、.xlsx、.pptx 三种文件类型。
    不支持的扩展名返回 0.0。

    输入:
        gt_path: GT 文件路径
        result_path: 待评估文件路径
        **kwargs: 传递给具体比对函数的额外参数

    输出:
        float，综合得分 0.0~1.0
    """
    ext = os.path.splitext(gt_path)[1].lower()

    try:
        if ext == ".docx":
            return compare_docx_comprehensive(gt_path, result_path, **kwargs)
        elif ext == ".xlsx":
            return compare_xlsx_comprehensive(gt_path, result_path, **kwargs)
        elif ext == ".pptx":
            return compare_pptx_comprehensive(gt_path, result_path, **kwargs)
        else:
            logger.warning("不支持的文件扩展名: %s", ext)
            return 0.0
    except Exception as exc:
        logger.error("综合比对异常: %s vs %s → %s", gt_path, result_path, exc)
        return 0.0

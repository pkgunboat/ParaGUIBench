"""从最终旧源机械提取的 Operation 检查最小传递闭包。"""

from __future__ import annotations
import glob
import hashlib
import logging
import math
import os
import re
import stat
import unicodedata
from typing import Any, Dict, Optional
from docx import Document
from docx.shared import Cm
from docx.text.paragraph import Paragraph

logger = logging.getLogger("eval.operation_checks.docx")


def _run_rgb(run) -> Optional[str]:
    """读取 Word run 的显式 RGB 颜色。

    输入:
        run: python-docx Run 对象。
    输出:
        六位大写 RGB 字符串；未显式设置或使用主题色时返回 None。
    """
    color = run.font.color
    if color is None or color.rgb is None:
        return None
    value = str(color.rgb).upper()
    return value[-6:] if len(value) >= 6 else value


def _is_red_run(run) -> bool:
    """判断 Word run 是否显式设为纯红色。

    输入:
        run: python-docx Run 对象。
    输出:
        RGB 为 FF0000 时返回 True，否则返回 False。
    """
    return _run_rgb(run) == "FF0000"


def _load_document(file_path: str) -> Optional[Document]:
    """
    安全加载 docx 文件。

    输入:
        file_path: docx 文件路径
    输出:
        Document 对象；加载失败返回 None
    """
    try:
        return Document(file_path)
    except Exception:
        logger.error("Operation 检查内部事件")
        return None


def _ok(reason: str = "通过") -> Dict[str, Any]:
    """功能：执行传递闭包内的 _ok 检查步骤。

    输入参数：
        由函数签名定义；调用值仅在 evaluator 内存中使用。
    输出返回值：
        由函数签名定义；公开聚合层会删除路径、内容、gold 与原始 reason。

    旧源语义说明：
    构造通过结果。"""
    return {"pass": True, "score": 1.0, "reason": reason}


def _fail(reason: str) -> Dict[str, Any]:
    """功能：执行传递闭包内的 _fail 检查步骤。

    输入参数：
        由函数签名定义；调用值仅在 evaluator 内存中使用。
    输出返回值：
        由函数签名定义；公开聚合层会删除路径、内容、gold 与原始 reason。

    旧源语义说明：
    构造失败结果。"""
    return {"pass": False, "score": 0.0, "reason": reason}


def _partial(score: float, reason: str) -> Dict[str, Any]:
    """功能：执行传递闭包内的 _partial 检查步骤。

    输入参数：
        由函数签名定义；调用值仅在 evaluator 内存中使用。
    输出返回值：
        由函数签名定义；公开聚合层会删除路径、内容、gold 与原始 reason。

    旧源语义说明：
    构造部分通过结果。严格阈值：仅当 score 等于 1.0 才算 pass。"""
    return {"pass": score >= 1.0 - 1e-09, "score": round(score, 4), "reason": reason}


def _config_error(reason: str) -> Dict[str, Any]:
    """功能：执行传递闭包内的 _config_error 检查步骤。

    输入参数：
        由函数签名定义；调用值仅在 evaluator 内存中使用。
    输出返回值：
        由函数签名定义；公开聚合层会删除路径、内容、gold 与原始 reason。

    旧源语义说明：
    评价器配置错误（缺参数等）：score=-1 哨兵，由上层冒泡为 evaluator_error。"""
    return {"pass": False, "score": -1.0, "status": "evaluator_error", "reason": reason}


def check_max_consecutive_blank_lines(file_path: str, params: dict) -> dict:
    """
    检查文档中连续空段落数是否超过允许的最大值。

    输入:
        file_path: docx 文件路径
        params:
            max_allowed (int): 允许的最大连续空行数，默认 1
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    max_allowed = params.get("max_allowed", 1)
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    max_found = 0
    current_streak = 0
    violations = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "":
            current_streak += 1
            if current_streak > max_found:
                max_found = current_streak
            if current_streak > max_allowed:
                violations.append(i)
        else:
            current_streak = 0
    if max_found <= max_allowed:
        return _ok(f"最大连续空行 {max_found} ≤ {max_allowed}")
    return _fail(
        f"发现连续空行 {max_found} 行（允许 {max_allowed}），共 {len(violations)} 处违规"
    )


def check_font_name(file_path: str, params: dict) -> dict:
    """
    检查文档中所有 run 的字体是否为指定字体名称。

    输入:
        file_path: docx 文件路径
        params:
            font_name (str): 期望的字体名称，如 "Times New Roman"
            threshold (float): 符合比例阈值，默认 0.9（90% 以上的 run 字体正确即通过）
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    expected_font = params.get("font_name", "")
    threshold = params.get("threshold", 0.9)
    if not expected_font:
        return _config_error("参数缺少 font_name")
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    total_runs = 0
    matched_runs = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            total_runs += 1
            run_font = run.font.name
            if run_font and run_font.lower() == expected_font.lower():
                matched_runs += 1
            elif run_font is None:
                style_font = (
                    para.style.font.name if para.style and para.style.font else None
                )
                if style_font and style_font.lower() == expected_font.lower():
                    matched_runs += 1
    if total_runs == 0:
        return _ok("文档无可检查的文本 run")
    ratio = matched_runs / total_runs
    if ratio >= threshold:
        return _ok(f"字体匹配率 {ratio:.1%}（{matched_runs}/{total_runs}）")
    return _partial(
        ratio,
        f"字体匹配率 {ratio:.1%}（{matched_runs}/{total_runs}），期望 ≥ {threshold:.0%}",
    )


_WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WORD_NAMESPACES = {"w": _WORD_NAMESPACE}
_STYLE_VALUE_ATTRIBUTE = f"{{{_WORD_NAMESPACE}}}val"
_STYLE_ID_ATTRIBUTE = f"{{{_WORD_NAMESPACE}}}styleId"
_STYLE_TYPE_ATTRIBUTE = f"{{{_WORD_NAMESPACE}}}type"
_LINE_SPACING_UNSET = object()
_LINE_SPACING_NON_MULTIPLE = object()
_LINE_SPACING_INVALID = object()
_TEXT_VISIBILITY_VISIBLE = object()
_TEXT_VISIBILITY_HIDDEN = object()
_TEXT_VISIBILITY_STYLE_INVALID = object()
_ON_OFF_UNSET = object()
_ON_OFF_INVALID = object()


def _line_spacing_from_properties(properties: Any) -> object:
    """从一层 OOXML 段落属性中严格解析行距语义。

    输入参数：
        properties：``w:pPr`` 元素或 ``None``。
    输出返回值：
        ``float`` 表示 ``auto`` 的行距倍数；当前层未声明
        ``w:line`` 时返回 ``_LINE_SPACING_UNSET``；单位为磅的
        ``exact/atLeast`` 返回 ``_LINE_SPACING_NON_MULTIPLE``；非法值返回
        ``_LINE_SPACING_INVALID``。
    """

    if properties is None:
        return _LINE_SPACING_UNSET
    spacings = tuple(properties.findall("./w:spacing", _WORD_NAMESPACES))
    if not spacings:
        return _LINE_SPACING_UNSET
    if len(spacings) != 1:
        return _LINE_SPACING_INVALID
    spacing = spacings[0]
    line = spacing.get(f"{{{_WORD_NAMESPACE}}}line")
    line_rule = spacing.get(f"{{{_WORD_NAMESPACE}}}lineRule", "auto")
    if line is None:
        return (
            _LINE_SPACING_UNSET
            if spacing.get(f"{{{_WORD_NAMESPACE}}}lineRule") is None
            else _LINE_SPACING_INVALID
        )
    try:
        line_value = int(line)
    except (TypeError, ValueError, OverflowError):
        return _LINE_SPACING_INVALID
    if line_value <= 0:
        return _LINE_SPACING_INVALID
    if line_rule == "auto":
        return line_value / 240.0
    if line_rule in {"exact", "atLeast"}:
        return _LINE_SPACING_NON_MULTIPLE
    return _LINE_SPACING_INVALID


def _docdefaults_line_spacing_multiple(document: Any) -> Optional[float]:
    """从 Word ``styles.xml`` 文档默认段落属性解析行距倍数。

    输入参数：
        document：python-docx ``Document`` 对象。
    输出返回值：
        ``w:docDefaults/w:pPrDefault/w:pPr/w:spacing`` 中的行距倍数；
        缺失、非法或非正数时返回 ``None``。
    """

    properties = document.styles.element.find(
        "./w:docDefaults/w:pPrDefault/w:pPr",
        _WORD_NAMESPACES,
    )
    value = _line_spacing_from_properties(properties)
    return value if isinstance(value, float) else None


def _document_body_paragraphs(document: Any) -> tuple[Paragraph, ...]:
    """枚举 Word 主文档中所有段落，包括表格单元格与嵌套容器。

    输入参数：
        document：python-docx ``Document`` 对象。
    输出返回值：
        按 OOXML 主文档顺序包装的段落元组；不读取页眉、页脚或外部部件。
    """

    paragraph_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
    return tuple(
        Paragraph(element, document)
        for element in document.element.body.iter(paragraph_tag)
    )


def _word_text_property_structure_is_unique(document: Any) -> bool:
    """统一验证 Word 文本、样式与默认属性容器的唯一性。

    输入参数：
        document：python-docx ``Document`` 对象。
    输出返回值：
        每个段落最多一个 ``pPr``，每个 run 最多一个 ``rPr``；
        每个 style 的 ``pPr/rPr/basedOn``、docDefaults 及其
        ``pPrDefault/rPrDefault/pPr/rPr`` 均最多一个；且所有
        ``pPr`` 中的 ``spacing/pStyle/rPr`` 与 ``rPr`` 中的
        ``rStyle/vanish/webHidden`` 均最多一个时返回 ``True``。
        任一 malformed 重复结构返回 ``False``，禁止 first-match
        解析产生的评分/可见性歧义。
    """

    def _at_most_one(parent: Any, child_name: str) -> bool:
        """验证一个父元素下指定 Word 子元素不重复。

        输入参数：
            parent：OOXML 父元素；child_name：不含命名空间的标签名。
        输出返回值：
            直接子元素数量不超过一个时返回 ``True``。
        """

        return len(parent.findall(f"./w:{child_name}", _WORD_NAMESPACES)) <= 1

    styles_root = document.styles.element
    doc_defaults = tuple(styles_root.findall("./w:docDefaults", _WORD_NAMESPACES))
    if len(doc_defaults) > 1:
        return False
    if doc_defaults:
        defaults = doc_defaults[0]
        if not _at_most_one(defaults, "pPrDefault") or not _at_most_one(
            defaults,
            "rPrDefault",
        ):
            return False
        paragraph_defaults = tuple(defaults.findall("./w:pPrDefault", _WORD_NAMESPACES))
        run_defaults = tuple(defaults.findall("./w:rPrDefault", _WORD_NAMESPACES))
        if paragraph_defaults and not _at_most_one(paragraph_defaults[0], "pPr"):
            return False
        if run_defaults and not _at_most_one(run_defaults[0], "rPr"):
            return False

    for style in styles_root.findall("./w:style", _WORD_NAMESPACES):
        if not all(
            _at_most_one(style, child_name) for child_name in ("pPr", "rPr", "basedOn")
        ):
            return False
    body = document.element.body
    for paragraph in body.iter(f"{{{_WORD_NAMESPACE}}}p"):
        if not _at_most_one(paragraph, "pPr"):
            return False
    for run in body.iter(f"{{{_WORD_NAMESPACE}}}r"):
        if not _at_most_one(run, "rPr"):
            return False

    paragraph_properties = tuple(styles_root.iter(f"{{{_WORD_NAMESPACE}}}pPr")) + tuple(
        body.iter(f"{{{_WORD_NAMESPACE}}}pPr")
    )
    for properties in paragraph_properties:
        if not all(
            _at_most_one(properties, child_name)
            for child_name in ("spacing", "pStyle", "rPr")
        ):
            return False
    run_properties = tuple(styles_root.iter(f"{{{_WORD_NAMESPACE}}}rPr")) + tuple(
        body.iter(f"{{{_WORD_NAMESPACE}}}rPr")
    )
    for properties in run_properties:
        if not all(
            _at_most_one(properties, child_name)
            for child_name in ("rStyle", "vanish", "webHidden")
        ):
            return False
    return True


def _style_elements(document: Any) -> Optional[dict[str, Any]]:
    """构造 OOXML ``styles.xml`` 中唯一且非空的样式索引。

    输入参数：
        document：python-docx ``Document`` 对象。
    输出返回值：
        从精确 ``w:styleId`` 到 ``w:style`` 元素的映射；任一空 ID
        或重复 ID 会破坏样式身份，此时返回 ``None`` 以 fail closed。
    """

    styles: dict[str, Any] = {}
    for style in document.styles.element.findall("./w:style", _WORD_NAMESPACES):
        style_id = style.get(_STYLE_ID_ATTRIBUTE)
        if not isinstance(style_id, str) or not style_id or style_id in styles:
            return None
        styles[style_id] = style
    return styles


def _default_paragraph_style_id(styles: dict[str, Any]) -> object:
    """解析唯一的 OOXML 默认段落样式 ID。

    输入参数：
        styles：已验证唯一的 style ID 到元素映射。
    输出返回值：
        唯一 ``w:type=paragraph,w:default=1`` 的 style ID；无默认样式时
        返回 ``_LINE_SPACING_UNSET``；多个默认样式时返回
        ``_LINE_SPACING_INVALID``。
    """

    type_attribute = f"{{{_WORD_NAMESPACE}}}type"
    default_attribute = f"{{{_WORD_NAMESPACE}}}default"
    candidates = tuple(
        style_id
        for style_id, style in styles.items()
        if style.get(type_attribute) == "paragraph"
        and style.get(default_attribute) in {"1", "true", "on"}
    )
    if len(candidates) > 1:
        return _LINE_SPACING_INVALID
    return candidates[0] if candidates else _LINE_SPACING_UNSET


def _style_chain_line_spacing(
    style_id: str,
    styles: dict[str, Any],
    expected_style_type: str,
) -> object:
    """沿 ``basedOn`` 链解析样式行距并验证整条引用闭包。

    输入参数：
        style_id：段落或可见 run 精确引用的 style ID；
        styles：已验证唯一的样式索引。
        expected_style_type：``paragraph`` 或 ``character``；整条
            ``basedOn`` 链必须保持与引用位置一致的样式类型。
    输出返回值：
        子样式优先的首个行距语义；整链未声明时返回
        ``_LINE_SPACING_UNSET``；未知 ID、循环、非法 basedOn 或超过
        64 层时返回 ``_LINE_SPACING_INVALID``。即使子样式已声明
        行距，也会继续审计其余引用链而不 fail open。
    """

    selected: object = _LINE_SPACING_UNSET
    visited: set[str] = set()
    current_id: str | None = style_id
    for _depth in range(64):
        if current_id is None:
            return selected
        if current_id in visited:
            return _LINE_SPACING_INVALID
        style = styles.get(current_id)
        if style is None:
            return _LINE_SPACING_INVALID
        if style.get(_STYLE_TYPE_ATTRIBUTE) != expected_style_type:
            return _LINE_SPACING_INVALID
        visited.add(current_id)
        value = _line_spacing_from_properties(style.find("./w:pPr", _WORD_NAMESPACES))
        if selected is _LINE_SPACING_UNSET and value is not _LINE_SPACING_UNSET:
            selected = value
        based_on = style.find("./w:basedOn", _WORD_NAMESPACES)
        if based_on is None:
            current_id = None
            continue
        parent_id = based_on.get(_STYLE_VALUE_ATTRIBUTE)
        if not isinstance(parent_id, str) or not parent_id:
            return _LINE_SPACING_INVALID
        current_id = parent_id
    return selected if current_id is None else _LINE_SPACING_INVALID


def _referenced_run_style_ids(paragraph: Any) -> object:
    """提取段落所有可见 run 引用的精确 ``rStyle`` 闭集。

    输入参数：
        paragraph：python-docx ``Paragraph`` 对象。
    输出返回值：
        按首次出现顺序去重的 style ID 元组；任一可见 run 的
        ``rStyle`` 为空时返回 ``_LINE_SPACING_INVALID``。
    """

    style_ids: list[str] = []
    for run in paragraph._p.findall(".//w:r", _WORD_NAMESPACES):
        visible_text = "".join(
            text.text or "" for text in run.findall(".//w:t", _WORD_NAMESPACES)
        )
        if not visible_text.strip():
            continue
        rstyle = run.find("./w:rPr/w:rStyle", _WORD_NAMESPACES)
        if rstyle is None:
            continue
        style_id = rstyle.get(_STYLE_VALUE_ATTRIBUTE)
        if not isinstance(style_id, str) or not style_id:
            return _LINE_SPACING_INVALID
        if style_id not in style_ids:
            style_ids.append(style_id)
    return tuple(style_ids)


def _effective_paragraph_line_spacing_multiple(
    paragraph: Any,
    document_default: Optional[float],
    styles: Optional[dict[str, Any]],
) -> Optional[float]:
    """按 OOXML 优先级读取段落当前有效行距倍数。

    输入参数：
        paragraph：python-docx ``Paragraph`` 对象；
        document_default：已从 ``docDefaults`` 解析的文档默认行距；
        styles：``styles.xml`` 中已验证唯一的样式索引。
    输出返回值：
        优先返回段落直接属性；未设置时合并 pStyle 与
        可见 run 的 rStyle ``basedOn`` 链；全部未设置时返回
        ``document_default``。未知、循环、超深或同层冲突任一出现均返回
        ``None`` 以 fail closed。
    """

    if styles is None:
        return None
    style_values: list[object] = []
    pstyle = paragraph._p.find("./w:pPr/w:pStyle", _WORD_NAMESPACES)
    if pstyle is not None:
        paragraph_style_id = pstyle.get(_STYLE_VALUE_ATTRIBUTE)
        if not isinstance(paragraph_style_id, str) or not paragraph_style_id:
            return None
    else:
        paragraph_style_id = _default_paragraph_style_id(styles)
        if paragraph_style_id is _LINE_SPACING_INVALID:
            return None
    if isinstance(paragraph_style_id, str):
        style_values.append(
            _style_chain_line_spacing(paragraph_style_id, styles, "paragraph")
        )
    run_style_ids = _referenced_run_style_ids(paragraph)
    if run_style_ids is _LINE_SPACING_INVALID:
        return None
    for run_style_id in run_style_ids:
        style_values.append(
            _style_chain_line_spacing(run_style_id, styles, "character")
        )
    if any(value is _LINE_SPACING_INVALID for value in style_values):
        return None
    direct = _line_spacing_from_properties(
        paragraph._p.find("./w:pPr", _WORD_NAMESPACES)
    )
    if direct in {_LINE_SPACING_INVALID, _LINE_SPACING_NON_MULTIPLE}:
        return None
    if isinstance(direct, float):
        return direct
    if any(value is _LINE_SPACING_NON_MULTIPLE for value in style_values):
        return None
    inherited_values = tuple(
        value for value in style_values if isinstance(value, float)
    )
    if inherited_values:
        selected = inherited_values[0]
        if any(abs(value - selected) >= 1e-9 for value in inherited_values[1:]):
            return None
        return selected
    return document_default


def _paragraph_text_visibility(
    paragraph: Any,
    document: Any,
    styles: Optional[dict[str, Any]],
) -> object:
    """聚合段落中所有非空文本节点的三态可见性。

    输入参数：
        paragraph：python-docx ``Paragraph``；document：当前 Word 文档；
        styles：经唯一性验证的样式索引。
    输出返回值：
        任一非空节点的样式可见性无法安全确定时返回
        ``_TEXT_VISIBILITY_STYLE_INVALID``；否则任一节点可见时返回
        ``_TEXT_VISIBILITY_VISIBLE``；只有删除、移出、明确隐藏或空文本
        时返回 ``_TEXT_VISIBILITY_HIDDEN``。``w:ins``、``w:moveTo`` 与
        ``w:fldSimple`` 中的可见文本不会被 ``Paragraph.text`` 漏掉。
    """

    states = tuple(
        _element_text_visibility(text_node, document, styles)
        for text_node in paragraph._p.iter(f"{{{_WORD_NAMESPACE}}}t")
        if (text_node.text or "").strip()
    )
    if _TEXT_VISIBILITY_STYLE_INVALID in states:
        return _TEXT_VISIBILITY_STYLE_INVALID
    if _TEXT_VISIBILITY_VISIBLE in states:
        return _TEXT_VISIBILITY_VISIBLE
    return _TEXT_VISIBILITY_HIDDEN


def check_line_spacing(file_path: str, params: dict) -> dict:
    """
    检查文档段落的行距是否为指定值。

    输入:
        file_path: docx 文件路径
        params:
            spacing (float): 期望的行距倍数，如 2.0 表示双倍行距
            threshold (float): 符合比例阈值，默认 0.9
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    expected_spacing = params.get("spacing")
    threshold = params.get("threshold", 0.9)
    if expected_spacing is None:
        return _config_error("参数缺少 spacing")
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    if not _word_text_property_structure_is_unique(doc):
        return _fail("Word 文本、样式或默认属性容器存在重复声明")
    document_default = _docdefaults_line_spacing_multiple(doc)
    styles = _style_elements(doc)
    total_paras = 0
    matched_paras = 0
    for para in _document_body_paragraphs(doc):
        text_visibility = _paragraph_text_visibility(para, doc, styles)
        if text_visibility is _TEXT_VISIBILITY_HIDDEN:
            continue
        total_paras += 1
        if text_visibility is _TEXT_VISIBILITY_STYLE_INVALID:
            continue
        actual = _effective_paragraph_line_spacing_multiple(
            para,
            document_default,
            styles,
        )
        if actual is None:
            continue
        if abs(actual - expected_spacing) < 0.05:
            matched_paras += 1
    if total_paras == 0:
        return _fail("文档无可检查的非空正文段落")
    ratio = matched_paras / total_paras
    if ratio >= 1.0 - 1e-9:
        return _ok(f"行距匹配率 {ratio:.1%}（{matched_paras}/{total_paras}）")
    return _partial(
        ratio,
        f"行距匹配率 {ratio:.1%}（{matched_paras}/{total_paras}），期望所有段落均为 {expected_spacing} 倍行距（canonical 阈值 {threshold:.0%}）",
    )


def check_heading_hierarchy(file_path: str, params: dict) -> dict:
    """
    检查文档的标题层级是否符合预期规则。

    输入:
        file_path: docx 文件路径
        params:
            rules (list[dict]): 标题规则列表，每项格式：
                {
                    "pattern": str,          # 正则表达式匹配标题文本
                    "expected_style": str,   # 期望的样式名称，如 "Heading 1"
                }
            threshold (float): 符合比例阈值，默认 0.8
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    rules = params.get("rules", [])
    threshold = params.get("threshold", 0.8)
    if not rules:
        return _config_error("参数缺少 rules")
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    total_checks = 0
    matched_checks = 0
    mismatches = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        for rule in rules:
            pattern = rule.get("pattern", "")
            expected_style = rule.get("expected_style", "")
            if not pattern or not expected_style:
                continue
            if re.search(pattern, text):
                total_checks += 1
                actual_style = para.style.name if para.style else ""
                if actual_style == expected_style:
                    matched_checks += 1
                else:
                    mismatches.append(
                        {
                            "text": text[:50],
                            "expected": expected_style,
                            "actual": actual_style,
                        }
                    )
    if total_checks == 0:
        return _fail("未匹配到需要检查的标题")
    ratio = matched_checks / total_checks
    if ratio >= threshold:
        return _ok(f"标题层级匹配率 {ratio:.1%}（{matched_checks}/{total_checks}）")
    mismatch_summary = "; ".join(
        (
            f"'{m['text']}' 期望 {m['expected']} 实际 {m['actual']}"
            for m in mismatches[:5]
        )
    )
    return _partial(ratio, f"标题层级匹配率 {ratio:.1%}，不匹配: {mismatch_summary}")


def check_has_toc(file_path: str, params: dict) -> dict:
    """
    检查文档是否包含目录（Table of Contents）。

    通过检测 TOC 域代码或特定样式（TOC Heading / TOC 1-9）来判断。

    输入:
        file_path: docx 文件路径
        params: {}（无额外参数）
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    body = doc.element.body
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    instr_texts = body.findall(".//w:instrText", namespaces)
    for instr in instr_texts:
        if instr.text and "TOC" in instr.text.upper():
            return _ok("检测到 TOC 域代码")
    toc_entry_pattern = re.compile("^(?:toc\\s*[1-9]|目录\\s*[1-9])$", re.IGNORECASE)
    toc_entries = [
        para
        for para in doc.paragraphs
        if para.text.strip()
        and para.style
        and para.style.name
        and toc_entry_pattern.match(para.style.name.strip())
    ]
    if len(toc_entries) >= 2:
        return _ok(f"检测到 {len(toc_entries)} 条 TOC 目录条目")
    sdt_elements = body.findall(".//w:sdt", namespaces)
    for sdt in sdt_elements:
        sdt_pr = sdt.find("w:sdtPr", namespaces)
        if sdt_pr is not None:
            doc_part = sdt_pr.find("w:docPartGallery", namespaces)
            if doc_part is not None:
                val = doc_part.get(f"{{{namespaces['w']}}}val", "")
                if "toc" in val.lower() or "table of contents" in val.lower():
                    entry_texts = [
                        "".join(node.itertext()).strip()
                        for node in sdt.findall(".//w:p", namespaces)
                    ]
                    if len([text for text in entry_texts if text]) >= 2:
                        return _ok("检测到含目录条目的 SDT 结构")
    return _fail("未检测到目录（TOC）")


def check_batchword002_tab_indent(file_path: str, params: dict) -> dict:
    """
    BatchoperationWord-002 专用：检查每个正文段落是否有标准首行缩进。

    判断方式（满足任一即视为有 tab 缩进）：
      1. 段落文本以 '\\t' 开头
      2. 段落中第一个 run 的文本以 '\\t' 开头
      3. 段落格式 first_line_indent 大于等于 min_indent_cm

    输入:
        file_path: docx 文件路径
        params:
            min_indent_cm (float): 标准首行缩进的最小厘米数，默认 0.5。
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    skip_styles = {
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Title",
        "Subtitle",
        "TOC Heading",
    }
    min_indent = Cm(float(params.get("min_indent_cm", 0.5)))
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    total_paras = 0
    tab_paras = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        if para.style and para.style.name in skip_styles:
            continue
        total_paras += 1
        if para.text.startswith("\t"):
            tab_paras += 1
        elif para.runs and para.runs[0].text.startswith("\t"):
            tab_paras += 1
        elif (
            para.paragraph_format.first_line_indent is not None
            and para.paragraph_format.first_line_indent >= min_indent
        ):
            tab_paras += 1
    if total_paras == 0:
        return _ok("文档无需检查的正文段落")
    ratio = tab_paras / total_paras
    if ratio >= 0.8:
        return _ok(f"Tab 缩进率 {ratio:.1%}（{tab_paras}/{total_paras}）")
    return _partial(ratio, f"Tab 缩进率 {ratio:.1%}（{tab_paras}/{total_paras}）")


def _normalize_table_value(value: Any) -> str:
    """归一化 Word 表格值，用于跨格式比较。

    输入:
        value: 表格单元格文本或任务配置中的期望值。
    输出:
        去空白、数字分组逗号并转小写后的字符串。
    """
    text = str(value).strip().casefold()
    text = re.sub("(?<=\\d),(?=\\d{3}\\b)", "", text)
    return re.sub("\\s+", "", text)


def check_table_contains_expected_values(file_path: str, params: dict) -> dict:
    """检查文档表格的结构与任务数据内容。

    输入:
        file_path: docx 文件路径。
        params:
            expected_values_by_file (dict[str, list]): 按文件名配置的必须出现值。
            min_rows (int): 单个表格最少行数，默认 1。
            min_columns (int): 单个表格最少列数，默认 1。
    输出:
        包含任务期望数据的表格存在时返回标准评价结果。
    """
    expected_by_file = params.get("expected_values_by_file", {})
    min_rows = int(params.get("min_rows", 1))
    min_columns = int(params.get("min_columns", 1))
    filename = os.path.basename(file_path)
    expected_values = expected_by_file.get(filename)
    if not expected_values:
        return _config_error(f"未配置 {filename} 的期望表格值")
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    qualifying_tables = [
        table
        for table in doc.tables
        if len(table.rows) >= min_rows and len(table.columns) >= min_columns
    ]
    if not qualifying_tables:
        return _fail(f"未找到至少 {min_rows}行×{min_columns}列的表格")
    table_values = {
        _normalize_table_value(cell.text)
        for table in qualifying_tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    }
    expected_normalized = [_normalize_table_value(value) for value in expected_values]
    missing = [
        original
        for original, normalized in zip(expected_values, expected_normalized)
        if normalized not in table_values
    ]
    if not missing:
        return _ok(f"表格结构正确且包含全部 {len(expected_values)} 个期望值")
    ratio = (len(expected_values) - len(missing)) / len(expected_values)
    return _partial(
        ratio,
        f"表格缺少 {len(missing)}/{len(expected_values)} 个期望值: {', '.join(map(str, missing[:5]))}",
    )


def check_vowels_colored_red(file_path: str, params: dict) -> dict:
    """
    检查文档中元音字母（aeiouAEIOU）是否被标记为红色。

    输入:
        file_path: docx 文件路径
        params:
            threshold (float): 符合比例阈值，默认 0.8
            max_non_vowel_red_ratio (float): 辅音字母被误标红的最大比例，
                默认 0.05。
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    threshold = params.get("threshold", 0.8)
    max_non_vowel_red_ratio = params.get("max_non_vowel_red_ratio", 0.05)
    vowels = set("aeiouAEIOU")
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    total_vowels = 0
    red_vowels = 0
    total_non_vowels = 0
    red_non_vowels = 0
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            if not text:
                continue
            for char in text:
                if char in vowels:
                    total_vowels += 1
                    if _is_red_run(run):
                        red_vowels += 1
                elif char.isalpha():
                    total_non_vowels += 1
                    if _is_red_run(run):
                        red_non_vowels += 1
    if total_vowels == 0:
        return _ok("文档无元音字母可检查")
    recall = red_vowels / total_vowels
    false_positive_ratio = (
        red_non_vowels / total_non_vowels if total_non_vowels else 0.0
    )
    specificity = 1.0 - false_positive_ratio
    score = min(recall, specificity)
    if recall >= threshold and false_positive_ratio <= max_non_vowel_red_ratio:
        return _ok(
            f"元音红色标记率 {recall:.1%}，非元音误标红率 {false_positive_ratio:.1%}"
        )
    return _partial(
        score, f"元音红色率 {recall:.1%}，非元音误标红率 {false_positive_ratio:.1%}"
    )


def check_uppercase_words_have_parentheses(file_path: str, params: dict) -> dict:
    """
    检查文档中纯大写的单词（如 MAC）后是否都有括号。

    输入:
        file_path: docx 文件路径
        params:
            threshold (float): 符合比例阈值，默认 0.9
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    threshold = params.get("threshold", 0.9)
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    import re

    uppercase_word_pattern = re.compile("\\b[A-Z]{2,}\\b")
    total_uppercase = 0
    with_parentheses = 0
    details = []
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + " "
    matches = uppercase_word_pattern.findall(full_text)
    total_uppercase = len(matches)
    for para in doc.paragraphs:
        para_text = para.text
        for match in uppercase_word_pattern.finditer(para_text):
            word = match.group()
            end_pos = match.end()
            if re.match("\\s*\\(", para_text[end_pos:]):
                with_parentheses += 1
            else:
                details.append(word)
    if total_uppercase == 0:
        return _ok("文档无纯大写单词可检查")
    ratio = with_parentheses / total_uppercase
    if ratio >= threshold:
        return _ok(
            f"大写单词括号率 {ratio:.1%}（{with_parentheses}/{total_uppercase}）"
        )
    return _partial(
        ratio,
        f"大写词括号率 {ratio:.1%}（{with_parentheses}/{total_uppercase}），无括号: {', '.join(details[:3])}",
    )


def check_highlighted_words_capitalized(file_path: str, params: dict) -> dict:
    """
    检查黄色高亮的词是否有大写字母开头。

    输入:
        file_path: docx 文件路径
        params:
            highlight_color (str): 高亮颜色，默认 "FFFF00"（黄色）
            threshold (float): 符合比例阈值，默认 0.8
            require_bold (bool): 是否同时要求高亮文本加粗，默认 False。
            allowed_leading_articles (list[str]): 地名可接受的小写冠词，
                默认 ["the", "a", "an"]。
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    highlight_color = params.get("highlight_color", "FFFF00").upper()
    threshold = params.get("threshold", 0.8)
    require_bold = params.get("require_bold", False)
    allowed_articles = {
        article.lower()
        for article in params.get("allowed_leading_articles", ["the", "a", "an"])
    }
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    total_highlighted = 0
    capitalized_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            if not text:
                continue
            if run.font.highlight_color:
                hl = str(run.font.highlight_color).upper()
                if (
                    highlight_color in hl
                    or "YELLOW" in hl
                    or run.font.highlight_color == 7
                ):
                    total_highlighted += 1
                    words = re.findall("[A-Za-z][A-Za-z'\\-]*", text.strip())
                    if words and words[0].lower() in allowed_articles:
                        words = words[1:]
                    capitalized = bool(words and words[0][0].isupper())
                    correctly_bold = not require_bold or run.font.bold is True
                    if capitalized and correctly_bold:
                        capitalized_count += 1
    if total_highlighted == 0:
        return _fail("未找到黄色高亮文本")
    ratio = capitalized_count / total_highlighted
    if ratio >= threshold:
        return _ok(
            f"高亮词大写开头率 {ratio:.1%}（{capitalized_count}/{total_highlighted}）"
        )
    return _partial(
        ratio, f"大写开头率 {ratio:.1%}（{capitalized_count}/{total_highlighted}）"
    )


def check_misspelled_words_highlighted(file_path: str, params: dict) -> dict:
    """
    检查特定拼写错误词被黄色高亮且不存在额外高亮词。

    输入:
        file_path: docx 文件路径
        params:
            expected_highlights (dict): 期望被高亮的词及其所在文档的映射
                例如: {"intrenational": "travel", "conmference": "travel"}
            highlight_color (str): 高亮颜色，默认 "FFFF00"
    输出:
        期望词全部命中且黄色高亮词集合无额外词时返回满分；出现任一
        非预期高亮词返回零分，缺少期望词时按命中比例返回部分分。
    """
    expected_highlights = params.get("expected_highlights", {})
    highlight_color = params.get("highlight_color", "FFFF00").upper()
    if not expected_highlights:
        return _config_error("参数缺少 expected_highlights")
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    expected_words = {str(word).casefold() for word in expected_highlights}
    highlighted_words: set[str] = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text or not run.font.highlight_color:
                continue
            actual_color = str(run.font.highlight_color).upper()
            if (
                highlight_color in actual_color
                or "YELLOW" in actual_color
                or run.font.highlight_color == 7
            ):
                highlighted_words.update(
                    token.casefold()
                    for token in re.findall(
                        r"[A-Za-z]+(?:['’\-][A-Za-z]+)*",
                        run.text,
                    )
                )
    unexpected_words = highlighted_words - expected_words
    if unexpected_words:
        return _fail(f"检测到 {len(unexpected_words)} 个非预期黄色高亮词")
    matched_words = expected_words & highlighted_words
    total = len(expected_words)
    matched = len(matched_words)
    if matched == total:
        return _ok(f"全部 {total} 个错误词都已黄色高亮且无额外高亮词")
    ratio = matched / total
    return _partial(ratio, f"{matched}/{total} 错误词已高亮")


def check_heading_palette_and_references(result_dir: str, params: dict) -> dict:
    """检查多文档标题配色及互引文档的颜色约束。

    此检查以每个文档第一个非空段落为主标题，因为任务 fixture
    的标题是 Normal/Normal (Web) 加粗段落，而非 Heading 1..3 样式。

    输入:
        result_dir: Agent 产出根目录。
        params:
            expected_files (list[str]): 参与着色的 docx 文件名。
            palette (list[str]): 允许的六位 RGB 颜色列表。
    输出:
        所有标题均着色且互相引用的文档颜色不同时返回通过。
    """
    expected_files = params.get("expected_files", [])
    palette = {
        str(color).upper().lstrip("#")[-6:]
        for color in params.get(
            "palette", ["FF0000", "FFFF00", "0000FF", "00B050", "00FF00", "008000"]
        )
    }
    if not expected_files:
        return _config_error("参数缺少 expected_files")
    documents = {}
    for filename in expected_files:
        matches = sorted(
            (
                path
                for path in glob.glob(
                    os.path.join(result_dir, "**", filename), recursive=True
                )
                if os.path.isfile(path)
            )
        )
        if len(matches) != 1:
            return _fail(f"{filename} 期望唯一文件，实际找到 {len(matches)} 个")
        document = _load_document(matches[0])
        if document is None:
            return _fail(f"无法打开文件: {filename}")
        documents[filename] = document
    colors = {}
    for filename, document in documents.items():
        title = next(
            (paragraph for paragraph in document.paragraphs if paragraph.text.strip()),
            None,
        )
        if title is None:
            return _fail(f"{filename} 无非空主标题")
        title_colors = set()
        for run in title.runs:
            if not run.text.strip():
                continue
            color = _run_rgb(run)
            if color is None and title.style is not None:
                style_color = title.style.font.color
                if style_color is not None and style_color.rgb is not None:
                    color = str(style_color.rgb).upper()[-6:]
            if color is None:
                return _fail(f"{filename} 主标题存在未显式着色文本")
            title_colors.add(color)
        if len(title_colors) != 1:
            return _fail(f"{filename} 主标题使用了 {len(title_colors)} 种颜色")
        color = next(iter(title_colors))
        if color not in palette:
            return _fail(f"{filename} 主标题颜色 {color} 不在允许色板中")
        colors[filename] = color
    references = []
    for source_name, document in documents.items():
        source_text = "\n".join((paragraph.text for paragraph in document.paragraphs))
        for target_name in expected_files:
            if target_name == source_name:
                continue
            target_stem = os.path.splitext(target_name)[0]
            if re.search(f"\\b{re.escape(target_stem)}\\b", source_text, re.IGNORECASE):
                references.append((source_name, target_name))
                if colors[source_name] == colors[target_name]:
                    return _fail(f"互引文档 {source_name} 与 {target_name} 颜色相同")
    if not references:
        return _fail("未在文档中识别到互引关系")
    return _ok(
        f"{len(colors)} 个文档主标题使用允许色板，{len(references)} 条互引关系颜色均不同"
    )


def _on_off_property_value(properties: Any, property_name: str) -> object:
    """严格解析一层 ``w:rPr`` 的 OnOff 隐藏属性。

    输入参数：
        properties：``w:rPr`` 元素或 ``None``；property_name：
            ``vanish`` 或 ``webHidden``。
    输出返回值：
        返回布尔值、``_ON_OFF_UNSET`` 或 ``_ON_OFF_INVALID``。
        无 ``w:val`` 按 OOXML OnOff 语义为真；重复元素或未知字面值
        fail closed 为无效。
    """

    if property_name not in {"vanish", "webHidden"}:
        return _ON_OFF_INVALID
    if properties is None:
        return _ON_OFF_UNSET
    markers = tuple(properties.findall(f"./w:{property_name}", _WORD_NAMESPACES))
    if not markers:
        return _ON_OFF_UNSET
    if len(markers) != 1:
        return _ON_OFF_INVALID
    raw_value = markers[0].get(_STYLE_VALUE_ATTRIBUTE)
    if raw_value is None:
        return True
    normalized = raw_value.casefold()
    if normalized in {"1", "true", "on"}:
        return True
    if normalized in {"0", "false", "off"}:
        return False
    return _ON_OFF_INVALID


def _style_chain_on_off_value(
    style_id: str,
    styles: dict[str, Any],
    expected_style_type: str,
    property_name: str,
) -> object:
    """按子样式优先级解析一条样式链的 OnOff 值。

    输入参数：
        style_id：被 ``pStyle`` 或 ``rStyle`` 引用的样式 ID；
        styles：经唯一性验证的样式索引；
        expected_style_type：引用位置要求的 ``paragraph`` 或
            ``character`` 类型；property_name：待解析的隐藏属性。
    输出返回值：
        返回子样式优先的首个布尔值或 ``_ON_OFF_UNSET``。
        为避免已选中属性后跳过恶意引用，整条 basedOn 闭包仍会
        审计；未知、跨类型、重复/非法属性、循环或超深均返回
        ``_ON_OFF_INVALID``。
    """

    selected: object = _ON_OFF_UNSET
    visited: set[str] = set()
    current_id: str | None = style_id
    for _depth in range(64):
        if current_id is None:
            return selected
        if current_id in visited:
            return _ON_OFF_INVALID
        style = styles.get(current_id)
        if style is None or style.get(_STYLE_TYPE_ATTRIBUTE) != expected_style_type:
            return _ON_OFF_INVALID
        visited.add(current_id)
        value = _on_off_property_value(
            style.find("./w:rPr", _WORD_NAMESPACES),
            property_name,
        )
        if value is _ON_OFF_INVALID:
            return _ON_OFF_INVALID
        if selected is _ON_OFF_UNSET and value is not _ON_OFF_UNSET:
            selected = value
        based_on = style.find("./w:basedOn", _WORD_NAMESPACES)
        if based_on is None:
            current_id = None
            continue
        parent_id = based_on.get(_STYLE_VALUE_ATTRIBUTE)
        if not isinstance(parent_id, str) or not parent_id:
            return _ON_OFF_INVALID
        current_id = parent_id
    return selected if current_id is None else _ON_OFF_INVALID


def _run_is_visible_evidence(
    run: Any,
    paragraph: Any,
    document: Any,
    styles: Optional[dict[str, Any]],
) -> object:
    """以三态结果审计文本或 drawing 所属 run。

    输入参数：
        run：承载证据的 ``w:r`` 元素；paragraph：所属 ``w:p``；
        document：python-docx ``Document``；styles：唯一样式索引。
    输出返回值：
        返回 ``_TEXT_VISIBILITY_VISIBLE``、``_HIDDEN`` 或
        ``_STYLE_INVALID``。每个隐藏属性按 direct run → rStyle
        → pStyle → docDefaults 的优先级选值；所有引用链仍完整审计。
    """

    if styles is None:
        return _TEXT_VISIBILITY_STYLE_INVALID
    paragraph_properties = paragraph.find("./w:pPr", _WORD_NAMESPACES)
    style_invalid = False
    paragraph_styles = (
        tuple(paragraph_properties.findall("./w:pStyle", _WORD_NAMESPACES))
        if paragraph_properties is not None
        else ()
    )
    if len(paragraph_styles) > 1:
        style_invalid = True
        paragraph_style_id: object = _LINE_SPACING_INVALID
    elif paragraph_styles:
        paragraph_style_id = paragraph_styles[0].get(_STYLE_VALUE_ATTRIBUTE)
        if not isinstance(paragraph_style_id, str) or not paragraph_style_id:
            style_invalid = True
    else:
        paragraph_style_id = _default_paragraph_style_id(styles)
        if paragraph_style_id is _LINE_SPACING_INVALID:
            style_invalid = True
    run_styles = tuple(run.findall("./w:rPr/w:rStyle", _WORD_NAMESPACES))
    if len(run_styles) > 1:
        style_invalid = True
        run_style_id: object = _LINE_SPACING_INVALID
    elif run_styles:
        run_style_id = run_styles[0].get(_STYLE_VALUE_ATTRIBUTE)
        if not isinstance(run_style_id, str) or not run_style_id:
            style_invalid = True
    else:
        run_style_id = _LINE_SPACING_UNSET

    direct_properties = run.find("./w:rPr", _WORD_NAMESPACES)
    default_properties = document.styles.element.find(
        "./w:docDefaults/w:rPrDefault/w:rPr",
        _WORD_NAMESPACES,
    )
    hidden = False
    for property_name in ("vanish", "webHidden"):
        direct_value = _on_off_property_value(direct_properties, property_name)
        run_style_value: object = _ON_OFF_UNSET
        if isinstance(run_style_id, str):
            run_style_value = _style_chain_on_off_value(
                run_style_id,
                styles,
                "character",
                property_name,
            )
        paragraph_style_value: object = _ON_OFF_UNSET
        if isinstance(paragraph_style_id, str):
            paragraph_style_value = _style_chain_on_off_value(
                paragraph_style_id,
                styles,
                "paragraph",
                property_name,
            )
        default_value = _on_off_property_value(default_properties, property_name)
        values = (
            direct_value,
            run_style_value,
            paragraph_style_value,
            default_value,
        )
        if _ON_OFF_INVALID in values:
            style_invalid = True
        selected = next(
            (
                value
                for value in values
                if value not in {_ON_OFF_UNSET, _ON_OFF_INVALID}
            ),
            False,
        )
        if selected is True:
            hidden = True
    if hidden:
        return _TEXT_VISIBILITY_HIDDEN
    return _TEXT_VISIBILITY_STYLE_INVALID if style_invalid else _TEXT_VISIBILITY_VISIBLE


def _element_text_visibility(
    element: Any,
    document: Any,
    styles: Optional[dict[str, Any]],
) -> object:
    """将文本节点或 drawing 映射到三态可见性。

    输入参数：
        element：``w:t`` 或 ``w:drawing`` 元素；document：Word 文档；
        styles：唯一样式索引。
    输出返回值：
        节点不属于删除/移出修订，且可定位到 run 和段落时
        返回 run 的三态结果；删除/移出或非法容器按不可见处理。
    """

    run = None
    paragraph = None
    non_visible_tags = {
        f"{{{_WORD_NAMESPACE}}}del",
        f"{{{_WORD_NAMESPACE}}}moveFrom",
    }
    for ancestor in element.iterancestors():
        if ancestor.tag in non_visible_tags:
            return _TEXT_VISIBILITY_HIDDEN
        if run is None and ancestor.tag == f"{{{_WORD_NAMESPACE}}}r":
            run = ancestor
        if ancestor.tag == f"{{{_WORD_NAMESPACE}}}p":
            paragraph = ancestor
            break
    if run is None or paragraph is None:
        return _TEXT_VISIBILITY_HIDDEN
    return _run_is_visible_evidence(run, paragraph, document, styles)


def _element_is_visible_evidence(
    element: Any,
    document: Any,
    styles: Optional[dict[str, Any]],
) -> bool:
    """验证节点是明确可见、且样式图完整的证据。

    输入参数：
        element：``w:t`` 或 ``w:drawing``；document：Word 文档；
        styles：经唯一性验证的样式索引。
    输出返回值：
        仅三态结果为 ``_TEXT_VISIBILITY_VISIBLE`` 时返回 ``True``；
        隐藏或样式异常都 fail closed。该严格包装供 Word010 证据使用。
    """

    return (
        _element_text_visibility(element, document, styles) is _TEXT_VISIBILITY_VISIBLE
    )


def _single_image_slot_precedes_single_drawing(document: Any) -> bool:
    """验证主文档精确一个 ``image Slot`` 且精确一个 drawing 位于其后。

    输入参数：
        document：python-docx ``Document`` 对象。
    输出返回值：
        占位文本在单个 ``w:t`` 或多个 run 中拆分都可识别；
        全文只出现一次精确、可见的 ``[image Slot]`` 段落，只有
        一个可见 ``w:drawing``，且 drawing 的 OOXML 顺序严格在
        占位符末尾之后时返回 ``True``；其余情况 fail closed。
    """

    if not _word_text_property_structure_is_unique(document):
        return False
    body = document.element.body
    traversal_order = {element: index for index, element in enumerate(body.iter())}
    styles = _style_elements(document)
    occurrence_count = 0
    placeholder_end_orders: list[int] = []
    for paragraph in body.iter(f"{{{_WORD_NAMESPACE}}}p"):
        text_nodes = tuple(paragraph.iter(f"{{{_WORD_NAMESPACE}}}t"))
        combined = "".join(node.text or "" for node in text_nodes)
        occurrence_count += combined.count("image Slot")
        if combined.strip() != "[image Slot]":
            continue
        visible_nodes = tuple(node for node in text_nodes if (node.text or "").strip())
        if not visible_nodes or not all(
            _element_is_visible_evidence(node, document, styles)
            for node in visible_nodes
        ):
            continue
        end_offset = combined.index("[image Slot]") + len("[image Slot]") - 1
        traversed = 0
        for node in text_nodes:
            next_offset = traversed + len(node.text or "")
            if end_offset < next_offset:
                placeholder_end_orders.append(traversal_order[node])
                break
            traversed = next_offset
    drawings = tuple(body.iterfind(".//w:drawing", _WORD_NAMESPACES))
    legacy_pictures = tuple(body.iterfind(".//w:pict", _WORD_NAMESPACES))
    if (
        occurrence_count != 1
        or len(placeholder_end_orders) != 1
        or len(drawings) != 1
        or legacy_pictures
        or not _element_is_visible_evidence(drawings[0], document, styles)
    ):
        return False
    return traversal_order[drawings[0]] > placeholder_end_orders[0]


def check_image_name_matches_doc(file_path: str, params: dict) -> dict:
    """
    检查文档中插入的图片是否为同名源图，并可校验宽度。

    Word/LibreOffice 会将嵌入媒体重命名为 image1.ext，因此不再比较
    DOCX 包内部件名，而是比较嵌入二进制与文档旁同名源图的 SHA-256。

    输入:
        file_path: docx 文件路径
        params:
            expected_width_cm (float|None): 期望图片宽度；默认不校验。
            width_tolerance_cm (float): 宽度绝对容差，默认 0.15cm。
            source_extensions (list[str]): 可接受的源图扩展名。
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    expected_width_cm = params.get("expected_width_cm")
    width_tolerance_cm = float(params.get("width_tolerance_cm", 0.15))
    raw_source_extensions = params.get(
        "source_extensions", [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff"]
    )
    if (
        not isinstance(raw_source_extensions, list)
        or not raw_source_extensions
        or not all(
            isinstance(extension, str)
            and re.fullmatch(r"\.[A-Za-z0-9]{1,10}", extension)
            for extension in raw_source_extensions
        )
    ):
        return _config_error("源图扩展名配置无效")
    source_extensions = tuple(
        dict.fromkeys(extension.lower() for extension in raw_source_extensions)
    )
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    if not _single_image_slot_precedes_single_drawing(doc):
        return _fail("文档必须只有一个 image Slot，且单一图片必须位于其后")
    doc_stem = os.path.splitext(os.path.basename(file_path))[0]
    source_directory = os.path.join(os.path.dirname(file_path), "images")
    target_key = unicodedata.normalize("NFKC", doc_stem).casefold()
    source_candidates = []
    try:
        source_directory_stat = os.lstat(source_directory)
        if not stat.S_ISDIR(source_directory_stat.st_mode):
            return _fail("受控源图目录不是真实目录")
        with os.scandir(source_directory) as entries:
            for entry in entries:
                candidate_stem, candidate_extension = os.path.splitext(entry.name)
                if (
                    candidate_extension.lower() not in source_extensions
                    or unicodedata.normalize("NFKC", candidate_stem).casefold()
                    != target_key
                ):
                    continue
                if not entry.is_file(follow_symlinks=False):
                    return _fail("同名源图不是常规文件")
                source_candidates.append(entry.path)
    except OSError:
        return _fail("无法读取受控源图目录")
    if len(source_candidates) != 1:
        return _fail(f"期望唯一同名源图，实际找到 {len(source_candidates)} 个")
    if os.path.splitext(os.path.basename(source_candidates[0]))[0] != doc_stem:
        return _fail("同名源图存在大小写或 Unicode 身份冲突")
    with open(source_candidates[0], "rb") as stream:
        source_digest = hashlib.sha256(stream.read()).hexdigest()
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    drawing = doc.element.body.find(".//w:drawing", namespaces)
    if drawing is None:
        return _fail("文档中无可检查图片")
    inline_containers = tuple(drawing.findall("./wp:inline", namespaces))
    anchor_containers = tuple(drawing.findall("./wp:anchor", namespaces))
    if len(inline_containers) != 1 or anchor_containers:
        return _fail("占位符后的图片必须属于唯一 inline drawing 容器")
    container = inline_containers[0]
    drawing_properties = tuple(container.findall("./wp:docPr", namespaces))
    if len(drawing_properties) != 1:
        return _fail("图片 drawing 必须只有一个可见性声明")
    hidden_value = drawing_properties[0].get("hidden")
    if hidden_value is not None and hidden_value.casefold() not in {
        "0",
        "false",
        "off",
    }:
        return _fail("图片 drawing 不得声明隐藏")

    graphic_data = tuple(container.findall("./a:graphic/a:graphicData", namespaces))
    if len(graphic_data) != 1 or graphic_data[0].get("uri") != namespaces["pic"]:
        return _fail("图片 drawing 必须使用唯一正式 picture 图形数据")
    pictures = tuple(graphic_data[0].findall("./pic:pic", namespaces))
    if len(pictures) != 1:
        return _fail("图片 drawing 必须只有一个 picture 对象")
    picture = pictures[0]
    blip_fills = tuple(picture.findall("./pic:blipFill", namespaces))
    if len(blip_fills) != 1:
        return _fail("图片 picture 必须只有一个 blipFill 容器")
    blip_fill = blip_fills[0]
    allowed_blip_fill_children = {
        f"{{{namespaces['a']}}}blip",
        f"{{{namespaces['a']}}}srcRect",
        f"{{{namespaces['a']}}}stretch",
    }
    if any(child.tag not in allowed_blip_fill_children for child in blip_fill):
        return _fail("图片 blipFill 不得携带 tile 或额外填充结构")
    stretches = tuple(blip_fill.findall("./a:stretch", namespaces))
    if len(stretches) != 1 or stretches[0].attrib:
        return _fail("图片 blipFill 必须只有一个标准 stretch 声明")
    fill_rectangles = tuple(stretches[0].findall("./a:fillRect", namespaces))
    if len(fill_rectangles) != 1 or len(stretches[0]) != 1:
        return _fail("图片 stretch 必须只有一个 fillRect")
    fill_rectangle = fill_rectangles[0]
    if any(name not in {"l", "t", "r", "b"} for name in fill_rectangle.attrib):
        return _fail("图片 fillRect 属性无效")
    for attribute_name in ("l", "t", "r", "b"):
        try:
            fill_value = int(fill_rectangle.get(attribute_name, "0"))
        except (TypeError, ValueError, OverflowError):
            return _fail("图片 fillRect 边界无效")
        if fill_value != 0:
            return _fail("图片 fillRect 不得改变原图填充边界")
    inner_properties = tuple(picture.findall("./pic:nvPicPr/pic:cNvPr", namespaces))
    if len(inner_properties) != 1:
        return _fail("图片 picture 必须只有一个内部可见性声明")
    inner_hidden_value = inner_properties[0].get("hidden")
    if inner_hidden_value is not None and inner_hidden_value.casefold() not in {
        "0",
        "false",
        "off",
    }:
        return _fail("图片 picture 不得声明隐藏")

    extents = tuple(container.findall("./wp:extent", namespaces))
    inner_extents = tuple(picture.findall("./pic:spPr/a:xfrm/a:ext", namespaces))
    blips = tuple(blip_fill.findall("./a:blip", namespaces))
    all_blips = tuple(container.findall(".//a:blip", namespaces))
    if (
        len(extents) != 1
        or len(inner_extents) != 1
        or len(blips) != 1
        or len(all_blips) != 1
        or all_blips[0] is not blips[0]
    ):
        return _fail("图片 drawing 必须只有一个宽度与一个嵌入对象")
    if len(blips[0]) != 0:
        return _fail("嵌入图片不得携带透明度或其他视觉效果")
    source_rectangles = tuple(blip_fill.findall("./a:srcRect", namespaces))
    if len(source_rectangles) > 1:
        return _fail("嵌入图片裁剪声明不唯一")
    if source_rectangles:
        for attribute_name in ("l", "t", "r", "b"):
            raw_crop = source_rectangles[0].get(attribute_name, "0")
            try:
                crop_value = int(raw_crop)
            except (TypeError, ValueError, OverflowError):
                return _fail("嵌入图片裁剪声明无效")
            if crop_value != 0:
                return _fail("嵌入图片不得裁剪原图")
    if blips[0].get(f"{{{namespaces['r']}}}link") is not None:
        return _fail("图片 drawing 不得携带外部链接")
    relationship_id = blips[0].get(f"{{{namespaces['r']}}}embed")
    relationship = doc.part.rels.get(relationship_id) if relationship_id else None
    if (
        relationship is None
        or relationship.is_external
        or relationship.reltype
        != "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
    ):
        return _fail("图片 drawing 的嵌入关系无效")
    source_extension = os.path.splitext(source_candidates[0])[1].lower()
    expected_content_type = {
        ".bmp": "image/bmp",
        ".gif": "image/gif",
        ".jpeg": "image/jpeg",
        ".jpg": "image/jpeg",
        ".png": "image/png",
        ".tif": "image/tiff",
        ".tiff": "image/tiff",
    }.get(source_extension)
    if expected_content_type is None:
        return _config_error("源图扩展名缺少固定媒体类型映射")
    if relationship.target_part.content_type != expected_content_type:
        return _fail("图片 drawing 的嵌入媒体类型与源图不一致")
    digest = hashlib.sha256(relationship.target_part.blob).hexdigest()
    if digest != source_digest:
        return _fail("嵌入图片与文档同名源图内容不一致")
    try:
        outer_width = int(extents[0].get("cx", "0"))
        outer_height = int(extents[0].get("cy", "0"))
        inner_width = int(inner_extents[0].get("cx", "0"))
        inner_height = int(inner_extents[0].get("cy", "0"))
    except (TypeError, ValueError, OverflowError):
        return _fail("图片几何尺寸无效")
    if (
        outer_width <= 0
        or outer_height <= 0
        or inner_width != outer_width
        or inner_height != outer_height
    ):
        return _fail("图片内外几何尺寸必须一致且为正数")
    try:
        embedded_image = relationship.target_part.image
        pixel_width = int(embedded_image.px_width)
        pixel_height = int(embedded_image.px_height)
        horizontal_dpi = int(embedded_image.horz_dpi)
        vertical_dpi = int(embedded_image.vert_dpi)
        if min(pixel_width, pixel_height, horizontal_dpi, vertical_dpi) <= 0:
            return _fail("嵌入图片缺少有效像素或 DPI 尺寸")
        expected_height = (
            outer_width * pixel_height * horizontal_dpi / (pixel_width * vertical_dpi)
        )
    except (AttributeError, TypeError, ValueError, OverflowError, ZeroDivisionError):
        return _fail("无法解析嵌入图片的原始纵横比")
    aspect_tolerance = max(2.0, expected_height * 0.001)
    if (
        not math.isfinite(expected_height)
        or expected_height <= 0
        or abs(outer_height - expected_height) > aspect_tolerance
    ):
        return _fail("图片几何尺寸必须保留源图纵横比")
    if expected_width_cm is not None:
        try:
            expected_width = float(expected_width_cm)
            tolerance_width = float(width_tolerance_cm)
            if (
                not math.isfinite(expected_width)
                or expected_width <= 0
                or not math.isfinite(tolerance_width)
                or tolerance_width < 0
            ):
                return _config_error("图片宽度或容差配置无效")
            expected_emu = int(Cm(expected_width))
            tolerance_emu = int(Cm(tolerance_width))
        except (TypeError, ValueError, OverflowError):
            return _config_error("图片宽度或容差配置无效")
        if abs(outer_width - expected_emu) > tolerance_emu:
            return _fail(f"未找到宽度约 {float(expected_width_cm):g}cm 的匹配图片")
    return _ok("已嵌入内容与宽度匹配的同名源图")


def check_docx_word_count(file_path: str, params: dict) -> dict:
    """
    检查Word文档的字数是否大于指定值。

    输入:
        file_path: docx 文件路径
        params:
            min_words (int): 最少字数要求，默认 100
            exclude_headings (bool): 是否排除标题/Title 样式段落的词数，
                默认 False（保持历史行为）。开放写作任务建议开启，
                避免模板自带的标题词数白送基础分。
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    min_words = params.get("min_words", 100)
    exclude_headings = params.get("exclude_headings", False)
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    word_count = 0
    for para in doc.paragraphs:
        style_name = para.style.name or "" if para.style is not None else ""
        if exclude_headings and (
            style_name.startswith("Heading") or style_name == "Title"
        ):
            continue
        words = para.text.split()
        word_count += len(words)
    scope = "正文" if exclude_headings else "文档"
    if word_count >= min_words:
        return _ok(f"{scope}字数 {word_count} >= {min_words}")
    ratio = word_count / min_words
    return _partial(ratio, f"{scope}字数 {word_count} < {min_words}")


def check_headings_have_body(file_path: str, params: dict) -> dict:
    """
    检查文档中每个标题（Heading 样式）之后是否有实质正文。

    用于"补全各小节正文"类开放写作任务：仅靠总词数无法约束内容写在
    哪里——agent 把全部文字堆在一个小节下、或写在标题外的任意位置都
    能通过词数检查。本检查按小节统计：每个 Heading 与下一个 Heading
    （或文档结尾）之间的非标题段落计为该小节正文，
    得分 = 正文词数达标的小节数 / 小节总数。

    输入:
        file_path: docx 文件路径
        params:
            min_words_per_heading (int): 每个小节正文最少词数，默认 15
            heading_style_prefix (str): 标题样式前缀，默认 "Heading"
                （Title 样式视为文档题目，不计入小节）
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    min_words = params.get("min_words_per_heading", 15)
    heading_prefix = params.get("heading_style_prefix", "Heading")
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    sections = []
    current = None
    for para in doc.paragraphs:
        style_name = para.style.name or "" if para.style is not None else ""
        if style_name.startswith(heading_prefix):
            if current is not None:
                sections.append(tuple(current))
            current = [para.text.strip(), 0]
        elif style_name != "Title" and current is not None:
            current[1] += len(para.text.split())
    if current is not None:
        sections.append(tuple(current))
    if not sections:
        return _fail(
            f"文档中没有 '{heading_prefix}' 样式的标题（可能被删除或样式丢失）"
        )
    satisfied = [s for s in sections if s[1] >= min_words]
    ratio = len(satisfied) / len(sections)
    lacking = [f"'{h[:30]}'({w}词)" for h, w in sections if w < min_words]
    if ratio >= 1.0:
        return _ok(f"全部 {len(sections)} 个小节均有 >= {min_words} 词正文")
    return _partial(
        ratio,
        f"{len(satisfied)}/{len(sections)} 个小节正文达标，不足的小节: {', '.join(lacking[:5])}",
    )


def check_docx_has_hyperlink(file_path: str, params: dict) -> dict:
    """检查 Word 文档的超链接数量、提示段落覆盖率与目标。

    输入:
        file_path: docx 文件路径。
        params:
            threshold (float): 兼容旧规则的链接数/段落数阈值，默认 0.5。
            cue_pattern (str|None): 需要覆盖的提示段落正则；默认识别
                "see ... .pptx Page N" 形式。
            require_cue_coverage (bool): 是否要求每个提示段落都有链接，
                默认 False。
            validate_cited_target (bool): 是否要求链接目标文件名与段落
                引用的 pptx 名称一致，默认 False。
    输出:
        {"pass": bool, "score": float, "reason": str}
    """
    threshold = params.get("threshold", 0.5)
    require_cue_coverage = params.get("require_cue_coverage", False)
    validate_cited_target = params.get("validate_cited_target", False)
    cue_pattern = re.compile(
        params.get(
            "cue_pattern",
            "(?P<target>[A-Za-z0-9_.\\-]+\\.pptx)\\s+Page\\s+(?P<page>\\d+)",
        ),
        re.IGNORECASE,
    )
    doc = _load_document(file_path)
    if doc is None:
        return _fail(f"无法打开文件: {file_path}")
    body = doc.element.body
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    hyperlinks = body.findall(".//w:hyperlink", namespaces)
    link_count = len(hyperlinks)
    if link_count == 0:
        return _fail("文档中不存在超链接")
    if require_cue_coverage:
        cue_count = 0
        covered_count = 0
        failures = []
        for paragraph in body.findall(".//w:p", namespaces):
            paragraph_text = "".join(
                (node.text or "" for node in paragraph.findall(".//w:t", namespaces))
            )
            citation = cue_pattern.search(paragraph_text)
            if citation is None:
                continue
            cue_count += 1
            paragraph_links = paragraph.findall(".//w:hyperlink", namespaces)
            targets = []
            for hyperlink in paragraph_links:
                relationship_id = hyperlink.get(f"{{{namespaces['r']}}}id")
                if relationship_id and relationship_id in doc.part.rels:
                    targets.append(str(doc.part.rels[relationship_id].target_ref))
            cited_name = os.path.basename(
                citation.groupdict().get("target", "").strip()
            )
            target_matches = any(
                (
                    os.path.basename(target.split("#", 1)[0]).casefold()
                    == cited_name.casefold()
                    for target in targets
                )
            )
            if paragraph_links and (not validate_cited_target or target_matches):
                covered_count += 1
            else:
                failures.append(cited_name or paragraph_text[:40])
        if cue_count == 0:
            return _fail("未识别到需要插入链接的 PPT 页面提示")
        if covered_count == cue_count:
            return _ok(f"全部 {cue_count} 个 PPT 页面提示均有正确文件链接")
        return _partial(
            covered_count / cue_count,
            f"仅覆盖 {covered_count}/{cue_count} 个提示段落，缺失或目标错误: {', '.join(failures[:3])}",
        )
    total_paras = len(
        [paragraph for paragraph in doc.paragraphs if paragraph.text.strip()]
    )
    if total_paras == 0:
        return _ok(f"文档无段落，但包含 {link_count} 个超链接")
    ratio = min(1.0, link_count / total_paras)
    if ratio >= threshold:
        return _ok(f"文档包含 {link_count} 个超链接")
    return _partial(ratio, f"仅 {link_count} 个超链接")


DOCX_CHECKS = {
    "check_heading_hierarchy": check_heading_hierarchy,
    "check_batchword002_tab_indent": check_batchword002_tab_indent,
    "check_max_consecutive_blank_lines": check_max_consecutive_blank_lines,
    "check_highlighted_words_capitalized": check_highlighted_words_capitalized,
    "check_misspelled_words_highlighted": check_misspelled_words_highlighted,
    "check_has_toc": check_has_toc,
    "check_table_contains_expected_values": check_table_contains_expected_values,
    "check_vowels_colored_red": check_vowels_colored_red,
    "check_font_name": check_font_name,
    "check_line_spacing": check_line_spacing,
    "check_image_name_matches_doc": check_image_name_matches_doc,
    "check_heading_palette_and_references": check_heading_palette_and_references,
    "check_uppercase_words_have_parentheses": check_uppercase_words_have_parentheses,
    "check_docx_has_hyperlink": check_docx_has_hyperlink,
    "check_headings_have_body": check_headings_have_body,
    "check_docx_word_count": check_docx_word_count,
}

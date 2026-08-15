"""Word-009/010 经 AttemptRunner 与 RunStore 的纵向闭集测试。"""

from __future__ import annotations

import hashlib
from dataclasses import replace
import json
from pathlib import Path
import tempfile
from typing import Any

import pytest

from paraguibench.agents import AgentRunResult
from paraguibench.benchmark import PreparedTask
from paraguibench.evaluation.operation import (
    OPERATION_PROTOCOL_ID,
    OperationEvaluationError,
    WordTextBaseline,
    WordTextInputFile,
    capture_word_text_baseline,
)
from paraguibench.integrations.osworld.operation_artifacts import (
    OperationArtifactSnapshot,
)
from paraguibench.runstore import EvaluationOutcome, RunStore
from paraguibench.runtime.attempt_runner import AttemptRunner
from paraguibench.runtime.evaluators import build_task_evaluator
from tests.runstore._audit import (
    synthetic_run_version_vector,
    synthetic_task_audit,
)


pytest.importorskip("docx")

from docx import Document


_REPO_ROOT = Path(__file__).resolve().parents[2]
_WORD009_TASK_ID = "Operation-FileOperate-BatchOperationWord-009"


class _SnapshotEnvironment:
    """为 AttemptRunner 提供已冻结 Operation 快照的最小环境。"""

    def __init__(
        self,
        snapshot: OperationArtifactSnapshot,
        *,
        baseline: WordTextBaseline | None = None,
    ) -> None:
        """绑定当前 Attempt 唯一快照。

        输入参数：
            snapshot：已由测试构造的临时 artifact 快照；
            baseline：可选显式 pre DTO；省略时构造正式身份的
            合成 Word-009 baseline。
        输出返回值：
            无；环境在 ``close`` 阶段取得清理责任。
        """

        self._snapshot = snapshot
        self._baseline = baseline or _formal_word009_baseline()

    def start(self) -> None:
        """启动合成环境。

        输入参数：无。
        输出返回值：无；本合成环境无外部资源。
        """

    def prepare(self, task: dict[str, Any]) -> None:
        """验证 AttemptRunner 传入的 trusted task 身份。

        输入参数：
            task：完整 canonical Operation 任务。
        输出返回值：无；身份漂移由断言显式暴露。
        """

        assert task["task_id"] == self._snapshot.task_id

    def operation_artifact_snapshot(
        self,
        task_id: str,
        protocol_id: str,
    ) -> OperationArtifactSnapshot:
        """返回与 evaluator 请求身份一致的冻结快照。

        输入参数：
            task_id/protocol_id：evaluator adapter 构造时固定的任务与协议。
        输出返回值：
            当前环境拥有的 ``OperationArtifactSnapshot``。
        """

        assert task_id == self._snapshot.task_id
        assert protocol_id == self._snapshot.protocol_id
        return self._snapshot

    def operation_word_text_baseline(
        self,
        task_id: str,
        protocol_id: str,
    ) -> WordTextBaseline:
        """返回 runtime evaluator 请求的 prepare 前 typed baseline。

        输入参数：
            task_id/protocol_id：evaluator adapter 固定的任务与协议。
        输出返回值：
            构造阶段注入的同一 evaluator-only DTO。
        """

        assert task_id == self._snapshot.task_id
        assert protocol_id == self._snapshot.protocol_id
        return self._baseline

    def close(self) -> None:
        """清理当前 Attempt 拥有的快照目录。

        输入参数：无。
        输出返回值：无；重复调用由快照幂等处理。
        """

        self._snapshot.close()


class _SnapshotCallSpyEnvironment(_SnapshotEnvironment):
    """记录 evaluator 是否在 baseline 失败前读取 post 快照。"""

    def __init__(
        self,
        snapshot: OperationArtifactSnapshot,
        *,
        baseline: WordTextBaseline,
    ) -> None:
        """绑定对抗 baseline 与延迟读取的 post 快照。

        输入参数：
            snapshot：若顺序错误才会被读取的 post 快照；
            baseline：需在任何 post I/O 前拒绝的 typed DTO。
        输出返回值：
            无；初始化读取计数为零。
        """

        super().__init__(snapshot, baseline=baseline)
        self.snapshot_calls = 0

    def operation_artifact_snapshot(
        self,
        task_id: str,
        protocol_id: str,
    ) -> OperationArtifactSnapshot:
        """记录并返回 post 快照。

        输入参数：
            task_id/protocol_id：evaluator 请求的固定身份。
        输出返回值：
            父类拥有的 ``OperationArtifactSnapshot``。
        """

        self.snapshot_calls += 1
        return super().operation_artifact_snapshot(task_id, protocol_id)


class _MissingBaselineSpyEnvironment(_SnapshotCallSpyEnvironment):
    """模拟未装配 typed baseline seam 的环境。"""

    operation_word_text_baseline = None


class _ClaimingAgent:
    """返回不得参与 Operation 评价或持久化的终端文本。"""

    def __init__(self, final_output: str, *, termination: str = "finished") -> None:
        """保存对抗性 Agent final text。

        输入参数：
            final_output：宣称完成且包含敏感哨兵的文本。
            termination：Agent 声明的终止类型；默认为稳定的
                ``finished``，负向用例可传入自由文本。
        输出返回值：无。
        """

        self._final_output = final_output
        self._termination = termination

    def run(
        self,
        task_view: dict[str, Any],
        environment: object,
    ) -> AgentRunResult:
        """返回一步结束的合法 Agent 结果。

        输入参数：
            task_view：不含 eval_rules 的 Agent 投影；environment：未读取的存活环境。
        输出返回值：
            包含测试哨兵 final text 的 ``AgentRunResult``。
        """

        del environment
        assert task_view["task_id"] == _WORD009_TASK_ID
        assert "eval_rules" not in task_view
        return AgentRunResult(
            final_output=self._final_output,
            step_count=1,
            termination=self._termination,
        )


def _load_word009_task() -> dict[str, Any]:
    """读取 Word-009 canonical 任务。

    输入参数：无。
    输出返回值：完整 trusted task 映射。
    """

    path = _REPO_ROOT / "benchmark/tasks" / f"{_WORD009_TASK_ID}.json"
    return json.loads(path.read_text(encoding="utf-8"))


def _formal_word009_baseline() -> WordTextBaseline:
    """构造具有正式 manifest/路径身份的合成 pre DTO。

    输入参数：无。
    输出返回值：
        四份可解析 DOCX 的 typed baseline；临时文件在返回前
        已关闭，DTO 不持有原文或文件句柄。
    """

    manifest_path = (
        _REPO_ROOT / "benchmark/assets/manifests" / f"{_WORD009_TASK_ID}.json"
    )
    manifest_payload = manifest_path.read_bytes()
    manifest = json.loads(manifest_payload)
    with tempfile.TemporaryDirectory(prefix="paraguibench-word009-pre-") as temporary:
        # tempfile 在 macOS 返回 /var 系统别名；测试先固定
        # 可信刚创建目录的 canonical root，再验证严格 nofollow。
        root = Path(temporary).resolve(strict=True)
        files: list[WordTextInputFile] = []
        for entry in manifest["files"]:
            document = Document()
            document.add_paragraph("ORIGINAL VISIBLE BODY")
            path = root / entry["path"]
            document.save(path)
            payload = path.read_bytes()
            files.append(
                WordTextInputFile(
                    path=entry["path"],
                    size=len(payload),
                    sha256=hashlib.sha256(payload).hexdigest(),
                    is_docx=True,
                )
            )
        return capture_word_text_baseline(
            task_id=_WORD009_TASK_ID,
            protocol_id=OPERATION_PROTOCOL_ID,
            manifest_sha256=hashlib.sha256(manifest_payload).hexdigest(),
            source_root=root,
            files=tuple(files),
        )


def _prepared_word009_task() -> PreparedTask:
    """构造 Word-009 trusted/agent/audit 三投影任务。

    输入参数：无。
    输出返回值：
        evaluator 可见完整规则、Agent 仅见指令的 ``PreparedTask``。
    """

    task = _load_word009_task()
    return PreparedTask(
        trusted_task=task,
        agent_task={"task_id": task["task_id"], "instruction": task["instruction"]},
        audit_metadata=synthetic_task_audit(
            task["task_id"],
            task_uid=task["task_uid"],
            task_type=task["task_type"],
            task_source=task["task_source"],
            task_tag=task["task_tag"],
        ),
    )


def _start_attempt(store: RunStore, prepared: PreparedTask):
    """在合成 RunStore 中建立 Word-009 Attempt。

    输入参数：
        store：空的测试 RunStore；prepared：Word-009 三投影任务。
    输出返回值：
        已登记且可交给 AttemptRunner 的 ``TaskAttempt``。
    """

    store.start_run(
        run_id="run-word009-hardening",
        run_record={"environment_id": "synthetic-osworld"},
        version_vector=synthetic_run_version_vector(),
    )
    return store.start_attempt(
        run_id="run-word009-hardening",
        task_id=_WORD009_TASK_ID,
        attempt_id="attempt-001",
        task_record=prepared.audit_metadata,
    )


def _missing_document_snapshot() -> OperationArtifactSnapshot:
    """构造只含三份正确双倍行距文档的 Word-009 快照。

    输入参数：无。
    输出返回值：
        文件数固定为 3、缺少第四份 pinned 文档的可清理快照。
    """

    temporary_directory = tempfile.TemporaryDirectory(
        prefix="paraguibench-word009-hardening-"
    )
    root = Path(temporary_directory.name)
    names = (
        "Introduction to Artificial Intelligence.docx",
        "Research on Multi.docx",
        "The Quiet Station.docx",
    )
    for name in names:
        document = Document()
        paragraph = document.add_paragraph("visible body")
        paragraph.paragraph_format.line_spacing = 2.0
        document.save(root / name)
    return OperationArtifactSnapshot(
        task_id=_WORD009_TASK_ID,
        protocol_id=OPERATION_PROTOCOL_ID,
        file_count=len(names),
        temporary_directory=temporary_directory,
    )


def _complete_document_snapshot(text: str) -> OperationArtifactSnapshot:
    """构造四路径闭集完整的 Word-009 post 快照。

    输入参数：
        text：每份文档的唯一可见正文。
    输出返回值：
        四文档均为直接双倍行距的 owned 临时快照。
    """

    temporary_directory = tempfile.TemporaryDirectory(
        prefix="paraguibench-word009-complete-"
    )
    root = Path(temporary_directory.name)
    names = tuple(entry["path"] for entry in _load_word009_manifest()["files"])
    for name in names:
        document = Document()
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.line_spacing = 2.0
        document.save(root / name)
    return OperationArtifactSnapshot(
        task_id=_WORD009_TASK_ID,
        protocol_id=OPERATION_PROTOCOL_ID,
        file_count=len(names),
        temporary_directory=temporary_directory,
    )


def _load_word009_manifest() -> dict[str, Any]:
    """读取 Word-009 正式 input manifest。

    输入参数：无。
    输出返回值：包含固定四 DOCX 路径顺序的 JSON 映射。
    """

    path = _REPO_ROOT / "benchmark/assets/manifests" / f"{_WORD009_TASK_ID}.json"
    return json.loads(path.read_text(encoding="utf-8"))


def test_missing_document_final_text_cannot_change_fixed_runtime_failure(
    tmp_path: Path,
) -> None:
    """验证 AttemptRunner/RunStore 保留固定分母与原因码，忽略 final text。

    输入参数：
        tmp_path：pytest 提供的 RunStore 隔离根。
    输出返回值：
        无；三文档快照必须以 0 分 FAILED 落盘，固定 contract
        reason 可持久化，文件名、路径与 Agent final text 不得落盘。
    """

    prepared = _prepared_word009_task()
    snapshot = _missing_document_snapshot()
    store = RunStore(tmp_path)
    attempt = _start_attempt(store, prepared)
    evaluator = build_task_evaluator(
        prepared.trusted_task,
        evaluation_protocol=OPERATION_PROTOCOL_ID,
    )
    final_sentinel = "CLAIMED COMPLETE PRIVATE WORD009 FINAL"

    result = AttemptRunner(store).run(
        attempt=attempt,
        prepared_task=prepared,
        environment=_SnapshotEnvironment(snapshot),
        agent=_ClaimingAgent(final_sentinel),
        evaluator=evaluator,
    )

    assert result.evaluation_outcome is EvaluationOutcome.FAILED
    assert result.score == 0.0
    persisted = b"\n".join(
        path.read_bytes() for path in tmp_path.rglob("*") if path.is_file()
    )
    assert b"ARTIFACT_CONTRACT_MISMATCH" in persisted
    for forbidden in (
        final_sentinel.encode("utf-8"),
        b"Introduction to Artificial Intelligence.docx",
        b"The Silent Library.docx",
        str(snapshot).encode("utf-8"),
    ):
        assert forbidden not in persisted


def test_text_fidelity_mismatch_is_fixed_failed_zero_and_private(
    tmp_path: Path,
) -> None:
    """验证可比较的正文改字以固定 FAIL/0 脱敏落盘。

    输入参数：
        tmp_path：RunStore 私有根。
    输出返回值：
        无；post 四文档路径和双倍行距都正确，但与 pre
        typed 正文不同时必须 ``TEXT_FIDELITY_MISMATCH``、FAILED/0；
        Agent final text、正文、文件名与 baseline repr 均不得持久化。
    """

    prepared = _prepared_word009_task()
    private_text = "TAMPERED PRIVATE VISIBLE BODY"
    snapshot = _complete_document_snapshot(private_text)
    store = RunStore(tmp_path)
    attempt = _start_attempt(store, prepared)
    evaluator = build_task_evaluator(
        prepared.trusted_task,
        evaluation_protocol=OPERATION_PROTOCOL_ID,
    )
    final_sentinel = "PRIVATE FINAL CLAIM FOR FIDELITY"

    result = AttemptRunner(store).run(
        attempt=attempt,
        prepared_task=prepared,
        environment=_SnapshotEnvironment(snapshot),
        agent=_ClaimingAgent(final_sentinel),
        evaluator=evaluator,
    )

    assert result.evaluation_outcome is EvaluationOutcome.FAILED
    assert result.score == 0.0
    persisted = b"\n".join(
        path.read_bytes() for path in tmp_path.rglob("*") if path.is_file()
    )
    assert b"TEXT_FIDELITY_MISMATCH" in persisted
    for forbidden in (
        private_text.encode("utf-8"),
        final_sentinel.encode("utf-8"),
        b"Introduction to Artificial Intelligence.docx",
        b"WordTextBaseline",
    ):
        assert forbidden not in persisted


def test_forged_baseline_shape_is_error_null_and_persists_only_type(
    tmp_path: Path,
) -> None:
    """验证伪造 typed DTO 只产生固定 ERROR/null 和异常类型。

    输入参数：
        tmp_path：RunStore 与 post 快照私有根。
    输出返回值：
        无；baseline 内部 digest 伪造为 ``None`` 时 evaluator
        以 ``OperationEvaluationError`` 结束，RunStore 不含错误码、
        DTO repr、文本/路径或 Agent final text，分数保持 null。
    """

    prepared = _prepared_word009_task()
    snapshot = _complete_document_snapshot("ORIGINAL VISIBLE BODY")
    baseline = _formal_word009_baseline()
    forged_document = replace(baseline.documents[0], digest=None)
    forged = replace(
        baseline,
        documents=(forged_document, *baseline.documents[1:]),
    )
    store = RunStore(tmp_path / "runstore")
    attempt = _start_attempt(store, prepared)
    evaluator = build_task_evaluator(
        prepared.trusted_task,
        evaluation_protocol=OPERATION_PROTOCOL_ID,
    )
    final_sentinel = "FORGED BASELINE PRIVATE FINAL"

    environment = _SnapshotCallSpyEnvironment(snapshot, baseline=forged)
    with pytest.raises(OperationEvaluationError) as captured:
        AttemptRunner(store).run(
            attempt=attempt,
            prepared_task=prepared,
            environment=environment,
            agent=_ClaimingAgent(final_sentinel),
            evaluator=evaluator,
        )

    assert captured.value.code == "WORD_TEXT_FIDELITY_INVALID"
    assert environment.snapshot_calls == 0
    persisted = b"\n".join(
        path.read_bytes()
        for path in (tmp_path / "runstore").rglob("*")
        if path.is_file()
    )
    assert b"OperationEvaluationError" in persisted
    for forbidden in (
        b"WORD_TEXT_FIDELITY_INVALID",
        b"WordTextBaseline",
        final_sentinel.encode("utf-8"),
        b"Introduction to Artificial Intelligence.docx",
    ):
        assert forbidden not in persisted


def test_missing_baseline_seam_errors_before_post_snapshot(
    tmp_path: Path,
) -> None:
    """验证 baseline seam 缺失时不得触发 post 捕获。

    输入参数：
        tmp_path：RunStore 与合成 post 快照的私有根。
    输出返回值：
        无；Attempt 必须 ERROR/null，post snapshot 读取计数为零，
        RunStore 仅保留异常类型而不保留 Agent final text。
    """

    prepared = _prepared_word009_task()
    snapshot = _complete_document_snapshot("ORIGINAL VISIBLE BODY")
    environment = _MissingBaselineSpyEnvironment(
        snapshot,
        baseline=_formal_word009_baseline(),
    )
    store = RunStore(tmp_path / "runstore")
    attempt = _start_attempt(store, prepared)
    evaluator = build_task_evaluator(
        prepared.trusted_task,
        evaluation_protocol=OPERATION_PROTOCOL_ID,
    )
    final_sentinel = "MISSING BASELINE PRIVATE FINAL"

    with pytest.raises(TypeError):
        AttemptRunner(store).run(
            attempt=attempt,
            prepared_task=prepared,
            environment=environment,
            agent=_ClaimingAgent(final_sentinel),
            evaluator=evaluator,
        )

    assert environment.snapshot_calls == 0
    persisted = b"\n".join(
        path.read_bytes()
        for path in (tmp_path / "runstore").rglob("*")
        if path.is_file()
    )
    assert b"TypeError" in persisted
    assert final_sentinel.encode("utf-8") not in persisted


def test_agent_controlled_termination_free_text_never_reaches_runstore(
    tmp_path: Path,
) -> None:
    """验证 Agent 不能把 final text 复制到 termination 后持久化。

    输入参数：
        tmp_path：pytest 提供的 RunStore 隔离根。
    输出返回值：
        无；AttemptRunner 必须在 worker completed 事件之前拒绝非固定
        termination，RunStore 只能保留异常类型，不得保留哨兵。
    """

    prepared = _prepared_word009_task()
    snapshot = _missing_document_snapshot()
    store = RunStore(tmp_path)
    attempt = _start_attempt(store, prepared)
    evaluator = build_task_evaluator(
        prepared.trusted_task,
        evaluation_protocol=OPERATION_PROTOCOL_ID,
    )
    final_sentinel = "PRIVATE AGENT FINAL TEXT SENTINEL"
    termination_sentinel = f"finished {final_sentinel}"

    with pytest.raises(ValueError):
        AttemptRunner(store).run(
            attempt=attempt,
            prepared_task=prepared,
            environment=_SnapshotEnvironment(snapshot),
            agent=_ClaimingAgent(
                final_sentinel,
                termination=termination_sentinel,
            ),
            evaluator=evaluator,
        )

    persisted = b"\n".join(
        path.read_bytes() for path in tmp_path.rglob("*") if path.is_file()
    )
    assert b"ValueError" in persisted
    assert final_sentinel.encode("utf-8") not in persisted
    assert termination_sentinel.encode("utf-8") not in persisted


def test_symlink_evaluator_error_persists_only_exception_type(
    tmp_path: Path,
) -> None:
    """验证 symlink/escape 错误经 AttemptRunner 后仅持久化脱敏异常类型。

    输入参数：
        tmp_path：pytest 提供的 RunStore 与根外目标父目录。
    输出返回值：
        无；评价须抛出固定 ``OperationEvaluationError``，RunStore
        不得含敏感链接名、宿主路径、原因正文或 Agent final text。
    """

    prepared = _prepared_word009_task()
    outside = tmp_path / "PRIVATE-OUTSIDE-TARGET"
    outside.write_bytes(b"private")
    temporary_directory = tempfile.TemporaryDirectory(
        prefix="paraguibench-word009-symlink-"
    )
    root = Path(temporary_directory.name)
    sensitive_name = "PRIVATE-LINK.docx"
    (root / sensitive_name).symlink_to(outside)
    snapshot = OperationArtifactSnapshot(
        task_id=_WORD009_TASK_ID,
        protocol_id=OPERATION_PROTOCOL_ID,
        file_count=1,
        temporary_directory=temporary_directory,
    )
    store = RunStore(tmp_path / "runstore")
    attempt = _start_attempt(store, prepared)
    evaluator = build_task_evaluator(
        prepared.trusted_task,
        evaluation_protocol=OPERATION_PROTOCOL_ID,
    )
    final_sentinel = "PRIVATE SYMLINK FINAL CLAIM"

    with pytest.raises(OperationEvaluationError) as captured:
        AttemptRunner(store).run(
            attempt=attempt,
            prepared_task=prepared,
            environment=_SnapshotEnvironment(snapshot),
            agent=_ClaimingAgent(final_sentinel),
            evaluator=evaluator,
        )

    assert captured.value.code == "ARTIFACT_SYMLINK_REJECTED"
    persisted = b"\n".join(
        path.read_bytes()
        for path in (tmp_path / "runstore").rglob("*")
        if path.is_file()
    )
    assert b"OperationEvaluationError" in persisted
    for forbidden in (
        sensitive_name.encode("utf-8"),
        str(outside).encode("utf-8"),
        b"ARTIFACT_SYMLINK_REJECTED",
        final_sentinel.encode("utf-8"),
    ):
        assert forbidden not in persisted

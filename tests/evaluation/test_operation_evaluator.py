"""Operation eval-rules 原生协议的闭包与安全行为测试。"""

from __future__ import annotations

import hashlib
import json
from copy import deepcopy
from collections import Counter
from dataclasses import asdict
from pathlib import Path
import tomllib
import zipfile

import pytest

import paraguibench.evaluation.operation.evaluator as operation_evaluator_module
from paraguibench.evaluation.operation import (
    OPERATION_CHECK_CONTRACTS,
    OPERATION_PROTOCOL_ID,
    OPERATION_TASK_RULES,
    OperationEvaluationError,
    WordAbbreviationBaseline,
    WordTextBaseline,
    WordTextInputFile,
    capture_word_text_baseline,
    evaluate_operation_artifacts,
    operation_word_text_input_contract,
)


_REPO_ROOT = Path(__file__).resolve().parents[2]

_WORD004_EXPECTED_BY_FILE = {
    "center.docx": ("sesssion",),
    "episode.docx": ("childhoood",),
    "experience.docx": ("visiters", "severall"),
    "hall.docx": ("prasenting",),
    "travel.docx": ("intrenational", "conmference"),
}


def _canonical_sha256(value: object) -> str:
    """计算与生产规则目录一致的 canonical JSON 摘要。

    输入参数：
        value：要规范化并摘要的 JSON 兼容对象。
    输出返回值：
        小写十六进制 SHA-256 字符串。
    """

    payload = json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _write_minimal_ooxml(path: Path) -> None:
    """写入足以通过容器安全门的最小合成 OOXML fixture。

    输入参数：
        path：扩展名为 docx、xlsx 或 pptx 的测试输出路径。
    输出返回值：
        无；写入 Content Types 与对应根 document member。
    """

    required = {
        ".docx": "word/document.xml",
        ".xlsx": "xl/workbook.xml",
        ".pptx": "ppt/presentation.xml",
    }[path.suffix]
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", b"<Types/>")
        archive.writestr(required, b"<root/>")


def _write_real_xlsx(path: Path) -> None:
    """写入供 canonical XLSX checks 读取的真实合成工作簿。

    输入参数：
        path：目标 xlsx 路径。
    输出返回值：
        无；创建 54 行、4 列数值与文本混合的有效工作簿。
    """

    openpyxl = pytest.importorskip("openpyxl")
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    for row in range(1, 55):
        worksheet.append([row, row * 10, f"value-{row}", row * 100])
    workbook.save(path)
    workbook.close()


def _write_real_docx(path: Path) -> None:
    """写入供 canonical DOCX checks 读取的真实合成文档。

    输入参数：
        path：目标 docx 路径。
    输出返回值：
        无；创建标题、足量正文与一个小表格的有效文档。
    """

    docx = pytest.importorskip("docx")
    document = docx.Document()
    document.add_heading("1. Project Background", level=1)
    document.add_paragraph("word " * 40)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "city"
    table.cell(0, 1).text = "value"
    table.cell(1, 0).text = "sample"
    table.cell(1, 1).text = "1"
    document.save(path)


def _write_word004_docx(
    path: Path,
    misspellings: tuple[str, ...],
    highlighted_words: frozenset[str],
    *,
    highlight_all: bool = False,
) -> None:
    """写入可精确控制拼错词高亮状态的合成 Word artifact。

    输入参数：
        path：与 canonical 逐文件规则同名的输出 DOCX 路径。
        misspellings：写入该文档正文的拼错词序列。
        highlighted_words：其中需要设置黄色高亮的词集合。
        highlight_all：是否把普通上下文和连接词也全部设为黄色高亮。
    输出返回值：
        无；文档包含未高亮普通文本及独立的拼错词 run。
    """

    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_COLOR_INDEX

    document = docx.Document()
    paragraph = document.add_paragraph()
    prefix = paragraph.add_run("ordinary context before ")
    if highlight_all:
        prefix.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for index, word in enumerate(misspellings):
        if index:
            connector = paragraph.add_run(" and ")
            if highlight_all:
                connector.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run = paragraph.add_run(word)
        if highlight_all or word in highlighted_words:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    suffix = paragraph.add_run(" ordinary context after")
    if highlight_all:
        suffix.font.highlight_color = WD_COLOR_INDEX.YELLOW
    document.save(path)


def _write_real_pptx(path: Path) -> None:
    """写入供 canonical PPTX checks 读取的真实合成演示文稿。

    输入参数：
        path：目标 pptx 路径。
    输出返回值：
        无；创建一张空白幻灯片的有效 OOXML 演示文稿。
    """

    pptx = pytest.importorskip("pptx")
    presentation = pptx.Presentation()
    presentation.slides.add_slide(presentation.slide_layouts[6])
    presentation.save(path)


def _write_artifact_by_suffix(path: Path) -> None:
    """按 canonical 文件扩展名写入安全、可解析的 smoke artifact。

    输入参数：
        path：要创建的 PDF、DOCX、XLSX、PPTX 或 HTML 路径。
    输出返回值：
        无；父目录自动创建，未知扩展名写入非空普通文件。
    """

    path.parent.mkdir(parents=True, exist_ok=True)
    if path.suffix == ".pdf":
        path.write_bytes(b"%PDF-1.7\n%%EOF\n")
    elif path.suffix == ".docx":
        _write_real_docx(path)
    elif path.suffix == ".xlsx":
        _write_real_xlsx(path)
    elif path.suffix == ".pptx":
        _write_real_pptx(path)
    elif path.suffix == ".html":
        path.write_text(
            "<html><table><tr><td>1</td></tr></table></html>", encoding="utf-8"
        )
    else:
        path.write_bytes(b"fixture")


def _prepare_task_smoke_artifacts(root: Path, task: dict[str, object]) -> None:
    """根据固定规则参数生成足以到达检查主体的最小 artifact 集。

    输入参数：
        root：单个 canonical 任务的隔离 artifact 根目录。
        task：仓库读取的 canonical task object。
    输出返回值：
        无；只生成合成 Office/PDF/HTML，不复制任务 gold 或真实数据。
    """

    task_rule = OPERATION_TASK_RULES[task["task_id"]]
    if task["task_id"] == "Operation-FileOperate-BatchOperationWord-012":
        from tests.evaluation.test_operation_word012_abbreviation_semantics import (
            _EXPECTED_TEXT,
            _write_documents,
        )

        _write_documents(root, _EXPECTED_TEXT)
        return
    rules = task["eval_rules"]
    checks = set(task_rule.check_ids)
    if "check_named_files_exist" in checks:
        for name in rules[0]["params"]["filenames"]:
            _write_artifact_by_suffix(root / name)
        return
    if "check_files_in_same_folder" in checks:
        for index, group in enumerate(rules[0]["params"]["file_groups"], start=1):
            for name in group:
                _write_artifact_by_suffix(root / f"topic-{index}" / name)
        return
    if "check_html_files_for_xlsx" in checks:
        _write_real_xlsx(root / "output.xlsx")
        _write_artifact_by_suffix(root / "output.html")
        return
    if "check_sorted_copies_preserve_rows" in checks:
        params = rules[0]["params"]
        _write_real_xlsx(root / params["source_filename"])
        return
    if "check_heading_palette_and_references" in checks:
        for name in rules[0]["params"]["expected_files"]:
            _write_real_docx(root / name)
        return
    if "check_table_contains_expected_values" in checks:
        for name in rules[0]["params"]["expected_values_by_file"]:
            _write_real_docx(root / name)
        return
    if "check_values_scaled_from_source" in checks:
        for name in rules[0]["params"]["source_values_by_file"]:
            _write_real_xlsx(root / name)
        return
    if "check_misspelled_words_highlighted" in checks:
        for rule in rules:
            expected_words = tuple(rule["params"]["expected_highlights"])
            _write_word004_docx(
                root / rule["file_pattern"],
                expected_words,
                frozenset(),
            )
        return
    if "check_combinationdocs003_source_table_insert" in checks:
        for name in (
            "McDonalds_Monthly_Data.xlsx",
            "store1.xlsx",
            "store2.xlsx",
        ):
            _write_real_xlsx(root / name)
        _write_real_pptx(root / "McDonalds_powerpoint_report.pptx")
        return
    kinds = set(task_rule.artifact_kinds)
    if "xlsx" in kinds:
        _write_real_xlsx(root / "output.xlsx")
    elif "docx" in kinds:
        _write_real_docx(root / "output.docx")
    elif "pptx" in kinds:
        _write_real_pptx(root / "output.pptx")
    else:
        raise AssertionError(f"smoke fixture 未覆盖 {task['task_id']}")


def _word_text_smoke_baseline(
    root: Path,
    task_id: str,
) -> WordTextBaseline | None:
    """为 Word-009/010 闭包 smoke 构造 formal-identity pre DTO。

    输入参数：
        root：当前 task post 根；task_id：canonical Operation ID。
    输出返回值：
        非 009/010 返回 ``None``；目标任务返回与正式
        manifest SHA 及 DOCX 路径闭集绑定的合成 typed baseline。
    """

    contract = operation_word_text_input_contract(task_id)
    if contract is None:
        return None
    source_root = root.parent / f"{root.name}-word-text-pre"
    source_root.mkdir()
    files: list[WordTextInputFile] = []
    for pinned_file in contract.files:
        if not pinned_file.path.casefold().endswith(".docx"):
            continue
        path = source_root / pinned_file.path
        path.parent.mkdir(parents=True, exist_ok=True)
        _write_real_docx(path)
        payload = path.read_bytes()
        files.append(
            WordTextInputFile(
                path=pinned_file.path,
                size=len(payload),
                sha256=hashlib.sha256(payload).hexdigest(),
                is_docx=True,
            )
        )
    return capture_word_text_baseline(
        task_id=task_id,
        protocol_id=OPERATION_PROTOCOL_ID,
        manifest_sha256=contract.manifest_sha256,
        source_root=source_root,
        files=tuple(files),
    )


def _word_abbreviation_smoke_baseline(
    root: Path,
    task_id: str,
) -> WordAbbreviationBaseline | None:
    """为 Word-012 闭包 smoke 构造 formal-identity pre DTO。

    输入参数：
        root：当前 task post 根；task_id：canonical Operation ID。
    输出返回值：
        非 Word-012 返回 ``None``；目标任务返回固定四路径、
        正式 manifest SHA 与逐处语境绑定的合成 typed DTO。
    """

    if task_id != "Operation-FileOperate-BatchOperationWord-012":
        return None
    from tests.evaluation.test_operation_word012_abbreviation_semantics import (
        _SOURCE_TEXT,
        _baseline,
        _write_documents,
    )

    source_root = root.parent / f"{root.name}-word-abbreviation-pre"
    source_root.mkdir()
    _write_documents(source_root, _SOURCE_TEXT)
    return _baseline(source_root)


def test_operation_catalog_exactly_pins_all_native_eval_rule_tasks() -> None:
    """验证 32 项原生 Operation 任务被目录完整、逐字节绑定。

    输入参数：
        无；读取仓库 canonical runtime-support 与 task JSON。
    输出返回值：
        无；目录数量、任务闭集、规则顺序和参数摘要任一漂移即失败。
    """

    manifest = json.loads(
        (_REPO_ROOT / "benchmark/manifests/runtime-support-v1.json").read_text(
            encoding="utf-8"
        )
    )
    expected_ids = {
        row["task_id"]
        for row in manifest["tasks"]
        if row["evaluation_protocol"] == OPERATION_PROTOCOL_ID
    }

    assert OPERATION_PROTOCOL_ID == "paraguibench.operation.eval-rules.v1"
    assert len(expected_ids) == 32
    assert set(OPERATION_TASK_RULES) == expected_ids
    assert sum(len(rule.check_ids) for rule in OPERATION_TASK_RULES.values()) == 41
    for task_id in sorted(expected_ids):
        task = json.loads(
            (_REPO_ROOT / "benchmark/tasks" / f"{task_id}.json").read_text(
                encoding="utf-8"
            )
        )
        catalog_rule = OPERATION_TASK_RULES[task_id]
        assert catalog_rule.task_id == task_id
        assert catalog_rule.rule_set_sha256 == _canonical_sha256(task["eval_rules"])
        assert catalog_rule.check_ids == tuple(
            rule["check"] for rule in task["eval_rules"]
        )


def test_check_closure_and_optional_dependencies_are_machine_verifiable() -> None:
    """验证 33-check 闭集可达且 Office 依赖只存在于 optional extra。

    输入参数：
        无；读取固定 task 目录、公开检查 contract 与 ``pyproject.toml``。
    输出返回值：
        无；检查遗漏、额外旧代码或基础依赖污染均导致失败。
    """

    expected_checks = {
        check_id
        for task_rule in OPERATION_TASK_RULES.values()
        for check_id in task_rule.check_ids
    }
    assert len(expected_checks) == 33
    assert set(OPERATION_CHECK_CONTRACTS) == expected_checks
    dependency_counts = Counter(
        dependency
        for contract in OPERATION_CHECK_CONTRACTS.values()
        for dependency in contract.dependencies
    )
    assert set(dependency_counts) == {
        "stdlib",
        "openpyxl",
        "python-docx",
        "python-pptx",
        "Pillow",
    }
    assert dependency_counts == {
        "stdlib": 2,
        "openpyxl": 13,
        "python-docx": 16,
        "python-pptx": 3,
        "Pillow": 1,
    }
    project = tomllib.loads((_REPO_ROOT / "pyproject.toml").read_text(encoding="utf-8"))
    assert project["project"]["dependencies"] == []
    assert set(project["project"]["optional-dependencies"]["operation"]) == {
        "openpyxl>=3.1.5,<4",
        "python-docx>=1.1.2,<2",
        "python-pptx>=1.0.2,<2",
        "Pillow>=11,<13",
    }


def test_named_pdf_task_passes_without_persisting_paths_or_gold(tmp_path: Path) -> None:
    """验证纯文件存在性任务端到端通过且结果严格脱敏。

    输入参数：
        tmp_path：pytest 提供的隔离 Agent artifact 根目录。
    输出返回值：
        无；五个有效 PDF 应满分，结构化结果不得含路径、文件名或 gold。
    """

    task_path = (
        _REPO_ROOT / "benchmark/tasks/Operation-FileOperate-CombinationDocs-005.json"
    )
    task = json.loads(task_path.read_text(encoding="utf-8"))
    for filename in task["eval_rules"][0]["params"]["filenames"]:
        (tmp_path / filename).write_bytes(b"%PDF-1.7\n%%EOF\n")

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.protocol_id == OPERATION_PROTOCOL_ID
    assert evaluation.passed is True
    assert evaluation.score == 1.0
    assert evaluation.evaluated_rule_count == 1
    assert evaluation.passed_rule_count == 1
    assert evaluation.artifact_count == 5
    serialized = json.dumps(asdict(evaluation), ensure_ascii=False, sort_keys=True)
    assert str(tmp_path) not in serialized
    assert "Business_Report.pdf" not in serialized
    assert "filenames" not in serialized


def test_word004_exact_expected_highlights_score_one(tmp_path: Path) -> None:
    """验证五份文档仅高亮各自拼错词时 Word-004 严格满分。

    输入参数：
        tmp_path：pytest 提供的隔离合成 Word artifact 根目录。
    输出返回值：
        无；七个词按五份文件正确归属并黄色高亮时任务得分必须为 1。
    """

    task = json.loads(
        (
            _REPO_ROOT
            / "benchmark/tasks/Operation-FileOperate-BatchOperationWord-004.json"
        ).read_text(encoding="utf-8")
    )
    for filename, misspellings in _WORD004_EXPECTED_BY_FILE.items():
        _write_word004_docx(
            tmp_path / filename,
            misspellings,
            frozenset(misspellings),
        )

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is True
    assert evaluation.score == 1.0
    assert evaluation.evaluated_rule_count == 5
    assert evaluation.passed_rule_count == 5
    assert evaluation.artifact_count == 5


def test_word004_noop_scores_zero(tmp_path: Path) -> None:
    """验证五份文档均未高亮时 Word-004 得分严格为零。

    输入参数：
        tmp_path：pytest 提供的隔离合成 Word artifact 根目录。
    输出返回值：
        无；文中虽含七个目标词但没有黄色高亮时不能获得部分分。
    """

    task = json.loads(
        (
            _REPO_ROOT
            / "benchmark/tasks/Operation-FileOperate-BatchOperationWord-004.json"
        ).read_text(encoding="utf-8")
    )
    for filename, misspellings in _WORD004_EXPECTED_BY_FILE.items():
        _write_word004_docx(tmp_path / filename, misspellings, frozenset())

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.passed_rule_count == 0
    assert evaluation.failed_rule_count == 5


def test_word004_all_yellow_scores_zero(tmp_path: Path) -> None:
    """验证全文全部黄色高亮不能冒充 Word-004 正确产物。

    输入参数：
        tmp_path：pytest 提供的隔离合成 Word artifact 根目录。
    输出返回值：
        无；目标词召回虽完整，只要普通词也被黄色高亮，五条规则均须失败。
    """

    task = json.loads(
        (
            _REPO_ROOT
            / "benchmark/tasks/Operation-FileOperate-BatchOperationWord-004.json"
        ).read_text(encoding="utf-8")
    )
    for filename, misspellings in _WORD004_EXPECTED_BY_FILE.items():
        _write_word004_docx(
            tmp_path / filename,
            misspellings,
            frozenset(misspellings),
            highlight_all=True,
        )

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.passed_rule_count == 0
    assert evaluation.failed_rule_count == 5


def test_word004_highlights_in_the_wrong_documents_score_zero(
    tmp_path: Path,
) -> None:
    """验证把正确词轮换到错误文件中高亮时任务严格得零分。

    输入参数：
        tmp_path：pytest 提供的隔离合成 Word artifact 根目录。
    输出返回值：
        无；五份文件各自只高亮另一文件的目标词时不能跨文件抵消。
    """

    task = json.loads(
        (
            _REPO_ROOT
            / "benchmark/tasks/Operation-FileOperate-BatchOperationWord-004.json"
        ).read_text(encoding="utf-8")
    )
    filenames = tuple(_WORD004_EXPECTED_BY_FILE)
    for index, filename in enumerate(filenames):
        wrong_filename = filenames[(index + 1) % len(filenames)]
        wrong_words = _WORD004_EXPECTED_BY_FILE[wrong_filename]
        _write_word004_docx(
            tmp_path / filename,
            wrong_words,
            frozenset(wrong_words),
        )

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.passed_rule_count == 0
    assert evaluation.failed_rule_count == 5


def test_high_compression_ratio_ooxml_is_rejected_before_any_check(
    tmp_path: Path,
) -> None:
    """验证高压缩比 OOXML 在第三方解析器之前被统一拒绝。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；压缩炸弹必须形成固定 evaluator 错误且异常不回显路径。
    """

    task_path = (
        _REPO_ROOT / "benchmark/tasks/Operation-FileOperate-CombinationDocs-008.json"
    )
    task = json.loads(task_path.read_text(encoding="utf-8"))
    archive_path = tmp_path / task["eval_rules"][0]["params"]["filenames"][0]
    with zipfile.ZipFile(archive_path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", b"<Types/>")
        archive.writestr("word/document.xml", b"0" * (1024 * 1024))

    with pytest.raises(OperationEvaluationError) as raised:
        evaluate_operation_artifacts(tmp_path, task)

    assert raised.value.code == "ARCHIVE_COMPRESSION_RATIO_EXCEEDED"
    assert str(tmp_path) not in str(raised.value)


def test_ooxml_member_budget_is_aggregated_across_artifact_tree(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """验证多个各自合法的 OOXML 不能绕过整树解压预算。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根。
        monkeypatch：把整树 member 上限收紧为 3，避免构造
            大型压缩包才能验证跨文件聚合语义。
    输出返回值：
        无；两个各含 2 个 member 的容器必须在 Office
        parser 前以固定脱敏错误拒绝。
    """

    task_path = (
        _REPO_ROOT / "benchmark/tasks/Operation-FileOperate-CombinationDocs-008.json"
    )
    task = json.loads(task_path.read_text(encoding="utf-8"))
    expected_name = task["eval_rules"][0]["params"]["filenames"][0]
    _write_minimal_ooxml(tmp_path / expected_name)
    _write_minimal_ooxml(tmp_path / "additional.docx")
    monkeypatch.setattr(
        operation_evaluator_module,
        "_MAX_ARCHIVE_TREE_MEMBERS",
        3,
    )

    with pytest.raises(OperationEvaluationError) as raised:
        evaluate_operation_artifacts(tmp_path, task)

    assert raised.value.code == "ARCHIVE_TREE_MEMBER_LIMIT_EXCEEDED"
    assert str(tmp_path) not in str(raised.value)


def test_ooxml_uncompressed_budget_is_aggregated_across_artifact_tree(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """验证多个 OOXML 共享整棵 artifact 树的解压字节预算。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根。
        monkeypatch：将整树解压上限收紧为 20 字节，
            保留单文件合法、两文件聚合超限的边界。
    输出返回值：
        无；第二个容器在读取 XML payload 前必须以
        固定脱敏错误拒绝。
    """

    task_path = (
        _REPO_ROOT / "benchmark/tasks/Operation-FileOperate-CombinationDocs-008.json"
    )
    task = json.loads(task_path.read_text(encoding="utf-8"))
    expected_name = task["eval_rules"][0]["params"]["filenames"][0]
    _write_minimal_ooxml(tmp_path / expected_name)
    _write_minimal_ooxml(tmp_path / "additional.docx")
    monkeypatch.setattr(
        operation_evaluator_module,
        "_MAX_ARCHIVE_TREE_UNCOMPRESSED_BYTES",
        20,
    )

    with pytest.raises(OperationEvaluationError) as raised:
        evaluate_operation_artifacts(tmp_path, task)

    assert raised.value.code == "ARCHIVE_TREE_UNCOMPRESSED_LIMIT_EXCEEDED"
    assert str(tmp_path) not in str(raised.value)


def test_rule_parameter_tampering_is_rejected_before_artifact_access(
    tmp_path: Path,
) -> None:
    """验证完整规则摘要阻止参数、gold 或检查名的运行时替换。

    输入参数：
        tmp_path：保持为空的隔离目录，用于证明先校验规则身份。
    输出返回值：
        无；修改一个 canonical filename 后必须返回固定身份错误。
    """

    task_path = (
        _REPO_ROOT / "benchmark/tasks/Operation-FileOperate-CombinationDocs-005.json"
    )
    task = deepcopy(json.loads(task_path.read_text(encoding="utf-8")))
    task["eval_rules"][0]["params"]["filenames"].append("injected.pdf")

    with pytest.raises(OperationEvaluationError) as raised:
        evaluate_operation_artifacts(tmp_path, task)

    assert raised.value.code == "RULE_SET_IDENTITY_MISMATCH"
    assert "injected.pdf" not in str(raised.value)


def test_artifact_tree_symlink_is_rejected_without_following_it(tmp_path: Path) -> None:
    """验证 artifact 树中的文件 symlink 不会被存在性规则读取。

    输入参数：
        tmp_path：包含一个指向同目录普通文件的合成 symlink。
    输出返回值：
        无；即使目标合法，评价器也必须以固定错误拒绝整棵树。
    """

    task_path = (
        _REPO_ROOT / "benchmark/tasks/Operation-FileOperate-CombinationDocs-005.json"
    )
    task = json.loads(task_path.read_text(encoding="utf-8"))
    target = tmp_path / "target.pdf"
    target.write_bytes(b"%PDF-1.7\n%%EOF\n")
    (tmp_path / task["eval_rules"][0]["params"]["filenames"][0]).symlink_to(target)

    with pytest.raises(OperationEvaluationError) as raised:
        evaluate_operation_artifacts(tmp_path, task)

    assert raised.value.code == "ARTIFACT_SYMLINK_REJECTED"
    assert str(target) not in str(raised.value)


@pytest.mark.filterwarnings("ignore:Duplicate name:UserWarning")
@pytest.mark.parametrize(
    ("member_name", "payload", "expected_code"),
    [
        ("../word/document.xml", b"<root/>", "ARCHIVE_PATH_REJECTED"),
        (
            "[Content_Types].xml",
            b"<Types/>",
            "ARCHIVE_DUPLICATE_MEMBER_REJECTED",
        ),
        ("word/vbaProject.bin", b"macro", "ARCHIVE_MACRO_REJECTED"),
        (
            "word/document.xml",
            b"<!DOCTYPE root [<!ENTITY x 'x'>]><root>&x;</root>",
            "ARCHIVE_ACTIVE_XML_REJECTED",
        ),
        (
            "word/document.xml",
            "<!DOCTYPE root [<!ENTITY x 'x'>]><root>&x;</root>".encode("utf-16"),
            "ARCHIVE_ACTIVE_XML_REJECTED",
        ),
    ],
)
def test_ooxml_escape_macro_and_active_xml_are_rejected(
    tmp_path: Path,
    member_name: str,
    payload: bytes,
    expected_code: str,
) -> None:
    """验证 OOXML 路径逃逸、宏与 DTD/entity 都在解析前失败。

    输入参数：
        tmp_path：pytest 隔离目录；member_name/payload 为恶意 ZIP member；
        expected_code：对应的固定安全错误码。
    输出返回值：
        无；三类输入都不能到达 Office parser，错误不得回显 member 名。
    """

    task_path = (
        _REPO_ROOT / "benchmark/tasks/Operation-FileOperate-CombinationDocs-008.json"
    )
    task = json.loads(task_path.read_text(encoding="utf-8"))
    path = tmp_path / task["eval_rules"][0]["params"]["filenames"][0]
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", b"<Types/>")
        archive.writestr(member_name, payload)

    with pytest.raises(OperationEvaluationError) as raised:
        evaluate_operation_artifacts(tmp_path, task)

    assert raised.value.code == expected_code
    assert member_name not in str(raised.value)


def test_same_folder_task_passes_as_a_pure_directory_check(tmp_path: Path) -> None:
    """验证目录级规则由固定分派器执行且不需要 Office 解析依赖。

    输入参数：
        tmp_path：pytest 提供的隔离 Agent artifact 根目录。
    输出返回值：
        无；两个主题组各自位于独立子目录时任务应满分。
    """

    task_path = (
        _REPO_ROOT / "benchmark/tasks/Operation-FileOperate-CombinationDocs-007.json"
    )
    task = json.loads(task_path.read_text(encoding="utf-8"))
    groups = task["eval_rules"][0]["params"]["file_groups"]
    for index, group in enumerate(groups, start=1):
        folder = tmp_path / f"topic-{index}"
        folder.mkdir()
        for filename in group:
            _write_minimal_ooxml(folder / filename)

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is True
    assert evaluation.score == 1.0
    assert evaluation.artifact_count == sum(map(len, groups))
    assert evaluation.rule_results[0].check_id == "check_files_in_same_folder"


def test_all_32_canonical_tasks_reach_their_real_check_closure(
    tmp_path: Path,
) -> None:
    """验证 32 项 canonical 任务均可到达真实检查闭包并返回分数。

    输入参数：
        tmp_path：pytest 提供的 32 个隔离 artifact 根目录父目录。
    输出返回值：
        无；每项都经过真实 Office parser 与 canonical 参数，任何 evaluator
        error 都表示规则、依赖或最小 fixture 发生回归。
    """

    pytest.importorskip("openpyxl")
    pytest.importorskip("docx")
    pytest.importorskip("pptx")
    errors: dict[str, str] = {}
    for task_id in sorted(OPERATION_TASK_RULES):
        task = json.loads(
            (_REPO_ROOT / "benchmark/tasks" / f"{task_id}.json").read_text(
                encoding="utf-8"
            )
        )
        task_root = tmp_path / task_id
        task_root.mkdir()
        _prepare_task_smoke_artifacts(task_root, task)
        input_text_baseline = _word_text_smoke_baseline(task_root, task_id)
        input_abbreviation_baseline = _word_abbreviation_smoke_baseline(
            task_root,
            task_id,
        )
        try:
            evaluation = evaluate_operation_artifacts(
                task_root,
                task,
                input_text_baseline=input_text_baseline,
                input_abbreviation_baseline=input_abbreviation_baseline,
            )
        except OperationEvaluationError as exc:
            errors[task_id] = exc.code
        else:
            assert evaluation.evaluated_rule_count == len(task["eval_rules"])
            assert 0.0 <= evaluation.score <= 1.0

    assert errors == {}

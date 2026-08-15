"""Word-009/010 Operation 评价器的闭集安全回归。"""

from __future__ import annotations

import base64
import hashlib
import json
import os
from pathlib import Path
import tempfile
from typing import Any

import pytest

from paraguibench.evaluation.operation import (
    OPERATION_PROTOCOL_ID,
    OperationEvaluationError,
    WordTextFidelityError,
    WordTextInputFile,
    capture_word_text_baseline,
    evaluate_operation_artifacts as _evaluate_operation_artifacts,
)
import paraguibench.evaluation.operation.evaluator as operation_evaluator_module


pytest.importorskip("docx")

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from paraguibench.evaluation.operation.checks.docx import (
    check_image_name_matches_doc,
)


_REPO_ROOT = Path(__file__).resolve().parents[2]
_WORD009_TASK_ID = "Operation-FileOperate-BatchOperationWord-009"
_WORD010_TASK_ID = "Operation-FileOperate-BatchOperationWord-010"
_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8A"
    "AQUBAScY42YAAAAASUVORK5CYII="
)


def _load_task(task_id: str) -> dict[str, Any]:
    """读取仓库中的正式 Operation 任务。

    输入参数：
        task_id：对应 ``benchmark/tasks/<task_id>.json`` 的任务 ID。
    输出返回值：
        包含完整 evaluator 规则的可变任务映射。
    """

    path = _REPO_ROOT / "benchmark/tasks" / f"{task_id}.json"
    return json.loads(path.read_text(encoding="utf-8"))


def _word_text_baseline_for_test(
    result_root: Path,
    task: dict[str, Any],
):
    """为旧检查原语回归构造当前 post 的合成 pre 快照。

    输入参数：
        result_root：本测试的 artifact 根；task：Word-009/010 canonical 任务。
    输出返回值：
        具有正式 manifest 身份和精确 DOCX 路径闭集的 typed baseline。
        该辅助仅用于保留旧行距/图片原语测试，生产 baseline
        仍只由 environment 在 guest 访问前构造。
    """

    contract = operation_evaluator_module._PINNED_ARTIFACT_CONTRACTS[task["task_id"]]
    document_paths = tuple(
        file.path for file in contract.files if file.path.endswith(".docx")
    )

    def _capture(source_root: Path):
        """从合成 pre 根生成并返回不持有文件句柄的 DTO。"""

        files: list[WordTextInputFile] = []
        for relative_path in document_paths:
            payload = (source_root / relative_path).read_bytes()
            files.append(
                WordTextInputFile(
                    path=relative_path,
                    size=len(payload),
                    sha256=hashlib.sha256(payload).hexdigest(),
                    is_docx=True,
                )
            )
        return capture_word_text_baseline(
            task_id=task["task_id"],
            protocol_id=OPERATION_PROTOCOL_ID,
            manifest_sha256=contract.manifest_sha256,
            source_root=source_root,
            files=tuple(files),
        )

    if all((result_root / path).is_file() for path in document_paths):
        try:
            return _capture(result_root)
        except WordTextFidelityError:
            pass
    with tempfile.TemporaryDirectory(
        prefix="word-text-pre-",
        dir=result_root.parent,
    ) as temporary:
        source_root = Path(temporary)
        for path in document_paths:
            _write_double_spaced_docx(source_root / path, text="BASELINE")
        return _capture(source_root)


def evaluate_operation_artifacts(
    result_root: Path,
    task: dict[str, Any],
):
    """使旧原语测试通过正式 typed baseline 调用生产 evaluator。

    输入参数：
        result_root/task：与生产 evaluator 一致的 artifact 根和 task。
    输出返回值：
        生产 ``OperationEvaluation``；不使用 Agent final text。
    """

    baseline = _word_text_baseline_for_test(result_root, task)
    return _evaluate_operation_artifacts(
        result_root,
        task,
        input_text_baseline=baseline,
    )


def _write_double_spaced_docx(path: Path, text: str = "visible body") -> None:
    """写入一份具有可见正文和直接双倍行距的 DOCX。

    输入参数：
        path：目标 DOCX 路径；text：写入段落的可见文本。
    输出返回值：
        无；父目录存在时写入有效 OOXML 文档。
    """

    document = Document()
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.line_spacing = 2.0
    document.save(path)


def _write_nine_of_ten_double_spaced_docx(path: Path) -> None:
    """写入 9 个双倍段落与 1 个单倍段落的 DOCX。

    输入参数：
        path：目标 DOCX 路径。
    输出返回值：
        无；生成比例精确为 90% 但不满足 ``Change all`` 的对抗文档。
    """

    document = Document()
    for index in range(10):
        paragraph = document.add_paragraph(f"visible body {index}")
        paragraph.paragraph_format.line_spacing = 2.0 if index < 9 else 1.0
    document.save(path)


def _write_visible_inserted_single_spacing_docx(path: Path) -> None:
    """写入含修订插入可见单倍段落的 DOCX。

    输入参数：
        path：目标 DOCX 路径。
    输出返回值：
        无；首段为普通双倍段落，次段的可见文本位于
        ``w:ins/w:r/w:t`` 中且行距为 1.0。
    """

    document = Document()
    first = document.add_paragraph("ordinary visible double")
    first.paragraph_format.line_spacing = 2.0
    inserted_paragraph = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    properties.append(spacing)
    inserted_paragraph.append(properties)
    insertion = OxmlElement("w:ins")
    insertion.set(qn("w:id"), "1")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "VISIBLE INSERTED SINGLE"
    run.append(text)
    insertion.append(run)
    inserted_paragraph.append(insertion)
    document.element.body.insert(-1, inserted_paragraph)
    document.save(path)


def _write_hidden_single_spacing_after_visible_double_docx(path: Path) -> None:
    """写入可见双倍段落与隐藏单倍段落的 DOCX。

    输入参数：
        path：目标 DOCX 路径。
    输出返回值：
        无；用于验证 ``w:vanish`` 文本不被纳入可见正文分母。
    """

    document = Document()
    visible = document.add_paragraph("ordinary visible double")
    visible.paragraph_format.line_spacing = 2.0
    hidden = document.add_paragraph()
    hidden.paragraph_format.line_spacing = 1.0
    hidden_run = hidden.add_run("hidden single spacing")
    hidden_run.font.hidden = True
    document.save(path)


def _write_visible_single_with_unknown_style_after_double_docx(
    path: Path,
    *,
    reference_kind: str,
) -> None:
    """写入未知样式不得缩减可见分母的 DOCX。

    输入参数：
        path：目标 DOCX 路径；reference_kind：``pStyle`` 或
            ``rStyle``，决定未知引用的位置。
    输出返回值：
        无；首段可见且双倍，次段可见且单倍，并附加一个
        不存在的样式引用。
    """

    if reference_kind not in {"pStyle", "rStyle"}:
        raise ValueError("reference_kind 必须是 pStyle 或 rStyle")
    document = Document()
    visible = document.add_paragraph("ordinary visible double")
    visible.paragraph_format.line_spacing = 2.0
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("visible single with unknown style")
    if reference_kind == "pStyle":
        style_reference = OxmlElement("w:pStyle")
        style_reference.set(qn("w:val"), "DOES_NOT_EXIST")
        paragraph._p.get_or_add_pPr().insert(0, style_reference)
    else:
        style_reference = OxmlElement("w:rStyle")
        style_reference.set(qn("w:val"), "DOES_NOT_EXIST")
        run._r.get_or_add_rPr().insert(0, style_reference)
    document.save(path)


def _write_direct_unhidden_single_over_hidden_style_after_double_docx(
    path: Path,
    *,
    reference_kind: str,
) -> None:
    """写入直接取消隐藏应覆盖样式隐藏的 DOCX。

    输入参数：
        path：目标 DOCX 路径；reference_kind：``pStyle`` 或
            ``rStyle``，决定隐藏属性的样式类型。
    输出返回值：
        无；首段双倍，次段单倍；次段样式声明隐藏，但
        run 直接声明 ``w:vanish w:val=0``，因此 Word 将其显示。
    """

    if reference_kind not in {"pStyle", "rStyle"}:
        raise ValueError("reference_kind 必须是 pStyle 或 rStyle")
    document = Document()
    visible = document.add_paragraph("ordinary visible double")
    visible.paragraph_format.line_spacing = 2.0
    if reference_kind == "pStyle":
        hidden_style = document.styles.add_style(
            "Hidden Paragraph Style",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        style_properties = OxmlElement("w:rPr")
        vanish = OxmlElement("w:vanish")
        style_properties.append(vanish)
        hidden_style._element.append(style_properties)
        paragraph = document.add_paragraph(style=hidden_style)
        run = paragraph.add_run("directly unhidden visible single")
    else:
        hidden_style = document.styles.add_style(
            "Hidden Character Style",
            WD_STYLE_TYPE.CHARACTER,
        )
        hidden_style.font.hidden = True
        paragraph = document.add_paragraph()
        run = paragraph.add_run("directly unhidden visible single")
        run.style = hidden_style
    paragraph.paragraph_format.line_spacing = 1.0
    run.font.hidden = False
    document.save(path)


def _manifest_entries(task: dict[str, Any]) -> tuple[dict[str, Any], ...]:
    """读取任务声明的正式 input manifest 条目。

    输入参数：
        task：包含仓库相对 ``asset_manifest`` 的 canonical 任务。
    输出返回值：
        按 manifest 顺序保留的不可变文件条目元组。
    """

    manifest_path = _REPO_ROOT / task["asset_manifest"]
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    return tuple(manifest["files"])


def _write_docdefaults_spaced_docx(path: Path, *, line_rule: str) -> None:
    """写入仅由 ``docDefaults`` 声明 480 行距值的 DOCX。

    输入参数：
        path：目标 DOCX 路径；line_rule：OOXML ``w:lineRule`` 字面值。
    输出返回值：
        无；段落本身和样式都不设行距，便于验证默认层规则。
    """

    document = Document()
    document.add_paragraph("visible body")
    styles = document.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    ppr_default = doc_defaults.find(qn("w:pPrDefault"))
    if ppr_default is None:
        ppr_default = OxmlElement("w:pPrDefault")
        doc_defaults.append(ppr_default)
    ppr = ppr_default.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        ppr_default.append(ppr)
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:line"), "480")
    spacing.set(qn("w:lineRule"), line_rule)
    document.save(path)


def _write_unknown_style_with_double_default_docx(path: Path) -> None:
    """写入引用未定义段落样式、但 docDefaults 为双倍行距的 DOCX。

    输入参数：
        path：目标 DOCX 路径。
    输出返回值：
        无；生成一份 python-docx 可打开，但 ``w:pStyle`` 无对应
        ``w:style`` 定义的对抗文档。
    """

    document = Document()
    paragraph = document.add_paragraph("visible body")
    ppr = paragraph._p.get_or_add_pPr()
    pstyle = ppr.find(qn("w:pStyle"))
    if pstyle is None:
        pstyle = OxmlElement("w:pStyle")
        ppr.insert(0, pstyle)
    pstyle.set(qn("w:val"), "UndefinedParagraphStyle")
    styles = document.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    ppr_default = doc_defaults.find(qn("w:pPrDefault"))
    if ppr_default is None:
        ppr_default = OxmlElement("w:pPrDefault")
        doc_defaults.append(ppr_default)
    default_ppr = ppr_default.find(qn("w:pPr"))
    if default_ppr is None:
        default_ppr = OxmlElement("w:pPr")
        ppr_default.append(default_ppr)
    spacing = default_ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        default_ppr.append(spacing)
    spacing.set(qn("w:line"), "480")
    spacing.set(qn("w:lineRule"), "auto")
    document.save(path)


def _write_overdeep_style_chain_with_double_default_docx(
    path: Path,
    *,
    depth: int = 65,
) -> None:
    """写入超过 evaluator 固定上限的 ``basedOn`` 样式链文档。

    输入参数：
        path：目标 DOCX 路径；depth：引用链样式数，默认 65。
    输出返回值：
        无；所有样式均不声明行距，只有 docDefaults 声明双倍值。
    """

    document = Document()
    previous = None
    for index in range(depth):
        style = document.styles.add_style(
            f"Deep Paragraph Style {index:02d}",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        if previous is not None:
            style.base_style = previous
        previous = style
    document.add_paragraph("visible body", style=previous)
    styles = document.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    ppr_default = doc_defaults.find(qn("w:pPrDefault"))
    if ppr_default is None:
        ppr_default = OxmlElement("w:pPrDefault")
        doc_defaults.append(ppr_default)
    default_ppr = ppr_default.find(qn("w:pPr"))
    if default_ppr is None:
        default_ppr = OxmlElement("w:pPr")
        ppr_default.append(default_ppr)
    spacing = default_ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        default_ppr.append(spacing)
    spacing.set(qn("w:line"), "480")
    spacing.set(qn("w:lineRule"), "auto")
    document.save(path)


def _write_run_style_inherited_double_spacing_docx(path: Path) -> None:
    """写入仅由 run ``rStyle`` 的 ``basedOn`` 链声明双倍行距的 DOCX。

    输入参数：
        path：目标 DOCX 路径。
    输出返回值：
        无；段落无直接行距，字符样式的基样式 ``w:pPr``
        含有双倍行距，用于验证 rStyle 传递优先级。
    """

    document = Document()
    base_style = document.styles.add_style(
        "Double Run Base",
        WD_STYLE_TYPE.CHARACTER,
    )
    ppr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "480")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    base_style._element.append(ppr)
    derived_style = document.styles.add_style(
        "Double Run Derived",
        WD_STYLE_TYPE.CHARACTER,
    )
    derived_style.base_style = base_style
    run = document.add_paragraph().add_run("visible body")
    run.style = derived_style
    document.save(path)


def _write_character_style_referenced_as_pstyle_docx(path: Path) -> None:
    """写入把字符样式非法引用为 ``pStyle`` 的 DOCX。

    输入参数：
        path：目标 DOCX 路径。
    输出返回值：
        无；字符样式内伪造双倍 ``w:pPr``，段落的
        ``w:pStyle`` 跨类型引用该样式。
    """

    document = Document()
    spoof_style = document.styles.add_style(
        "Character Style PStyle Spoof",
        WD_STYLE_TYPE.CHARACTER,
    )
    style_ppr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "480")
    spacing.set(qn("w:lineRule"), "auto")
    style_ppr.append(spacing)
    spoof_style._element.append(style_ppr)
    paragraph = document.add_paragraph("visible body")
    pstyle = OxmlElement("w:pStyle")
    pstyle.set(qn("w:val"), spoof_style.style_id)
    paragraph._p.get_or_add_pPr().insert(0, pstyle)
    document.save(path)


def _write_paragraph_style_referenced_as_rstyle_docx(path: Path) -> None:
    """写入把段落样式非法引用为 ``rStyle`` 的 DOCX。

    输入参数：
        path：目标 DOCX 路径。
    输出返回值：
        无；段落样式声明双倍行距，可见 run 的 ``w:rStyle``
        跨类型引用该样式。
    """

    document = Document()
    spoof_style = document.styles.add_style(
        "Paragraph Style RStyle Spoof",
        WD_STYLE_TYPE.PARAGRAPH,
    )
    spoof_style.paragraph_format.line_spacing = 2.0
    run = document.add_paragraph().add_run("visible body")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), spoof_style.style_id)
    run._r.get_or_add_rPr().insert(0, rstyle)
    document.save(path)


def _write_cyclic_style_with_double_default_docx(path: Path) -> None:
    """写入段落样式 ``A -> B -> A`` 循环且 docDefaults 为双倍的 DOCX。

    输入参数：
        path：目标 DOCX 路径。
    输出返回值：
        无；生成 python-docx 可打开，但引用样式图含循环的对抗文档。
    """

    document = Document()
    style_a = document.styles.add_style("Cycle A", WD_STYLE_TYPE.PARAGRAPH)
    style_b = document.styles.add_style("Cycle B", WD_STYLE_TYPE.PARAGRAPH)
    style_a.base_style = style_b
    based_on = OxmlElement("w:basedOn")
    based_on.set(qn("w:val"), style_a.style_id)
    style_b._element.insert(0, based_on)
    document.add_paragraph("visible body", style=style_a)
    styles = document.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    assert doc_defaults is not None
    ppr_default = doc_defaults.find(qn("w:pPrDefault"))
    assert ppr_default is not None
    default_ppr = ppr_default.find(qn("w:pPr"))
    assert default_ppr is not None
    spacing = default_ppr.find(qn("w:spacing"))
    assert spacing is not None
    spacing.set(qn("w:line"), "480")
    spacing.set(qn("w:lineRule"), "auto")
    document.save(path)


def _write_direct_double_over_exact_style_docx(path: Path) -> None:
    """写入直接双倍行距覆盖样式固定磅值的 DOCX。

    输入参数：
        path：目标 DOCX 路径。
    输出返回值：
        无；段落样式声明 24pt ``exact``，段落直接层声明 2.0
        ``auto``，用于验证 direct 优先级。
    """

    document = Document()
    style = document.styles.add_style("Exact Style", WD_STYLE_TYPE.PARAGRAPH)
    style.paragraph_format.line_spacing = Pt(24)
    paragraph = document.add_paragraph("visible body", style=style)
    paragraph.paragraph_format.line_spacing = 2.0
    document.save(path)


def _write_duplicate_direct_spacing_docx(path: Path) -> None:
    """写入同一 ``w:pPr`` 含两个冲突 spacing 的 DOCX。

    输入参数：
        path：目标 DOCX 路径。
    输出返回值：
        无；首个 spacing 为双倍，第二个为单倍，用于验证
        malformed OOXML 不得因 evaluator 只读首个节点而通过。
    """

    document = Document()
    paragraph = document.add_paragraph("visible body")
    paragraph.paragraph_format.line_spacing = 2.0
    duplicate = OxmlElement("w:spacing")
    duplicate.set(qn("w:line"), "240")
    duplicate.set(qn("w:lineRule"), "auto")
    paragraph._p.get_or_add_pPr().append(duplicate)
    document.save(path)


def _write_duplicate_style_reference_docx(
    path: Path,
    *,
    reference_kind: str,
) -> None:
    """写入同一容器含重复 ``pStyle`` 或 ``rStyle`` 的 DOCX。

    输入参数：
        path：目标 DOCX 路径；reference_kind：``pStyle`` 或
            ``rStyle``。
    输出返回值：
        无；首个样式有效且提供双倍行距，第二个引用未知
        样式，用于要求 evaluator 对重复身份 fail closed。
    """

    if reference_kind not in {"pStyle", "rStyle"}:
        raise ValueError("reference_kind 必须是 pStyle 或 rStyle")
    document = Document()
    if reference_kind == "pStyle":
        valid_style = document.styles.add_style(
            "Valid Double Paragraph",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        valid_style.paragraph_format.line_spacing = 2.0
        paragraph = document.add_paragraph("visible body", style=valid_style)
        duplicate = OxmlElement("w:pStyle")
        duplicate.set(qn("w:val"), "DOES_NOT_EXIST")
        paragraph._p.get_or_add_pPr().append(duplicate)
    else:
        valid_style = document.styles.add_style(
            "Valid Double Character",
            WD_STYLE_TYPE.CHARACTER,
        )
        style_properties = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "480")
        spacing.set(qn("w:lineRule"), "auto")
        style_properties.append(spacing)
        valid_style._element.append(style_properties)
        run = document.add_paragraph().add_run("visible body")
        run.style = valid_style
        duplicate = OxmlElement("w:rStyle")
        duplicate.set(qn("w:val"), "DOES_NOT_EXIST")
        run._r.get_or_add_rPr().append(duplicate)
    document.save(path)


def _spacing_properties(line_value: str) -> Any:
    """构造只含一个 ``w:spacing`` 的 ``w:pPr``。

    输入参数：
        line_value：``auto`` 行距的 twip 字面值，例如 240/480。
    输出返回值：
        未挂载的 ``w:pPr`` OOXML 元素。
    """

    properties = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), line_value)
    spacing.set(qn("w:lineRule"), "auto")
    properties.append(spacing)
    return properties


def _write_duplicate_structural_container_docx(
    path: Path,
    *,
    malformed_kind: str,
) -> None:
    """写入重复 OOXML 属性容器的 Word009 对抗文档。

    输入参数：
        path：目标 DOCX 路径；malformed_kind：``basedOn``、
            ``style_pPr``、``docdefaults_pPr`` 或 ``direct_pPr``。
    输出返回值：
        无；每种结构的首节点声明双倍行距，重复节点导向单倍，
        用于验证 evaluator 不使用 first-match 解析 malformed OOXML。
    """

    allowed = {"basedOn", "style_pPr", "docdefaults_pPr", "direct_pPr"}
    if malformed_kind not in allowed:
        raise ValueError("malformed_kind 不在固定集合中")
    document = Document()
    if malformed_kind == "basedOn":
        good_base = document.styles.add_style(
            "Good Double Base",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        good_base.paragraph_format.line_spacing = 2.0
        bad_base = document.styles.add_style(
            "Bad Single Base",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        bad_base.paragraph_format.line_spacing = 1.0
        derived = document.styles.add_style(
            "Duplicate BasedOn Derived",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        derived.base_style = good_base
        duplicate = OxmlElement("w:basedOn")
        duplicate.set(qn("w:val"), bad_base.style_id)
        derived._element.append(duplicate)
        document.add_paragraph("visible body", style=derived)
    elif malformed_kind == "style_pPr":
        style = document.styles.add_style(
            "Duplicate PPr Style",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        style.paragraph_format.line_spacing = 2.0
        style._element.append(_spacing_properties("240"))
        document.add_paragraph("visible body", style=style)
    elif malformed_kind == "docdefaults_pPr":
        document.add_paragraph("visible body")
        styles = document.styles.element
        doc_defaults = styles.find(qn("w:docDefaults"))
        assert doc_defaults is not None
        ppr_default = doc_defaults.find(qn("w:pPrDefault"))
        assert ppr_default is not None
        first_properties = ppr_default.find(qn("w:pPr"))
        assert first_properties is not None
        first_spacing = first_properties.find(qn("w:spacing"))
        if first_spacing is None:
            first_spacing = OxmlElement("w:spacing")
            first_properties.append(first_spacing)
        first_spacing.set(qn("w:line"), "480")
        first_spacing.set(qn("w:lineRule"), "auto")
        ppr_default.append(_spacing_properties("240"))
    else:
        paragraph = document.add_paragraph("visible body")
        paragraph.paragraph_format.line_spacing = 2.0
        paragraph._p.insert(1, _spacing_properties("240"))
    document.save(path)


def test_word009_missing_unfinished_document_cannot_shrink_denominator(
    tmp_path: Path,
) -> None:
    """验证删除未完成文档不能把 Word-009 剩余文档平均成满分。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；只交付四份 pinned 文档中的三份时必须固定失败。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names[:-1]:
        _write_double_spaced_docx(tmp_path / filename)

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.reason_codes == ("ARTIFACT_CONTRACT_MISMATCH",)
    assert evaluation.rule_results[0].evaluated_artifact_count == len(expected_names)


def test_word009_missing_baseline_precedes_missing_document_failure(
    tmp_path: Path,
) -> None:
    """验证缺失 pre baseline 不能被 post 缺文件的 FAIL 绕过。

    输入参数：
        tmp_path：不含任何交付文档的 artifact 根。
    输出返回值：
        无；生产 evaluator 必须先报 baseline ERROR，不得返回
        ``ARTIFACT_CONTRACT_MISMATCH`` 零分。
    """

    task = _load_task(_WORD009_TASK_ID)

    with pytest.raises(OperationEvaluationError) as captured:
        _evaluate_operation_artifacts(tmp_path, task)

    assert captured.value.code == "WORD_TEXT_BASELINE_REQUIRED"


def test_word009_wrong_baseline_identity_precedes_missing_document_failure(
    tmp_path: Path,
) -> None:
    """验证跨 task baseline 与 post 缺文件并存时仍是 ERROR。

    输入参数：
        tmp_path：空 artifact 根与合成 pre 快照临时目录的父根。
    输出返回值：
        无；010 baseline 不能作为 009 证据，且不受 post 闭集影响。
    """

    task = _load_task(_WORD009_TASK_ID)
    wrong_baseline = _word_text_baseline_for_test(
        tmp_path,
        _load_task(_WORD010_TASK_ID),
    )

    with pytest.raises(OperationEvaluationError) as captured:
        _evaluate_operation_artifacts(
            tmp_path,
            task,
            input_text_baseline=wrong_baseline,
        )

    assert captured.value.code == "WORD_TEXT_FIDELITY_INVALID"


def test_word009_exact_four_document_contract_passes_with_fixed_denominator(
    tmp_path: Path,
) -> None:
    """验证四份 pinned 文档齐全且均为双倍行距时满分通过。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；正向路径必须保留固定四文档分母与空 reason code。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_double_spaced_docx(tmp_path / filename)

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is True
    assert evaluation.score == 1.0
    assert evaluation.reason_codes == ()
    assert evaluation.rule_results[0].evaluated_artifact_count == 4


def test_word009_change_all_rejects_canonical_ninety_percent_threshold_bypass(
    tmp_path: Path,
) -> None:
    """验证 canonical ``threshold=0.9`` 不能放宽 ``Change all`` 的全称约束。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；每份文档即使 10 段中有 9 段正确，只要 1 段不是
        双倍行距，formal evaluator 就必须严格失败。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_nine_of_ten_double_spaced_docx(tmp_path / filename)

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is False
    assert evaluation.score == 0.9
    assert evaluation.reason_codes == ("RULE_MISMATCH",)


def test_word009_visible_inserted_revision_text_stays_in_fixed_denominator(
    tmp_path: Path,
) -> None:
    """验证 ``w:ins`` 中的可见正文不能绕过全称行距约束。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；每份文档一个普通双倍段落和一个修订插入单倍
        段落时，formal evaluator 必须以 0.5 分失败。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_visible_inserted_single_spacing_docx(tmp_path / filename)

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is False
    assert evaluation.score == 0.5
    assert evaluation.reason_codes == ("RULE_MISMATCH",)


def test_word009_hidden_text_does_not_create_a_false_line_spacing_failure(
    tmp_path: Path,
) -> None:
    """验证完全隐藏的段落不进入可见正文行距分母。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；每份文档的唯一可见段落为双倍行距时应满分，
        ``w:vanish`` 单倍段落不得造成假失败。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_hidden_single_spacing_after_visible_double_docx(tmp_path / filename)

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is True
    assert evaluation.score == 1.0
    assert evaluation.reason_codes == ()


@pytest.mark.parametrize("reference_kind", ["pStyle", "rStyle"])
def test_word009_unknown_style_on_visible_text_cannot_shrink_denominator(
    tmp_path: Path,
    reference_kind: str,
) -> None:
    """验证样式图异常不能被误当作隐藏文本后跳过。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录；
        reference_kind：未知样式位于 ``pStyle`` 或 ``rStyle``。
    输出返回值：
        无；每份文档一个双倍段落和一个可见单倍段落时，
        未知样式必须 fail closed，不得让次段从分母消失。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_visible_single_with_unknown_style_after_double_docx(
            tmp_path / filename,
            reference_kind=reference_kind,
        )

    with pytest.raises(OperationEvaluationError) as captured:
        evaluate_operation_artifacts(tmp_path, task)

    assert captured.value.code == "WORD_TEXT_FIDELITY_INVALID"


@pytest.mark.parametrize("reference_kind", ["pStyle", "rStyle"])
def test_word009_direct_unhide_precedes_inherited_hidden_style(
    tmp_path: Path,
    reference_kind: str,
) -> None:
    """验证 run 直接 ``vanish=false`` 优先于样式继承的隐藏值。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录；
        reference_kind：隐藏值来自段落或字符样式。
    输出返回值：
        无；直接取消隐藏的单倍段落在 Word 中可见，必须进入
        分母并使 formal evaluator 以 0.5 分失败。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_direct_unhidden_single_over_hidden_style_after_double_docx(
            tmp_path / filename,
            reference_kind=reference_kind,
        )

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is False
    assert evaluation.score == 0.5
    assert evaluation.reason_codes == ("RULE_MISMATCH",)


def test_word009_extra_document_fails_exact_output_closure(tmp_path: Path) -> None:
    """验证在四份 pinned 文档外添加第五份 DOCX 也必须固定失败。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；额外文件不得进入动态平均，应返回固定 contract 零分。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in (*expected_names, "Unexpected.docx"):
        _write_double_spaced_docx(tmp_path / filename)

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.reason_codes == ("ARTIFACT_CONTRACT_MISMATCH",)
    assert evaluation.rule_results[0].evaluated_artifact_count == 4


def test_word009_empty_docx_is_a_real_fixed_denominator_failure(
    tmp_path: Path,
) -> None:
    """验证空 DOCX 不能利用零段落分母得分。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；四份文档路径齐全但其中一份无可见正文时，
        应按固定四文档分母返回 0.75 且严格失败。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names[:-1]:
        _write_double_spaced_docx(tmp_path / filename)
    _write_double_spaced_docx(tmp_path / expected_names[-1], text="")

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is False
    assert evaluation.score == 0.75
    assert evaluation.reason_codes == ("RULE_MISMATCH",)
    assert evaluation.rule_results[0].evaluated_artifact_count == len(expected_names)


@pytest.mark.parametrize("line_rule", ["exact", "atLeast"])
def test_word009_exact_or_at_least_points_are_not_double_spacing(
    tmp_path: Path,
    line_rule: str,
) -> None:
    """验证固定/最小磅值不能按 ``line/240`` 伪装双倍行距。

    输入参数：
        tmp_path：pytest 隔离 artifact 根；line_rule：``exact`` 或
        ``atLeast`` 两种非倍数 OOXML 语义。
    输出返回值：
        无；四份文档全部使用非 ``auto`` 480 值时须为零分。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_docdefaults_spaced_docx(tmp_path / filename, line_rule=line_rule)

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.reason_codes == ("RULE_MISMATCH",)


def test_word009_unknown_referenced_style_fails_closed(tmp_path: Path) -> None:
    """验证未定义 ``pStyle`` 不能回落到双倍 docDefaults 后通过。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；四份文档均引用未定义样式时应严格零分。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_unknown_style_with_double_default_docx(tmp_path / filename)

    with pytest.raises(OperationEvaluationError) as captured:
        evaluate_operation_artifacts(tmp_path, task)

    assert captured.value.code == "WORD_TEXT_FIDELITY_INVALID"


def test_word009_overdeep_style_chain_fails_closed(tmp_path: Path) -> None:
    """验证 65 层 ``basedOn`` 链不能在达到上限后回落 docDefaults。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；四份文档都使用超深链时应严格零分。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_overdeep_style_chain_with_double_default_docx(tmp_path / filename)

    with pytest.raises(OperationEvaluationError) as captured:
        evaluate_operation_artifacts(tmp_path, task)

    assert captured.value.code == "WORD_TEXT_FIDELITY_INVALID"


def test_word009_style_chain_at_64_layer_limit_remains_valid(tmp_path: Path) -> None:
    """验证精确 64 层 ``basedOn`` 链可完成审计并回落 docDefaults。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；64 层是可审计上限，链本身无行距且 docDefaults 为
        2.0 时四份文档应满分。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_overdeep_style_chain_with_double_default_docx(
            tmp_path / filename,
            depth=64,
        )

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is True
    assert evaluation.score == 1.0
    assert evaluation.reason_codes == ()


def test_word009_run_style_based_on_chain_precedes_docdefaults(
    tmp_path: Path,
) -> None:
    """验证可见 run 的 ``rStyle`` 可沿 ``basedOn`` 链提供有效行距。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；四份文档都只从 rStyle 基链继承双倍行距时应满分。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_run_style_inherited_double_spacing_docx(tmp_path / filename)

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is True
    assert evaluation.score == 1.0
    assert evaluation.reason_codes == ()


def test_word009_pstyle_rejects_character_style_type_confusion(
    tmp_path: Path,
) -> None:
    """验证 ``pStyle`` 不能借字符样式中的伪造行距通过。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；四份文档都让段落跨类型引用字符样式时必须零分。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_character_style_referenced_as_pstyle_docx(tmp_path / filename)

    with pytest.raises(OperationEvaluationError) as captured:
        evaluate_operation_artifacts(tmp_path, task)

    assert captured.value.code == "WORD_TEXT_FIDELITY_INVALID"


def test_word009_rstyle_rejects_paragraph_style_type_confusion(
    tmp_path: Path,
) -> None:
    """验证 ``rStyle`` 不能借段落样式中的行距通过。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；四份文档都让 run 跨类型引用段落样式时必须零分。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_paragraph_style_referenced_as_rstyle_docx(tmp_path / filename)

    with pytest.raises(OperationEvaluationError) as captured:
        evaluate_operation_artifacts(tmp_path, task)

    assert captured.value.code == "WORD_TEXT_FIDELITY_INVALID"


def test_word009_cyclic_style_graph_fails_closed(tmp_path: Path) -> None:
    """验证 ``basedOn`` 循环不能跳出审计后回落双倍 docDefaults。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；四份文档都引用循环样式图时应严格零分。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_cyclic_style_with_double_default_docx(tmp_path / filename)

    with pytest.raises(OperationEvaluationError) as captured:
        evaluate_operation_artifacts(tmp_path, task)

    assert captured.value.code == "WORD_TEXT_FIDELITY_INVALID"


def test_word009_direct_auto_spacing_precedes_exact_style_value(
    tmp_path: Path,
) -> None:
    """验证段落直接 ``auto`` 双倍值优先于样式中的 ``exact`` 磅值。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；四份文档的有效值均由 direct 2.0 决定，应满分通过。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_direct_double_over_exact_style_docx(tmp_path / filename)

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is True
    assert evaluation.score == 1.0
    assert evaluation.reason_codes == ()


def test_word009_duplicate_direct_spacing_nodes_fail_closed(tmp_path: Path) -> None:
    """验证同一段落的冲突 ``w:spacing`` 节点不得取首值通过。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；四份文档均含双倍+单倍重复直接声明时必须零分。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_duplicate_direct_spacing_docx(tmp_path / filename)

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.reason_codes == ("RULE_MISMATCH",)


@pytest.mark.parametrize("reference_kind", ["pStyle", "rStyle"])
def test_word009_duplicate_style_references_fail_closed(
    tmp_path: Path,
    reference_kind: str,
) -> None:
    """验证重复 ``pStyle``/``rStyle`` 不得通过取首引用绕过。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根；reference_kind：
            被重复的段落或字符样式引用。
    输出返回值：
        无；四份 malformed 文档必须零分且返回固定规则失配码。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_duplicate_style_reference_docx(
            tmp_path / filename,
            reference_kind=reference_kind,
        )

    with pytest.raises(OperationEvaluationError) as captured:
        evaluate_operation_artifacts(tmp_path, task)

    assert captured.value.code == "WORD_TEXT_FIDELITY_INVALID"


@pytest.mark.parametrize(
    "malformed_kind",
    ["basedOn", "style_pPr", "docdefaults_pPr", "direct_pPr"],
)
def test_word009_duplicate_property_containers_fail_closed(
    tmp_path: Path,
    malformed_kind: str,
) -> None:
    """验证样式图与默认/直接属性的重复容器统一失败。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根；malformed_kind：
            重复 basedOn、style pPr、docDefaults pPr 或 paragraph pPr。
    输出返回值：
        无；任一 first-match 冲突结构都必须按固定四文档分母零分。
    """

    task = _load_task(_WORD009_TASK_ID)
    expected_names = tuple(entry["path"] for entry in _manifest_entries(task))
    for filename in expected_names:
        _write_duplicate_structural_container_docx(
            tmp_path / filename,
            malformed_kind=malformed_kind,
        )

    with pytest.raises(OperationEvaluationError) as captured:
        evaluate_operation_artifacts(tmp_path, task)

    assert captured.value.code == "WORD_TEXT_FIDELITY_INVALID"


def test_word010_tampered_source_images_fail_formal_manifest_contract(
    tmp_path: Path,
) -> None:
    """验证与 pinned 正式 input manifest 大小/摘要不同的源图严格失败。

    输入参数：
        tmp_path：pytest 提供的隔离 artifact 根目录。
    输出返回值：
        无；即使 5 DOCX + ``images/`` 5 JPEG 路径闭集完整，
        任一源图字节不匹配也须用固定 contract reason 返回零分。
    """

    task = _load_task(_WORD010_TASK_ID)
    entries = _manifest_entries(task)
    for entry in entries:
        path = tmp_path / entry["path"]
        path.parent.mkdir(parents=True, exist_ok=True)
        if path.suffix.lower() == ".docx":
            document = Document()
            document.add_paragraph("[image Slot]")
            document.save(path)
        else:
            path.write_bytes(b"tampered-source-image")

    evaluation = evaluate_operation_artifacts(tmp_path, task)

    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.reason_codes == ("ARTIFACT_CONTRACT_MISMATCH",)
    assert evaluation.rule_results[0].evaluated_artifact_count == 5


@pytest.mark.parametrize("task_id", [_WORD009_TASK_ID, _WORD010_TASK_ID])
def test_pinned_evaluator_contract_exactly_matches_formal_input_manifest(
    task_id: str,
) -> None:
    """验证 evaluator 使用的路径、size、SHA 与正式 input manifest 逐项一致。

    输入参数：
        task_id：Word-009 或 Word-010 的固定 canonical 任务 ID。
    输出返回值：
        无；manifest 路径、整文件摘要、文件顺序与三元身份、
        DOCX 固定分母任一漂移均失败。
    """

    task = _load_task(task_id)
    contract = operation_evaluator_module._PINNED_ARTIFACT_CONTRACTS[task_id]
    assert task["asset_manifest"] == contract.manifest_reference
    manifest_path = _REPO_ROOT / contract.manifest_reference
    payload = manifest_path.read_bytes()
    manifest = json.loads(payload)

    assert hashlib.sha256(payload).hexdigest() == contract.manifest_sha256
    assert tuple(
        (entry["path"], entry["size"], entry["sha256"]) for entry in manifest["files"]
    ) == tuple((file.path, file.size, file.sha256) for file in contract.files)
    assert sum(file.path.endswith(".docx") for file in contract.files) == (
        contract.expected_document_count
    )


def test_word010_matching_image_before_placeholder_fails_order_contract(
    tmp_path: Path,
) -> None:
    """验证正确图片在 ``image Slot`` 之前时不能因字节和宽度匹配通过。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；图片内容与 5cm 宽度都正确但顺序在占位符前时必须失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph().add_run().add_picture(
        os.fspath(source),
        width=Cm(5),
    )
    document.add_paragraph("[image Slot]")
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_bytes_and_width_cannot_cross_inline_containers(
    tmp_path: Path,
) -> None:
    """验证同一 ``w:drawing`` 内不同 inline 的正确字节与 5cm 宽度不得笛卡尔拼接。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；正确图片仅 3cm，另一错图为 5cm，即使被合并到同一
        ``w:drawing``，也必须失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    matching_source = images_dir / "Cats.png"
    matching_source.write_bytes(_ONE_PIXEL_PNG)
    distractor = tmp_path / "distractor.png"
    distractor.write_bytes(_ONE_PIXEL_PNG + b"different-image-bytes")
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(os.fspath(matching_source), width=Cm(3))
    paragraph.add_run().add_picture(os.fspath(distractor), width=Cm(5))
    drawings = paragraph._p.findall(f".//{qn('w:drawing')}")
    assert len(drawings) == 2
    for child in tuple(drawings[1]):
        drawings[0].append(child)
    drawings[1].getparent().remove(drawings[1])
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_non_finite_extent_width_fails_closed(tmp_path: Path) -> None:
    """验证 ``wp:extent/@cx=NaN`` 不能利用浮点比较语义伪装 5cm。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；图片字节、占位符顺序和容器都正确，但宽度为非有限数时
        必须严格失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(os.fspath(source), width=Cm(5))
    extent = picture_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent"
    )
    assert extent is not None
    extent.set("cx", "NaN")
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_external_blip_link_cannot_accompany_pinned_embed(
    tmp_path: Path,
) -> None:
    """验证正确嵌入图片不得同时携带外部 ``r:link`` 越过 pinned 身份。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；``a:blip`` 的 ``r:embed`` 虽指向正确字节，只要另有
        ``r:link`` 就必须 fail closed。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(os.fspath(source), width=Cm(5))
    blip = picture_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    )
    assert blip is not None
    blip.set(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link",
        "rIdExternalAttacker",
    )
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_hidden_placeholder_text_is_not_visible_evidence(
    tmp_path: Path,
) -> None:
    """验证 ``w:vanish`` 隐藏的 image Slot 不能充当可见占位符。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；即使图片字节和宽度正确，隐藏占位符也必须失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    placeholder = document.add_paragraph().add_run("[image Slot]")
    placeholder.font.hidden = True
    document.add_paragraph().add_run().add_picture(
        os.fspath(source),
        width=Cm(5),
    )
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_hidden_drawing_is_not_visible_evidence(tmp_path: Path) -> None:
    """验证 ``wp:docPr/@hidden=1`` 的 drawing 不能充当插图证据。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；图片字节、5cm 宽度与顺序都正确时，只要
        drawing 声明隐藏就必须失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(os.fspath(source), width=Cm(5))
    doc_pr = picture_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr"
    )
    assert doc_pr is not None
    doc_pr.set("hidden", "1")
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_placeholder_must_be_an_exact_visible_paragraph(
    tmp_path: Path,
) -> None:
    """验证前后缀中夹带 ``image Slot`` 不等于精确占位符。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；``prefiximage SlotSuffix`` 后的正确图片必须失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("prefiximage SlotSuffix")
    document.add_paragraph().add_run().add_picture(
        os.fspath(source),
        width=Cm(5),
    )
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_fully_transparent_picture_is_not_visible_evidence(
    tmp_path: Path,
) -> None:
    """验证完全透明的嵌入图片不能充当可见 drawing。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；正确字节与 5cm 宽度的 ``a:blip`` 加入
        ``a:alphaModFix amt=0`` 后必须失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(os.fspath(source), width=Cm(5))
    blip = picture_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    )
    assert blip is not None
    alpha = OxmlElement("a:alphaModFix")
    alpha.set("amt", "0")
    blip.append(alpha)
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_inner_picture_extent_must_match_outer_five_cm_extent(
    tmp_path: Path,
) -> None:
    """验证内部图片几何不能与外层 5cm extent 矛盾。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；``wp:extent`` 保持 5cm，但同一 picture 的
        ``pic:spPr/a:xfrm/a:ext/@cx`` 为零时必须失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(os.fspath(source), width=Cm(5))
    inner_extent = picture_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr/"
        "{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm/"
        "{http://schemas.openxmlformats.org/drawingml/2006/main}ext"
    )
    assert inner_extent is not None
    inner_extent.set("cx", "0")
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_picture_height_must_preserve_source_aspect_ratio(
    tmp_path: Path,
) -> None:
    """验证内外 extent 一致仍不能把图片压成不可见细线。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；外层与内层宽度都为 5cm，但两层高度均篡改为
        1 EMU 时，必须因违反源图纵横比而失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(os.fspath(source), width=Cm(5))
    outer_extent = picture_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent"
    )
    inner_extent = picture_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr/"
        "{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm/"
        "{http://schemas.openxmlformats.org/drawingml/2006/main}ext"
    )
    assert outer_extent is not None
    assert inner_extent is not None
    outer_extent.set("cy", "1")
    inner_extent.set("cy", "1")
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_image_relationship_type_requires_exact_official_uri(
    tmp_path: Path,
) -> None:
    """验证以 ``/image`` 结尾的伪关系类型不得通过。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；嵌入字节正确但 relationship type 为
        ``urn:adversarial/image`` 时必须失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(os.fspath(source), width=Cm(5))
    blip = picture_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    )
    assert blip is not None
    relationship_id = blip.get(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    )
    assert relationship_id is not None
    document.part.rels[relationship_id]._reltype = "urn:adversarial/image"
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_inner_picture_hidden_flag_is_not_visible_evidence(
    tmp_path: Path,
) -> None:
    """验证 picture 内部 ``pic:cNvPr/@hidden`` 不能被忽略。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；外层 drawing 可见，但同一 picture 的非可见属性
        声明隐藏时必须失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(os.fspath(source), width=Cm(5))
    inner_properties = picture_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPr"
    )
    assert inner_properties is not None
    inner_properties.set("hidden", "1")
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_embedded_part_requires_matching_image_content_type(
    tmp_path: Path,
) -> None:
    """验证正确图片字节不能搭配非图片 content type 通过。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；嵌入关系和 blob 正确，但 target part 声明为
        ``application/octet-stream`` 时必须 fail closed。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(os.fspath(source), width=Cm(5))
    blip = picture_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    )
    assert blip is not None
    relationship_id = blip.get(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    )
    assert relationship_id is not None
    document.part.rels[
        relationship_id
    ].target_part._content_type = "application/octet-stream"
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_off_page_anchor_is_not_visible_sequential_evidence(
    tmp_path: Path,
) -> None:
    """验证移到页面外的 anchor 不能充当占位符后的可见图片。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；图片 blob、5cm 宽度、纵横比与 OOXML 顺序均正确，
        但 ``wp:anchor`` 的水平/垂直页面偏移均为十亿 EMU 时必须失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(os.fspath(source), width=Cm(5))
    inline = picture_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
    )
    assert inline is not None
    inline.tag = (
        "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor"
    )
    for axis in ("H", "V"):
        position = OxmlElement(f"wp:position{axis}")
        position.set("relativeFrom", "page")
        offset = OxmlElement("wp:posOffset")
        offset.text = "1000000000"
        position.append(offset)
        inline.insert(0, position)
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False


def test_word010_stretch_fill_rectangle_cannot_cover_the_page(
    tmp_path: Path,
) -> None:
    """验证篡改 ``a:fillRect`` 不能把 5cm inline 图片扩展覆盖页面。

    输入参数：
        tmp_path：pytest 提供的隔离文档根目录。
    输出返回值：
        无；blob、宽度、内外几何与纵横比均正确，但
        ``a:fillRect`` 四边均为 100000 时必须失败。
    """

    images_dir = tmp_path / "images"
    images_dir.mkdir()
    source = images_dir / "Cats.png"
    source.write_bytes(_ONE_PIXEL_PNG)
    document_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(os.fspath(source), width=Cm(5))
    fill_rectangle = picture_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}stretch/"
        "{http://schemas.openxmlformats.org/drawingml/2006/main}fillRect"
    )
    assert fill_rectangle is not None
    for attribute_name in ("l", "t", "r", "b"):
        fill_rectangle.set(attribute_name, "100000")
    document.save(document_path)

    result = check_image_name_matches_doc(
        os.fspath(document_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert result["pass"] is False

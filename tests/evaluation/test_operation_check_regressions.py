"""Operation 最小检查闭包的旧审计语义回归测试。"""

from __future__ import annotations

import base64
import os
from pathlib import Path

import pytest


pytest.importorskip("docx")
pytest.importorskip("openpyxl")

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor
from openpyxl import Workbook

from paraguibench.evaluation.operation.checks.docx import (
    check_batchword002_tab_indent,
    check_has_toc,
    check_heading_palette_and_references,
    check_image_name_matches_doc,
    check_line_spacing,
    check_uppercase_words_have_parentheses,
    check_vowels_colored_red,
)
from paraguibench.evaluation.operation.checks.file import check_named_files_exist
from paraguibench.evaluation.operation.checks.file_legacy import (
    check_html_files_for_xlsx,
)
from paraguibench.evaluation.operation.checks.xlsx import (
    _rows_digest,
    check_sorted_copies_preserve_rows,
    check_values_scaled_from_source,
)


def _save_workbook(path: Path, rows: list[list[object]]) -> None:
    """把合成值行保存为有效 xlsx。

    输入参数：
        path：输出工作簿路径；rows：按行组织的单元格值。
    输出返回值：
        无；工作簿写入指定路径并关闭。
    """

    workbook = Workbook()
    worksheet = workbook.active
    for row in rows:
        worksheet.append(row)
    workbook.save(path)
    workbook.close()


def test_html_export_requires_table_and_source_values(tmp_path: Path) -> None:
    """验证 HTML 占位页不能冒充 xlsx 转换产物。

    输入参数：
        tmp_path：合成 xlsx/HTML 文件目录。
    输出返回值：
        无；占位页失败，包含源表可见值的表格通过。
    """

    _save_workbook(tmp_path / "sales.xlsx", [["Name", "Amount"], ["Alice", 1234]])
    html_path = tmp_path / "sales.html"
    html_path.write_text("<html><body>placeholder</body></html>", encoding="utf-8")
    params = {"validate_content": True, "sample_cell_limit": 4, "min_match_ratio": 1.0}
    assert not check_html_files_for_xlsx(os.fspath(tmp_path), params)["pass"]
    html_path.write_text(
        "<html><body><table><tr><td>Name</td><td>Amount</td></tr>"
        "<tr><td>Alice</td><td>1,234</td></tr></table></body></html>",
        encoding="utf-8",
    )
    assert check_html_files_for_xlsx(os.fspath(tmp_path), params)["pass"]


def test_named_files_reject_zero_byte_pdf(tmp_path: Path) -> None:
    """验证同名空 PDF 不能通过存在性评价。

    输入参数：
        tmp_path：待检查文件目录。
    输出返回值：
        无；空文件失败，具有 PDF 魔数的非空文件通过。
    """

    pdf_path = tmp_path / "report.pdf"
    pdf_path.touch()
    params = {"filenames": ["report.pdf"], "validate_format": True}
    assert not check_named_files_exist(os.fspath(tmp_path), params)["pass"]
    pdf_path.write_bytes(b"%PDF-1.7\n%%EOF\n")
    assert check_named_files_exist(os.fspath(tmp_path), params)["pass"]


def test_word_structural_regressions(tmp_path: Path) -> None:
    """验证标准首行缩进、缩写空格与伪 TOC 的旧修复语义。

    输入参数：
        tmp_path：合成 docx 输出目录。
    输出返回值：
        无；缩进和缩写通过，仅有 TOC Heading 的伪目录失败。
    """

    path = tmp_path / "sample.docx"
    document = Document()
    paragraph = document.add_paragraph(
        "MAC (Media Access Control) and API (Application Programming Interface)"
    )
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    if "TOC Heading" not in [style.name for style in document.styles]:
        document.styles.add_style("TOC Heading", WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph("Table of Contents", style="TOC Heading")
    document.save(path)
    assert check_batchword002_tab_indent(os.fspath(path), {})["pass"]
    assert check_uppercase_words_have_parentheses(os.fspath(path), {})["pass"]
    assert not check_has_toc(os.fspath(path), {})["pass"]


def test_all_red_text_does_not_satisfy_vowel_task(tmp_path: Path) -> None:
    """验证整篇标红会因辅音误标而失败。

    输入参数：
        tmp_path：合成 docx 输出目录。
    输出返回值：
        无；即使元音召回满分也不能通过精确配色约束。
    """

    path = tmp_path / "red.docx"
    document = Document()
    run = document.add_paragraph().add_run("alphabet")
    run.font.color.rgb = RGBColor(255, 0, 0)
    document.save(path)
    assert not check_vowels_colored_red(os.fspath(path), {})["pass"]


def test_embedded_image_matches_source_bytes_and_width(tmp_path: Path) -> None:
    """验证嵌入图片按固定 images 布局的源二进制与 5cm 宽度匹配。

    输入参数：
        tmp_path：根目录下存放 docx，``images`` 子目录存放同名源图。
    输出返回值：
        无；真实源图而非 OOXML 内部 imageN 名决定通过状态。
    """

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8A"
        "AQUBAScY42YAAAAASUVORK5CYII="
    )
    images_dir = tmp_path / "images"
    images_dir.mkdir()
    image_path = images_dir / "Cats.png"
    image_path.write_bytes(png_bytes)
    docx_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    document.add_paragraph().add_run().add_picture(
        os.fspath(image_path),
        width=Cm(5),
    )
    document.save(docx_path)
    result = check_image_name_matches_doc(
        os.fspath(docx_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )
    assert result["pass"], result


def test_embedded_image_rejects_case_equivalent_source_collision(
    tmp_path: Path,
) -> None:
    """验证固定 images 布局拒绝文档 stem 的大小写等价候选冲突。

    输入参数：
        tmp_path：同时存在 ``Cats.png`` 与 ``CATS.jpg`` 的隔离目录。
    输出返回值：
        无；即使 DOCX 嵌入了其中一张，源图身份不唯一也必须失败。
    """

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8A"
        "AQUBAScY42YAAAAASUVORK5CYII="
    )
    images_dir = tmp_path / "images"
    images_dir.mkdir()
    exact_image = images_dir / "Cats.png"
    exact_image.write_bytes(png_bytes)
    (images_dir / "CATS.jpg").write_bytes(png_bytes)
    docx_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    document.add_paragraph().add_run().add_picture(
        os.fspath(exact_image),
        width=Cm(5),
    )
    document.save(docx_path)

    result = check_image_name_matches_doc(
        os.fspath(docx_path),
        {
            "expected_width_cm": 5,
            "source_extensions": [".png", ".jpg"],
        },
    )

    assert not result["pass"], result


def test_embedded_image_rejects_source_extension_path_escape(tmp_path: Path) -> None:
    """验证源图扩展名不能把受控查找逃逸到 images 目录外。

    输入参数：
        tmp_path：在根目录放置嵌入图，并构造可被 ``..`` 解析的中间目录。
    输出返回值：
        无；非单段点扩展名必须失败，不得读取根目录图片。
    """

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8A"
        "AQUBAScY42YAAAAASUVORK5CYII="
    )
    images_dir = tmp_path / "images"
    images_dir.mkdir()
    (images_dir / "Cats").mkdir()
    escaped_image = tmp_path / "escape.png"
    escaped_image.write_bytes(png_bytes)
    docx_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    document.add_paragraph().add_run().add_picture(
        os.fspath(escaped_image),
        width=Cm(5),
    )
    document.save(docx_path)

    result = check_image_name_matches_doc(
        os.fspath(docx_path),
        {
            "expected_width_cm": 5,
            "source_extensions": ["/../../escape.png"],
        },
    )

    assert not result["pass"], result


def test_embedded_image_rejects_unicode_equivalent_source_collision(
    tmp_path: Path,
) -> None:
    """验证固定 images 布局拒绝 Unicode 兼容归一后的 stem 冲突。

    输入参数：
        tmp_path：同时存在 ASCII ``Cats`` 与全角 ``Ｃats`` 源图的目录。
    输出返回值：
        无；归一后的图片身份不唯一时，即使嵌入正确图也必须失败。
    """

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8A"
        "AQUBAScY42YAAAAASUVORK5CYII="
    )
    images_dir = tmp_path / "images"
    images_dir.mkdir()
    exact_image = images_dir / "Cats.png"
    exact_image.write_bytes(png_bytes)
    (images_dir / "Ｃats.jpg").write_bytes(png_bytes)
    docx_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    document.add_paragraph().add_run().add_picture(
        os.fspath(exact_image),
        width=Cm(5),
    )
    document.save(docx_path)

    result = check_image_name_matches_doc(
        os.fspath(docx_path),
        {
            "expected_width_cm": 5,
            "source_extensions": [".png", ".jpg"],
        },
    )

    assert not result["pass"], result


def test_embedded_image_rejects_symlinked_images_directory(tmp_path: Path) -> None:
    """验证 images 目录不能通过符号链接逃逸 artifact 根。

    输入参数：
        tmp_path：包含指向根外源图目录的 ``images`` 符号链接。
    输出返回值：
        无；即使外部图片与 DOCX 嵌入内容相同也必须失败。
    """

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8A"
        "AQUBAScY42YAAAAASUVORK5CYII="
    )
    outside_images = tmp_path.parent / f"{tmp_path.name}-outside-images"
    outside_images.mkdir()
    outside_image = outside_images / "Cats.png"
    outside_image.write_bytes(png_bytes)
    (tmp_path / "images").symlink_to(outside_images, target_is_directory=True)
    docx_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    document.add_paragraph().add_run().add_picture(
        os.fspath(outside_image),
        width=Cm(5),
    )
    document.save(docx_path)

    result = check_image_name_matches_doc(
        os.fspath(docx_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert not result["pass"], result


def test_embedded_image_requires_matching_bytes_and_width_on_same_drawing(
    tmp_path: Path,
) -> None:
    """验证同名图片内容与 5cm 宽度必须属于同一个 drawing。

    输入参数：
        tmp_path：嵌入“正确内容+错宽度”和“错内容+5cm”两张图的 DOCX 目录。
    输出返回值：
        无；两个不同 drawing 的属性不得拼接成虚假通过。
    """

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8A"
        "AQUBAScY42YAAAAASUVORK5CYII="
    )
    images_dir = tmp_path / "images"
    images_dir.mkdir()
    matching_image = images_dir / "Cats.png"
    matching_image.write_bytes(png_bytes)
    distractor_image = tmp_path / "distractor.png"
    distractor_image.write_bytes(png_bytes + b"distractor")
    docx_path = tmp_path / "Cats.docx"
    document = Document()
    document.add_paragraph("[image Slot]")
    document.add_paragraph().add_run().add_picture(
        os.fspath(matching_image),
        width=Cm(3),
    )
    document.add_paragraph().add_run().add_picture(
        os.fspath(distractor_image),
        width=Cm(5),
    )
    document.save(docx_path)

    result = check_image_name_matches_doc(
        os.fspath(docx_path),
        {"expected_width_cm": 5, "source_extensions": [".png"]},
    )

    assert not result["pass"], result


def test_double_line_spacing_accepts_paragraph_style_inheritance(
    tmp_path: Path,
) -> None:
    """验证段落通过 OOXML 样式继承获得双倍行距时可通过。

    输入参数：
        tmp_path：合成 DOCX 输出目录。
    输出返回值：
        无；段落没有直接行距，但其有效样式设为 2.0 时必须通过。
    """

    path = tmp_path / "style-spacing.docx"
    document = Document()
    paragraph = document.add_paragraph("style inherited body")
    paragraph.style.paragraph_format.line_spacing = 2.0
    document.save(path)

    result = check_line_spacing(
        os.fspath(path),
        {"spacing": 2.0, "threshold": 0.9},
    )

    assert result["pass"], result


def test_double_line_spacing_accepts_base_style_chain_inheritance(
    tmp_path: Path,
) -> None:
    """验证派生段落样式可从 base style 继承双倍行距。

    输入参数：
        tmp_path：合成带两层段落样式链的 DOCX 输出目录。
    输出返回值：
        无；当前样式未设行距、基样式设为 2.0 时必须通过。
    """

    path = tmp_path / "base-style-spacing.docx"
    document = Document()
    base_style = document.styles.add_style("Double Base", WD_STYLE_TYPE.PARAGRAPH)
    base_style.paragraph_format.line_spacing = 2.0
    derived_style = document.styles.add_style(
        "Double Derived",
        WD_STYLE_TYPE.PARAGRAPH,
    )
    derived_style.base_style = base_style
    document.add_paragraph("base style inherited body", style=derived_style)
    document.save(path)

    result = check_line_spacing(
        os.fspath(path),
        {"spacing": 2.0, "threshold": 0.9},
    )

    assert result["pass"], result


def test_double_line_spacing_accepts_docdefaults_inheritance(tmp_path: Path) -> None:
    """验证段落可从 styles.xml ``docDefaults`` 继承双倍行距。

    输入参数：
        tmp_path：合成含 ``w:pPrDefault/w:spacing`` 的 DOCX 输出目录。
    输出返回值：
        无；段落和样式都未设行距、文档默认值为 2.0 时必须通过。
    """

    path = tmp_path / "docdefaults-spacing.docx"
    document = Document()
    document.add_paragraph("docDefaults inherited body")
    styles = document.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    ppr_default = doc_defaults.find(qn("w:pPrDefault"))
    if ppr_default is None:
        ppr_default = OxmlElement("w:pPrDefault")
        doc_defaults.append(ppr_default)
    ppr = ppr_default.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        ppr_default.append(ppr)
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:line"), "480")
    spacing.set(qn("w:lineRule"), "auto")
    document.save(path)

    result = check_line_spacing(
        os.fspath(path),
        {"spacing": 2.0, "threshold": 0.9},
    )

    assert result["pass"], result


def test_double_line_spacing_does_not_ignore_table_cell_paragraphs(
    tmp_path: Path,
) -> None:
    """验证只有表格单元格文本的文档不会因分母为零而虚假通过。

    输入参数：
        tmp_path：合成只含一个单倍行距表格段落的 DOCX 输出目录。
    输出返回值：
        无；表格内可见文本必须进入固定分母并使双倍行距规则失败。
    """

    path = tmp_path / "table-spacing.docx"
    document = Document()
    paragraph = document.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
    paragraph.text = "table body"
    paragraph.paragraph_format.line_spacing = 1.0
    document.save(path)

    result = check_line_spacing(
        os.fspath(path),
        {"spacing": 2.0, "threshold": 0.9},
    )

    assert not result["pass"], result


def test_referenced_documents_require_different_heading_colors(
    tmp_path: Path,
) -> None:
    """验证互引文档不能使用同一主标题颜色。

    输入参数：
        tmp_path：两个互引 docx 的隔离目录。
    输出返回值：
        无；异色通过，把第二份改为同色后失败。
    """

    for filename, reference, color in [
        ("Doc_A.docx", "Doc_B", RGBColor(255, 0, 0)),
        ("Doc_B.docx", "Doc_A", RGBColor(0, 0, 255)),
    ]:
        document = Document()
        title = document.add_paragraph().add_run(f"Title {filename}")
        title.font.bold = True
        title.font.color.rgb = color
        document.add_paragraph(f"References: see {reference} for details.")
        document.save(tmp_path / filename)
    params = {
        "expected_files": ["Doc_A.docx", "Doc_B.docx"],
        "palette": ["FF0000", "0000FF"],
    }
    assert check_heading_palette_and_references(os.fspath(tmp_path), params)["pass"]
    document = Document(tmp_path / "Doc_B.docx")
    document.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
    document.save(tmp_path / "Doc_B.docx")
    assert not check_heading_palette_and_references(os.fspath(tmp_path), params)["pass"]


def test_currency_scaling_rejects_wrong_divisor(tmp_path: Path) -> None:
    """验证浮点类型或除以 1000 不能冒充元转万元。

    输入参数：
        tmp_path：合成工作簿目录。
    输出返回值：
        无；错误倍率失败，按 10000 换算的值通过。
    """

    path = tmp_path / "store.xlsx"
    params = {
        "start_cell": "B2",
        "end_cell": "B3",
        "divisor": 10000,
        "source_values_by_file": {"store.xlsx": [10000, 25000]},
    }
    _save_workbook(path, [["month", "amount"], [1, 10.0], [2, 25.0]])
    assert not check_values_scaled_from_source(os.fspath(path), params)["pass"]
    _save_workbook(path, [["month", "amount"], [1, 1.0], [2, 2.5]])
    assert check_values_scaled_from_source(os.fspath(path), params)["pass"]


def test_sorted_copies_preserve_complete_source_rows(tmp_path: Path) -> None:
    """验证四个排序副本完整保留源行并一对一覆盖四列。

    输入参数：
        tmp_path：源工作簿与排序副本目录。
    输出返回值：
        无；完整副本通过，替换一个副本的行多重集后失败。
    """

    header = [["A", "B", "C", "D"]]
    rows = [
        (3, "c", 30, 300),
        (1, "d", 10, 400),
        (4, "a", 40, 100),
        (2, "b", 20, 200),
    ]
    _save_workbook(tmp_path / "source.xlsx", header + [list(row) for row in rows])
    for index in range(4):
        sorted_rows = sorted(rows, key=lambda row: row[index])
        _save_workbook(
            tmp_path / f"sort_{index}.xlsx",
            header + [list(row) for row in sorted_rows],
        )
    params = {
        "source_filename": "source.xlsx",
        "header_row": 1,
        "column_count": 4,
        "expected_data_rows": 4,
        "expected_total_rows": 4,
        "source_data_sha256": _rows_digest(rows),
        "required_sorts": [{"column": column, "order": "asc"} for column in "ABCD"],
    }
    assert check_sorted_copies_preserve_rows(os.fspath(tmp_path), params)["pass"]
    _save_workbook(
        tmp_path / "sort_0.xlsx",
        header + [[1, "x", 1, 1], [2, "y", 2, 2]],
    )
    assert not check_sorted_copies_preserve_rows(os.fspath(tmp_path), params)["pass"]

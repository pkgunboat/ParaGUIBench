"""Word-009/010 pre→post typed DOCX 正文保真协议测试。"""

from __future__ import annotations

import base64
from collections.abc import Callable
from dataclasses import replace
import hashlib
from pathlib import Path
import shutil
import xml.etree.ElementTree as ET
import zipfile

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from paraguibench.evaluation.operation.word_text_fidelity import (
    WordTextInputFile,
    WordTextFidelityError,
    capture_word_text_baseline,
    compare_word_text_fidelity,
    validate_word_text_baseline_identity,
)


_TASK_ID = "Operation-FileOperate-BatchOperationWord-009"
_WORD010_TASK_ID = "Operation-FileOperate-BatchOperationWord-010"
_MANIFEST_SHA256 = "1" * 64
_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8A"
    "AQUBAScY42YAAAAASUVORK5CYII="
)


def _write_document(path: Path, text: str) -> None:
    """写入一份带可见正文的最小 DOCX。

    输入参数：
        path：目标 DOCX 路径；text：唯一正文字符串。
    输出返回值：
        无；生成的文档可由生产 typed snapshot 解析。
    """

    document = Document()
    document.add_paragraph(text)
    document.save(path)


def _identity(path: Path) -> WordTextInputFile:
    """从同一个合成 DOCX 构造固定输入身份。

    输入参数：
        path：待固定的测试文档。
    输出返回值：
        包含相对文件名、大小、SHA-256 与 DOCX 类型的身份。
    """

    payload = path.read_bytes()
    return WordTextInputFile(
        path=path.name,
        size=len(payload),
        sha256=hashlib.sha256(payload).hexdigest(),
        is_docx=True,
    )


def _write_anchor_hyperlink_document(path: Path, anchor: str) -> None:
    """写入一份使用内部 bookmark anchor 的超链接 DOCX。

    输入参数：
        path：目标 DOCX；anchor：``w:hyperlink/@w:anchor`` 值。
    输出返回值：
        无；超链接可见文字固定为 ``LINK TEXT``。
    """

    document = Document()
    paragraph = document.add_paragraph()
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "LINK TEXT"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    document.save(path)


def _write_external_hyperlink_document(path: Path) -> None:
    """写入一份通过 external relationship 承载目标的超链接 DOCX。

    输入参数：
        path：目标 DOCX 路径。
    输出返回值：
        无；生成的超链接文字为 ``EXTERNAL LINK``。
    """

    document = Document()
    paragraph = document.add_paragraph()
    relationship_id = document.part.relate_to(
        "https://example.invalid/private",
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "EXTERNAL LINK"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    document.save(path)


def _rewrite_docx_xml(
    path: Path,
    member_name: str,
    mutate: Callable[[ET.Element], None],
) -> None:
    """在保留 ZIP 闭集的情况下改写一个 DOCX XML member。

    输入参数：
        path：待改写 DOCX；member_name：包内相对路径；
        mutate：接收 ``Element`` 根并就地修改的测试 callable。
    输出返回值：
        无；原子替换同一测试文件。
    """

    with zipfile.ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    rewritten: list[tuple[zipfile.ZipInfo, bytes]] = []
    found = False
    for info, payload in entries:
        if info.filename == member_name:
            root = ET.fromstring(payload)
            mutate(root)
            payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            found = True
        rewritten.append((info, payload))
    assert found
    temporary = path.with_suffix(".rewrite.docx")
    with zipfile.ZipFile(temporary, "w") as archive:
        for info, payload in rewritten:
            archive.writestr(info, payload)
    temporary.replace(path)


def _append_docx_member(path: Path, member_name: str, payload: bytes) -> None:
    """向合成 DOCX 追加一个不重名 member。

    输入参数：
        path：待改写 DOCX；member_name：新的 package 相对路径；
        payload：要写入的完整字节。
    输出返回值：
        无；保留所有现有 member 并原子替换测试文档。
    """

    with zipfile.ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    assert member_name not in {info.filename for info, _payload in entries}
    temporary = path.with_suffix(".append.docx")
    with zipfile.ZipFile(temporary, "w") as archive:
        for info, existing_payload in entries:
            archive.writestr(info, existing_payload)
        archive.writestr(member_name, payload)
    temporary.replace(path)


def test_word009_changed_body_text_fails_pre_post_fidelity(tmp_path: Path) -> None:
    """验证目标行距操作不能伪装被改写的正文。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；正式 pre snapshot 与改字后 post DOCX 必须不匹配。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    result_path = result_root / source_path.name
    _write_document(source_path, "ORIGINAL VISIBLE BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    _write_document(result_path, "TAMPERED VISIBLE BODY")

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False
    assert fidelity.document_count == 1


def test_same_text_moved_from_paragraph_to_table_cell_fails_fidelity(
    tmp_path: Path,
) -> None:
    """验证只拼接文字的假保真不能隐藏容器迁移。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；字符序列相同但从正文段落移入表格单元格时必须失败。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "SAME TEXT")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post = Document()
    post.add_table(rows=1, cols=1).cell(0, 0).text = "SAME TEXT"
    post.save(result_root / source_path.name)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_equivalent_run_split_preserves_typed_text_fidelity(tmp_path: Path) -> None:
    """验证不改文字与样式的 run 拆分不产生假失败。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；pre 单 run 和 post 同样式双 run 的连续正文应等价。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "SAME TEXT")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post = Document()
    paragraph = post.add_paragraph()
    paragraph.add_run("SAME ")
    paragraph.add_run("TEXT")
    post.save(result_root / source_path.name)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is True


def test_same_text_changed_to_hidden_fails_fidelity(tmp_path: Path) -> None:
    """验证 ``w:vanish`` 不能在文字未变时绕过保真。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；post 把同一正文改为隐藏后必须不匹配。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "VISIBLE BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post = Document()
    run = post.add_paragraph().add_run("VISIBLE BODY")
    run.font.hidden = True
    post.save(result_root / source_path.name)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_same_text_moved_into_deleted_revision_fails_fidelity(tmp_path: Path) -> None:
    """验证修订删除 wrapper 的语义不能被裸文字摘要忽略。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；精确文字未改但被包入 ``w:del`` 后必须失败。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "REVISION TEXT")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post = Document()
    paragraph = post.add_paragraph()
    run = paragraph.add_run("REVISION TEXT")
    paragraph._p.remove(run._r)
    deletion = OxmlElement("w:del")
    deletion.append(run._r)
    paragraph._p.append(deletion)
    post.save(result_root / source_path.name)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_deleted_header_text_fails_fidelity(tmp_path: Path) -> None:
    """验证主文档之外的页眉正文也属于固定 baseline。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；post 保留主文档却删除 header 文字时必须失败。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    source = Document()
    source.add_paragraph("BODY")
    source.sections[0].header.paragraphs[0].text = "PRIVATE HEADER"
    source.save(source_path)
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    _write_document(result_root / source_path.name, "BODY")

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_hyperlink_anchor_change_fails_fidelity(tmp_path: Path) -> None:
    """验证超链接可见文字未变时目标漂移仍不保真。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；``bookmark-a`` 改为 ``bookmark-b`` 必须失败。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_anchor_hyperlink_document(source_path, "bookmark-a")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    _write_anchor_hyperlink_document(result_root / source_path.name, "bookmark-b")

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_added_office_math_text_fails_fidelity(tmp_path: Path) -> None:
    """验证 ``m:t`` 数学文字不是 ``w:t`` 扫描的盲区。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；主文本不变但新增可见 Office Math 文字必须失败。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post = Document()
    paragraph = post.add_paragraph("BODY")
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "x+1"
    math_run.append(math_text)
    math.append(math_run)
    paragraph._p.append(math)
    post.save(result_root / source_path.name)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_altchunk_visible_text_channel_fails_closed(tmp_path: Path) -> None:
    """验证 evaluator 不会忽略尚未解析的 altChunk 文字载体。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；post 出现 ``w:altChunk`` 时必须抛固定脱敏解析错误。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post = Document()
    post.add_paragraph("BODY")
    post.element.body.insert(-1, OxmlElement("w:altChunk"))
    post.save(result_root / source_path.name)

    with pytest.raises(WordTextFidelityError) as captured:
        compare_word_text_fidelity(baseline, result_root)

    assert captured.value.code == "WORD_TEXT_UNSUPPORTED_CARRIER"


def test_unknown_namespace_text_carrier_fails_closed(tmp_path: Path) -> None:
    """验证未知 namespace 中的非空文字不会被静默忽略。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；post 主文档在原正文外追加未知 XML 文字叶时，
        evaluator 必须以固定错误码 ERROR，不得 matched=True。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _append_unknown_text(root: ET.Element) -> None:
        """在主文档 body 内追加未知可见文字叶。"""

        body = root.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body"
        )
        assert body is not None
        carrier = ET.Element("{urn:paraguibench:unknown}visibleText")
        carrier.text = "UNKNOWN VISIBLE TEXT"
        body.insert(-1, carrier)

    _rewrite_docx_xml(post_path, "word/document.xml", _append_unknown_text)

    with pytest.raises(WordTextFidelityError) as captured:
        compare_word_text_fidelity(baseline, result_root)

    assert captured.value.code == "WORD_TEXT_UNKNOWN_TEXT_CARRIER"


def test_unknown_namespace_wrapper_around_text_fails_closed(
    tmp_path: Path,
) -> None:
    """验证未知 wrapper 不能作为文字子树的透明通道。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；post 只把原 ``w:r/w:t`` 包进未知 namespace
        节点，字符不变时也必须固定 ERROR。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _wrap_run(root: ET.Element) -> None:
        """用无文字的未知节点包裹第一个 run。"""

        namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        paragraph = root.find(f".//{{{namespace}}}p")
        assert paragraph is not None
        run = paragraph.find(f"{{{namespace}}}r")
        assert run is not None
        paragraph.remove(run)
        wrapper = ET.Element("{urn:paraguibench:unknown}wrapper")
        wrapper.append(run)
        paragraph.append(wrapper)

    _rewrite_docx_xml(post_path, "word/document.xml", _wrap_run)

    with pytest.raises(WordTextFidelityError) as captured:
        compare_word_text_fidelity(baseline, result_root)

    assert captured.value.code == "WORD_TEXT_UNKNOWN_TEXT_CARRIER"


def test_alternate_content_branches_are_not_concatenated_as_plain_text(
    tmp_path: Path,
) -> None:
    """验证 mc:AlternateContent 两个互斥分支不会串联伪装正文。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；pre 的普通 ``AB`` 被改为 Choice=`A`、Fallback=`B`
        时，即使递归字符串拼接仍为 ``AB`` 也必须不匹配。
    """

    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    compatibility_namespace = (
        "http://schemas.openxmlformats.org/markup-compatibility/2006"
    )
    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "AB")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _replace_run_with_branches(root: ET.Element) -> None:
        """用 Choice/Fallback 两个互斥分支替换原 run。"""

        paragraph = root.find(f".//{{{word_namespace}}}p")
        assert paragraph is not None
        run = paragraph.find(f"{{{word_namespace}}}r")
        assert run is not None
        paragraph.remove(run)
        alternate = ET.Element(f"{{{compatibility_namespace}}}AlternateContent")
        choice = ET.SubElement(
            alternate,
            f"{{{compatibility_namespace}}}Choice",
        )
        choice.set("Requires", "w14")
        fallback = ET.SubElement(
            alternate,
            f"{{{compatibility_namespace}}}Fallback",
        )
        for branch, value in ((choice, "A"), (fallback, "B")):
            branch_run = ET.SubElement(branch, f"{{{word_namespace}}}r")
            text = ET.SubElement(branch_run, f"{{{word_namespace}}}t")
            text.text = value
        paragraph.append(alternate)

    _rewrite_docx_xml(
        post_path,
        "word/document.xml",
        _replace_run_with_branches,
    )

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_alternate_content_requires_drift_fails_fidelity(
    tmp_path: Path,
) -> None:
    """验证 mc:Choice/@Requires 属于分支选择语义。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；两个分支的标签与文字不变，只改
        ``Requires`` 值时 typed 语义仍必须不匹配。
    """

    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    compatibility_namespace = (
        "http://schemas.openxmlformats.org/markup-compatibility/2006"
    )
    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "AB")

    def _wrap_with_alternate_content(root: ET.Element) -> None:
        """写入含固定 Requires 的 Choice/Fallback 容器。"""

        paragraph = root.find(f".//{{{word_namespace}}}p")
        assert paragraph is not None
        run = paragraph.find(f"{{{word_namespace}}}r")
        assert run is not None
        paragraph.remove(run)
        alternate = ET.Element(f"{{{compatibility_namespace}}}AlternateContent")
        choice = ET.SubElement(
            alternate,
            f"{{{compatibility_namespace}}}Choice",
        )
        choice.set("Requires", "w14")
        fallback = ET.SubElement(
            alternate,
            f"{{{compatibility_namespace}}}Fallback",
        )
        choice.append(run)
        fallback_run = ET.SubElement(fallback, f"{{{word_namespace}}}r")
        fallback_text = ET.SubElement(
            fallback_run,
            f"{{{word_namespace}}}t",
        )
        fallback_text.text = "AB"
        paragraph.append(alternate)

    _rewrite_docx_xml(
        source_path,
        "word/document.xml",
        _wrap_with_alternate_content,
    )
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _change_requires(root: ET.Element) -> None:
        """只改写 Choice 的 Requires 字面值。"""

        choice = root.find(f".//{{{compatibility_namespace}}}Choice")
        assert choice is not None
        choice.set("Requires", "w15")

    _rewrite_docx_xml(post_path, "word/document.xml", _change_requires)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_list_numbering_semantics_change_fails_fidelity(tmp_path: Path) -> None:
    """验证字符未变时删除列表编号样式仍属于语义漂移。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；pre ``List Number`` 段落改为普通段落必须失败。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    source = Document()
    source.add_paragraph("NUMBERED ITEM", style="List Number")
    source.save(source_path)
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    _write_document(result_root / source_path.name, "NUMBERED ITEM")

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_referenced_style_visibility_change_fails_fidelity(tmp_path: Path) -> None:
    """验证仅修改被引用样式的 hidden 属性也会改变 typed 语义。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；run 的 ``rStyle`` ID 与文字均未变，但样式定义
        从可见改为隐藏时必须失败。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    for path, hidden in (
        (source_path, False),
        (result_root / source_path.name, True),
    ):
        document = Document()
        style = document.styles.add_style("Semantic Style", WD_STYLE_TYPE.CHARACTER)
        style.font.hidden = hidden
        run = document.add_paragraph().add_run("STYLE TEXT")
        run.style = style
        document.save(path)
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_numbering_definition_change_fails_fidelity(tmp_path: Path) -> None:
    """验证 document/style 引用未变时 numbering 定义漂移仍不保真。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；只把 ``word/numbering.xml`` 的 ``lvlText`` 改为另一标签
        格式时必须失败。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    source = Document()
    source.add_paragraph("NUMBERED ITEM", style="List Number")
    source.save(source_path)
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _change_label(root: ET.Element) -> None:
        """改写 ListNumber 引用的 abstract 列表标签。"""

        namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        abstract_id = None
        for instance in root.findall(f"{{{namespace}}}num"):
            if instance.get(f"{{{namespace}}}numId") != "5":
                continue
            reference = instance.find(f"{{{namespace}}}abstractNumId")
            assert reference is not None
            abstract_id = reference.get(f"{{{namespace}}}val")
            break
        assert abstract_id is not None
        for abstract in root.findall(f"{{{namespace}}}abstractNum"):
            if abstract.get(f"{{{namespace}}}abstractNumId") != abstract_id:
                continue
            level_text = abstract.find(f"{{{namespace}}}lvl/{{{namespace}}}lvlText")
            assert level_text is not None
            level_text.set(f"{{{namespace}}}val", "Section %1")
            return
        raise AssertionError("missing referenced abstract numbering")

    _rewrite_docx_xml(post_path, "word/numbering.xml", _change_label)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_unreferenced_numbering_definition_change_preserves_fidelity(
    tmp_path: Path,
) -> None:
    """验证未被可见段落引用的默认列表不造成假失败。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；普通段落没有 ``numPr``，仅改写 python-docx 附带
        但未引用的 ``numbering.xml/lvlText`` 时 typed 正文仍匹配。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "PLAIN BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _change_unused_label(root: ET.Element) -> None:
        """改写第一个未被正文引用的列表标签。"""

        namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        level_text = root.find(f".//{{{namespace}}}lvlText")
        assert level_text is not None
        level_text.set(f"{{{namespace}}}val", "UNUSED %1")

    _rewrite_docx_xml(post_path, "word/numbering.xml", _change_unused_label)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is True


def test_equivalent_numbering_id_remap_preserves_fidelity(tmp_path: Path) -> None:
    """验证列表 ID 等价重编号不造成假失败。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；``List Number`` 样式的 numId 与对应 ``w:num``
        同时从 5 改为 99，但实例和 abstract level 语义不变时匹配。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    source = Document()
    source.add_paragraph("NUMBERED ITEM", style="List Number")
    source.save(source_path)
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _remap_instance(root: ET.Element) -> None:
        """将 numId=5 的列表实例重编号为 99。"""

        for instance in root.findall(f"{{{namespace}}}num"):
            if instance.get(f"{{{namespace}}}numId") == "5":
                instance.set(f"{{{namespace}}}numId", "99")
                return
        raise AssertionError("missing numId=5")

    def _remap_style_reference(root: ET.Element) -> None:
        """将 ListNumber 样式的 numPr 引用同步改为 99。"""

        for style in root.findall(f"{{{namespace}}}style"):
            if style.get(f"{{{namespace}}}styleId") != "ListNumber":
                continue
            number_id = style.find(
                f"{{{namespace}}}pPr/{{{namespace}}}numPr/{{{namespace}}}numId"
            )
            assert number_id is not None
            number_id.set(f"{{{namespace}}}val", "99")
            return
        raise AssertionError("missing ListNumber style")

    _rewrite_docx_xml(post_path, "word/numbering.xml", _remap_instance)
    _rewrite_docx_xml(post_path, "word/styles.xml", _remap_style_reference)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is True


def test_word009_adding_only_double_spacing_preserves_fidelity(
    tmp_path: Path,
) -> None:
    """验证 009 的唯一目标差异不会被 pPr 空壳误判。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；pre 没有 ``w:pPr``，post 只新增双倍 ``w:spacing``
        时 typed 正文必须仍匹配。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post = Document()
    paragraph = post.add_paragraph("BODY")
    paragraph.paragraph_format.line_spacing = 2.0
    post.save(result_root / source_path.name)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is True


def test_word010_new_drawing_only_paragraph_preserves_text_fidelity(
    tmp_path: Path,
) -> None:
    """验证 010 允许为目标图片新增 drawing-only 段落。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录与图片。
    输出返回值：
        无；原文字和容器未变，仅多一个内联图片段落时
        typed 正文必须仍匹配。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "[image Slot]")
    baseline = capture_word_text_baseline(
        task_id=_WORD010_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    image_path = tmp_path / "Source.png"
    image_path.write_bytes(_ONE_PIXEL_PNG)
    post = Document()
    post.add_paragraph("[image Slot]")
    post.add_paragraph().add_run().add_picture(str(image_path), width=Cm(5))
    post.save(result_root / source_path.name)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is True


def test_word010_drawing_only_paragraph_rejects_unknown_text_carrier(
    tmp_path: Path,
) -> None:
    """验证 010 的 drawing-only 豁免不会隐藏未知文字。

    输入参数：
        tmp_path：pytest 提供的 pre/post 根与合成图片。
    输出返回值：
        无；合法内联图片段落中追加未知 namespace
        非空文字时必须 ERROR，不得被整段快捷跳过。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "[image Slot]")
    baseline = capture_word_text_baseline(
        task_id=_WORD010_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    image_path = tmp_path / "Source.png"
    image_path.write_bytes(_ONE_PIXEL_PNG)
    post_path = result_root / source_path.name
    post = Document()
    post.add_paragraph("[image Slot]")
    post.add_paragraph().add_run().add_picture(str(image_path), width=Cm(5))
    post.save(post_path)

    def _inject_unknown_drawing_text(root: ET.Element) -> None:
        """向最后一个图片段落注入未知文字叶。"""

        namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        paragraphs = root.findall(f".//{{{namespace}}}p")
        assert paragraphs
        carrier = ET.Element("{urn:paraguibench:unknown}visibleText")
        carrier.text = "INJECTED VISIBLE TEXT"
        paragraphs[-1].append(carrier)

    _rewrite_docx_xml(
        post_path,
        "word/document.xml",
        _inject_unknown_drawing_text,
    )

    with pytest.raises(WordTextFidelityError) as captured:
        compare_word_text_fidelity(baseline, result_root)

    assert captured.value.code == "WORD_TEXT_UNKNOWN_TEXT_CARRIER"


@pytest.mark.parametrize("mutation", ("wrong_type", "missing", "duplicate"))
def test_main_office_document_relationship_is_unique_and_canonical(
    tmp_path: Path,
    mutation: str,
) -> None:
    """验证 package root 唯一正式 main-document 关系。

    输入参数：
        tmp_path：私有 pre/post 根；mutation：把根关系的
        Type 改错、删除或复制为第二条正式边。
    输出返回值：
        无；即使目标仍指向 ``word/document.xml``，三类
        package 身份漂移也必须以固定 ERROR 拒绝。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _mutate_main_relationship(root: ET.Element) -> None:
        """对唯一 officeDocument 关系执行指定变体。"""

        relationships = [
            relationship
            for relationship in root
            if relationship.get("Type", "").endswith("/officeDocument")
        ]
        assert len(relationships) == 1
        relationship = relationships[0]
        if mutation == "wrong_type":
            relationship.set("Type", "urn:paraguibench:wrong-document")
        elif mutation == "missing":
            root.remove(relationship)
        else:
            duplicate = ET.fromstring(ET.tostring(relationship))
            duplicate.set("Id", "rIdParaGUIDuplicate")
            root.append(duplicate)

    _rewrite_docx_xml(post_path, "_rels/.rels", _mutate_main_relationship)

    with pytest.raises(WordTextFidelityError) as captured:
        compare_word_text_fidelity(baseline, result_root)

    assert captured.value.code == "WORD_TEXT_MAIN_RELATIONSHIP_INVALID"


@pytest.mark.parametrize("mutation", ("wrong_type", "missing", "duplicate"))
def test_main_document_content_type_is_unique_and_canonical(
    tmp_path: Path,
    mutation: str,
) -> None:
    """验证 Content Types 唯一正确声明 Word main part。

    输入参数：
        tmp_path：私有 pre/post 根；mutation：改错、删除或
        复制 ``/word/document.xml`` Override。
    输出返回值：
        无；主文档 content type 不再唯一且正式时必须
        固定 ERROR，不得仅依赖 ZIP member 名。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _mutate_main_override(root: ET.Element) -> None:
        """对 main-document Override 执行指定变体。"""

        overrides = [
            child for child in root if child.get("PartName") == "/word/document.xml"
        ]
        assert len(overrides) == 1
        override = overrides[0]
        if mutation == "wrong_type":
            override.set("ContentType", "application/x-invalid")
        elif mutation == "missing":
            root.remove(override)
        else:
            root.append(ET.fromstring(ET.tostring(override)))

    _rewrite_docx_xml(post_path, "[Content_Types].xml", _mutate_main_override)

    with pytest.raises(WordTextFidelityError) as captured:
        compare_word_text_fidelity(baseline, result_root)

    assert captured.value.code == "WORD_TEXT_CONTENT_TYPES_INVALID"


def test_reachable_header_content_type_drift_fails_fidelity(
    tmp_path: Path,
) -> None:
    """验证可达页眉 part 的有效 ContentType 属于文字语义。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；页眉 XML、文字和 relationship 均不变，仅把
        ``/word/header1.xml`` Override 改为无效类型时必须不匹配。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    document = Document()
    document.add_paragraph("BODY")
    document.sections[0].header.paragraphs[0].text = "HEADER"
    document.save(source_path)
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _change_header_content_type(root: ET.Element) -> None:
        """只改写页眉 Override 的 MIME 类型。"""

        overrides = [
            child for child in root if child.get("PartName") == "/word/header1.xml"
        ]
        assert len(overrides) == 1
        overrides[0].set("ContentType", "application/x-invalid-header")

    _rewrite_docx_xml(
        post_path,
        "[Content_Types].xml",
        _change_header_content_type,
    )

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_header_content_type_part_name_case_drift_fails_fidelity(
    tmp_path: Path,
) -> None:
    """验证 OPC Override PartName 不会被大小写折叠。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；仅把页眉 Override 的 PartName 改为大写路径，
        ZIP member 仍为小写时有效 ContentType 必须漂移。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    document = Document()
    document.add_paragraph("BODY")
    document.sections[0].header.paragraphs[0].text = "HEADER"
    document.save(source_path)
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _uppercase_header_part_name(root: ET.Element) -> None:
        """只改写页眉 Override PartName 的大小写。"""

        overrides = [
            child for child in root if child.get("PartName") == "/word/header1.xml"
        ]
        assert len(overrides) == 1
        overrides[0].set("PartName", "/WORD/HEADER1.XML")

    _rewrite_docx_xml(
        post_path,
        "[Content_Types].xml",
        _uppercase_header_part_name,
    )

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_header_part_root_qname_drift_fails_closed(tmp_path: Path) -> None:
    """验证可达页眉 part 必须保留 canonical 根 QName。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；仅把 ``header1.xml`` 根从 ``w:hdr`` 改成
        ``w:ftr``，关系、ContentType 和文字都不变时仍必须 ERROR。
    """

    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    document = Document()
    document.add_paragraph("BODY")
    document.sections[0].header.paragraphs[0].text = "HEADER"
    document.save(source_path)
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _change_header_root(root: ET.Element) -> None:
        """把页眉根 QName 改成 footer，保留全部子树。"""

        root.tag = f"{{{namespace}}}ftr"

    _rewrite_docx_xml(post_path, "word/header1.xml", _change_header_root)

    with pytest.raises(WordTextFidelityError) as captured:
        compare_word_text_fidelity(baseline, result_root)

    assert captured.value.code == "WORD_TEXT_PART_ROOT_INVALID"


def test_styles_and_numbering_root_qname_drift_fail_closed(
    tmp_path: Path,
) -> None:
    """验证语义辅助 part 不能伪装根 QName。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；``word/styles.xml`` 或 ``word/numbering.xml`` 只改
        根 QName 时都必须以同一固定结构 ERROR 拒绝。
    """

    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    document = Document()
    document.add_paragraph("NUMBERED BODY", style="List Number")
    document.save(source_path)
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    for member_name, wrong_root in (
        ("word/styles.xml", "numbering"),
        ("word/numbering.xml", "styles"),
    ):
        shutil.copy2(source_path, post_path)

        def _change_root(root: ET.Element) -> None:
            """仅改写当前语义 part 的根 QName。"""

            root.tag = f"{{{namespace}}}{wrong_root}"

        _rewrite_docx_xml(post_path, member_name, _change_root)
        with pytest.raises(WordTextFidelityError) as captured:
            compare_word_text_fidelity(baseline, result_root)
        assert captured.value.code == "WORD_TEXT_PART_ROOT_INVALID"


def test_simple_field_wrapper_changes_text_semantics(tmp_path: Path) -> None:
    """验证字符不变时 simple field wrapper 仍属语义漂移。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；把同一 ``w:r/w:t`` 包入带 ``w:instr`` 的
        ``w:fldSimple`` 后 typed 快照必须不匹配。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _wrap_run(root: ET.Element) -> None:
        """把第一个段落的文字 run 包入 simple field。"""

        namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        paragraph = root.find(f".//{{{namespace}}}p")
        assert paragraph is not None
        run = paragraph.find(f"{{{namespace}}}r")
        assert run is not None
        paragraph.remove(run)
        field = ET.Element(f"{{{namespace}}}fldSimple")
        field.set(f"{{{namespace}}}instr", "DATE")
        field.append(run)
        paragraph.append(field)

    _rewrite_docx_xml(post_path, "word/document.xml", _wrap_run)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_same_simple_field_wrapper_preserves_text_fidelity(tmp_path: Path) -> None:
    """验证未改变的 simple field 不产生假失败。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；pre/post 保留相同 ``w:fldSimple/@w:instr``
        与文字时必须匹配。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    document = Document()
    paragraph = document.add_paragraph()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "DATE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "BODY"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)
    document.save(source_path)
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    shutil.copy2(source_path, result_root / source_path.name)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is True


def test_footnote_entry_id_swap_changes_reference_semantics(
    tmp_path: Path,
) -> None:
    """验证脚注文字顺序不变时 entry ID 映射仍必须固定。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根。
    输出返回值：
        无；只交换脚注 A/B 容器的 ``w:id``，保持文字
        与 part 顺序不变，typed 快照仍必须不匹配。
    """

    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    package_relationship_namespace = (
        "http://schemas.openxmlformats.org/package/2006/relationships"
    )
    content_types_namespace = (
        "http://schemas.openxmlformats.org/package/2006/content-types"
    )
    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "BODY")

    def _add_body_reference(root: ET.Element) -> None:
        """向正文段落追加指向脚注 1 的引用。"""

        paragraph = root.find(f".//{{{word_namespace}}}p")
        assert paragraph is not None
        run = ET.SubElement(paragraph, f"{{{word_namespace}}}r")
        reference = ET.SubElement(
            run,
            f"{{{word_namespace}}}footnoteReference",
        )
        reference.set(f"{{{word_namespace}}}id", "1")

    def _add_footnote_relationship(root: ET.Element) -> None:
        """向 document relationships 追加 footnotes part 边。"""

        relationship = ET.SubElement(
            root,
            f"{{{package_relationship_namespace}}}Relationship",
        )
        relationship.set("Id", "rIdParaGUIFootnotes")
        relationship.set(
            "Type",
            (
                "http://schemas.openxmlformats.org/officeDocument/2006/"
                "relationships/footnotes"
            ),
        )
        relationship.set("Target", "footnotes.xml")

    def _add_footnote_content_type(root: ET.Element) -> None:
        """向 Content Types 追加 footnotes Override。"""

        override = ET.SubElement(
            root,
            f"{{{content_types_namespace}}}Override",
        )
        override.set("PartName", "/word/footnotes.xml")
        override.set(
            "ContentType",
            (
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.footnotes+xml"
            ),
        )

    _rewrite_docx_xml(source_path, "word/document.xml", _add_body_reference)
    _rewrite_docx_xml(
        source_path,
        "word/_rels/document.xml.rels",
        _add_footnote_relationship,
    )
    _rewrite_docx_xml(
        source_path,
        "[Content_Types].xml",
        _add_footnote_content_type,
    )
    footnotes = ET.Element(f"{{{word_namespace}}}footnotes")
    for note_id, note_text in (("1", "NOTE A"), ("2", "NOTE B")):
        footnote = ET.SubElement(
            footnotes,
            f"{{{word_namespace}}}footnote",
        )
        footnote.set(f"{{{word_namespace}}}id", note_id)
        paragraph = ET.SubElement(footnote, f"{{{word_namespace}}}p")
        run = ET.SubElement(paragraph, f"{{{word_namespace}}}r")
        text = ET.SubElement(run, f"{{{word_namespace}}}t")
        text.text = note_text
    _append_docx_member(
        source_path,
        "word/footnotes.xml",
        ET.tostring(footnotes, encoding="utf-8", xml_declaration=True),
    )
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _swap_footnote_ids(root: ET.Element) -> None:
        """交换两个脚注容器 ID，保持文字和容器顺序。"""

        entries = root.findall(f"{{{word_namespace}}}footnote")
        assert len(entries) == 2
        entries[0].set(f"{{{word_namespace}}}id", "2")
        entries[1].set(f"{{{word_namespace}}}id", "1")

    _rewrite_docx_xml(post_path, "word/footnotes.xml", _swap_footnote_ids)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is False


def test_external_hyperlink_relationship_fails_closed(tmp_path: Path) -> None:
    """验证 baseline 不会跟随或默认信任外部超链接关系。

    输入参数：
        tmp_path：pytest 提供的私有 input 根。
    输出返回值：
        无；固定 input 含 ``TargetMode=External`` 时必须以固定脱敏码
        ERROR，不得将 URL 写入 DTO 或异常。
    """

    source_root = tmp_path / "pre"
    source_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_external_hyperlink_document(source_path)

    with pytest.raises(WordTextFidelityError) as captured:
        capture_word_text_baseline(
            task_id=_TASK_ID,
            protocol_id="paraguibench.operation.eval-rules.v1",
            manifest_sha256=_MANIFEST_SHA256,
            source_root=source_root,
            files=(_identity(source_path),),
        )

    assert captured.value.code == "WORD_TEXT_EXTERNAL_RELATIONSHIP_REJECTED"
    assert "example" not in str(captured.value)


def test_word010_empty_style_wrapper_preserves_text_fidelity(
    tmp_path: Path,
) -> None:
    """验证不带文字语义属性的样式包装不产生假失败。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；post 仅多一层 basedOn Normal 的空段落样式时，
        typed 文字仍应匹配。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    _write_document(source_path, "STYLE-STABLE BODY")
    baseline = capture_word_text_baseline(
        task_id=_WORD010_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post = Document()
    style = post.styles.add_style("Empty Wrapper", WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = post.styles["Normal"]
    post.add_paragraph("STYLE-STABLE BODY", style=style)
    post.save(result_root / source_path.name)

    fidelity = compare_word_text_fidelity(baseline, result_root)

    assert fidelity.matched is True


def test_removed_header_relationship_fails_closed(tmp_path: Path) -> None:
    """验证保留孤儿 header part 不能伪装页眉仍可见。

    输入参数：
        tmp_path：pytest 提供的 pre/post 私有根目录。
    输出返回值：
        无；post 仅删除 document→header relationship，即使
        ``header1.xml`` 与 ``headerReference`` 仍在，也必须 ERROR。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    source_path = source_root / "Document.docx"
    document = Document()
    document.add_paragraph("BODY")
    document.sections[0].header.paragraphs[0].text = "HEADER"
    document.save(source_path)
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=(_identity(source_path),),
    )
    post_path = result_root / source_path.name
    shutil.copy2(source_path, post_path)

    def _remove_header_relationship(root: ET.Element) -> None:
        """删除唯一指向 header part 的 relationship。"""

        for relationship in tuple(root):
            if relationship.get("Target", "").endswith("header1.xml"):
                root.remove(relationship)

    _rewrite_docx_xml(
        post_path,
        "word/_rels/document.xml.rels",
        _remove_header_relationship,
    )

    with pytest.raises(WordTextFidelityError) as captured:
        compare_word_text_fidelity(baseline, result_root)

    assert captured.value.code == "WORD_TEXT_RELATIONSHIP_REFERENCE_UNRESOLVED"


@pytest.mark.parametrize(
    ("field", "value"),
    (
        ("task_id", _WORD010_TASK_ID),
        ("protocol_id", "wrong.protocol"),
        ("manifest_sha256", "2" * 64),
    ),
)
def test_formal_baseline_identity_rejects_wrong_binding(
    tmp_path: Path,
    field: str,
    value: str,
) -> None:
    """验证 baseline 不能跨 task/protocol/manifest 重用。

    输入参数：
        tmp_path：私有 pre 根；field/value：要伪造的身份字段。
    输出返回值：
        无；任一错误绑定均以同一固定错误码拒绝。
    """

    source_path = tmp_path / "Document.docx"
    _write_document(source_path, "BOUND BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=tmp_path,
        files=(_identity(source_path),),
    )
    forged = replace(baseline, **{field: value})

    with pytest.raises(WordTextFidelityError) as captured:
        validate_word_text_baseline_identity(
            forged,
            task_id=_TASK_ID,
            protocol_id="paraguibench.operation.eval-rules.v1",
            manifest_sha256=_MANIFEST_SHA256,
            document_paths=(source_path.name,),
        )

    assert captured.value.code == "WORD_TEXT_BASELINE_IDENTITY_INVALID"


def test_formal_baseline_identity_rejects_path_order_drift(
    tmp_path: Path,
) -> None:
    """验证路径换位不能通过 formal baseline 绑定。

    输入参数：
        tmp_path：私有 pre 根目录。
    输出返回值：
        无；baseline 内部顺序与固定 manifest 路径顺序不同时 ERROR。
    """

    first = tmp_path / "First.docx"
    second = tmp_path / "Second.docx"
    _write_document(first, "FIRST")
    _write_document(second, "SECOND")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=tmp_path,
        files=(_identity(first), _identity(second)),
    )

    with pytest.raises(WordTextFidelityError) as captured:
        validate_word_text_baseline_identity(
            baseline,
            task_id=_TASK_ID,
            protocol_id="paraguibench.operation.eval-rules.v1",
            manifest_sha256=_MANIFEST_SHA256,
            document_paths=(second.name, first.name),
        )

    assert captured.value.code == "WORD_TEXT_BASELINE_IDENTITY_INVALID"


def test_formal_baseline_identity_rejects_forged_internal_shape(
    tmp_path: Path,
) -> None:
    """验证伪造私有 snapshot 字段不泄漏原生 TypeError。

    输入参数：
        tmp_path：私有 pre 根目录。
    输出返回值：
        无；digest 被伪造为 ``None`` 时仍只抛固定 baseline
        身份错误，不将 repr、路径或摘要写入异常。
    """

    source_path = tmp_path / "Document.docx"
    _write_document(source_path, "SHAPE BODY")
    baseline = capture_word_text_baseline(
        task_id=_TASK_ID,
        protocol_id="paraguibench.operation.eval-rules.v1",
        manifest_sha256=_MANIFEST_SHA256,
        source_root=tmp_path,
        files=(_identity(source_path),),
    )
    forged_document = replace(baseline.documents[0], digest=None)
    forged = replace(baseline, documents=(forged_document,))

    with pytest.raises(WordTextFidelityError) as captured:
        validate_word_text_baseline_identity(
            forged,
            task_id=_TASK_ID,
            protocol_id="paraguibench.operation.eval-rules.v1",
            manifest_sha256=_MANIFEST_SHA256,
            document_paths=(source_path.name,),
        )

    assert captured.value.code == "WORD_TEXT_BASELINE_IDENTITY_INVALID"
    assert source_path.name not in str(captured.value)


def test_baseline_source_root_rejects_intermediate_symlink(
    tmp_path: Path,
) -> None:
    """验证 host cache 根的任一中间目录符号链接均被拒绝。

    输入参数：
        tmp_path：pytest 提供的真实根与 symlink 父路径。
    输出返回值：
        无；即使最终 cache 目录和 DOCX 都是普通节点，
        只要绝对根的中间父经过 symlink，held-fd 链就必须
        在读取正文前固定 ERROR。
    """

    real_parent = tmp_path / "real-parent"
    source_root = real_parent / "cache"
    source_root.mkdir(parents=True)
    source_path = source_root / "Document.docx"
    _write_document(source_path, "NOFOLLOW BODY")
    linked_parent = tmp_path / "linked-parent"
    linked_parent.symlink_to(real_parent, target_is_directory=True)

    with pytest.raises(WordTextFidelityError) as captured:
        capture_word_text_baseline(
            task_id=_TASK_ID,
            protocol_id="paraguibench.operation.eval-rules.v1",
            manifest_sha256=_MANIFEST_SHA256,
            source_root=linked_parent / "cache",
            files=(_identity(source_path),),
        )

    assert captured.value.code == "WORD_TEXT_ROOT_INVALID"

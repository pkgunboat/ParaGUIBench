"""Word-012 逐处语境缩写语义协议测试。"""

from __future__ import annotations

import hashlib
from dataclasses import replace
import io
import json
from pathlib import Path
from typing import Callable
import xml.etree.ElementTree as ET
import zipfile

from docx import Document
import pytest

from paraguibench.evaluation.operation import (
    OPERATION_PROTOCOL_ID,
    OperationEvaluationError,
    WordTextInputFile,
    capture_word_abbreviation_baseline,
    evaluate_operation_artifacts,
)


_REPO_ROOT = Path(__file__).resolve().parents[2]
_TASK_ID = "Operation-FileOperate-BatchOperationWord-012"
_MANIFEST_SHA256 = "00b56d5ab84094a98e70156f399881792fe01a649b945284705f79ec050bf1f2"
_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = f"{{{_WORD_NS}}}"
_SOURCE_TEXT = {
    "Clinical Procedure.docx": (
        "Clinical Procedure:\u00a0The anesthesiologist recorded the "
        "MAC (Minimum Alveolar Concentration)\u00a0every ten minutes. Keeping the\u00a0"
        "MAC\u00a0at 1.2% ensured the patient remained stable. Simultaneously, the "
        "nurse used a\u00a0MAC\u00a0tablet to log the patient’s heart rate. The\u00a0MAC\u00a0"
        "software crashed once, but the\u00a0MAC\u00a0value remained visible on the "
        "analog monitor."
    ),
    "Hardware Review.docx": (
        "Hardware Review:\u00a0The new MAC (Macintosh) models are revolutionizing the "
        "creative industry. Most professionals prefer the\u00a0MAC\u00a0for video editing, "
        "but they often forget to check the\u00a0MAC\u00a0address of their wireless "
        "peripherals. Without a valid\u00a0MAC, the Bluetooth mouse cannot pair with "
        "the\u00a0MAC\u00a0workstation."
    ),
    "Infrastructure Log.docx": (
        "Infrastructure Log:\u00a0The engineer first configured the Mac(Media Access "
        "Controller) address to ensure layer 2 connectivity. After verifying the\u00a0"
        "Mac in the router settings, he opened his\u00a0MAC\u00a0to download the latest "
        "firmware. He noticed that the\u00a0MAC\u00a0OS interface was much faster than his "
        "previous LT."
    ),
    "Security Protocol.docx": (
        "Security Protocol:\u00a0Each message is appended with a\u00a0MAC (Message "
        "Authentication Code)\u00a0to prevent tampering. While the\u00a0MAC\u00a0ensures data "
        "integrity, the SA discovered a vulnerability in the\u00a0MAC\u00a0address filtering "
        "of the firewall. The\u00a0MAC\u00a0was spoofed by an attacker using a modified "
        "network card."
    ),
}
_EXPECTED_TEXT = {
    "Clinical Procedure.docx": (
        "Clinical Procedure:\u00a0The anesthesiologist recorded the "
        "MAC (Minimum Alveolar Concentration)\u00a0every ten minutes. Keeping the\u00a0"
        "MAC (Minimum Alveolar Concentration)\u00a0at 1.2% ensured the patient remained "
        "stable. Simultaneously, the nurse used a\u00a0MAC (Macintosh)\u00a0tablet to log the "
        "patient’s heart rate. The\u00a0MAC (Macintosh)\u00a0software crashed once, but the\u00a0"
        "MAC (Minimum Alveolar Concentration)\u00a0value remained visible on the analog "
        "monitor."
    ),
    "Hardware Review.docx": (
        "Hardware Review:\u00a0The new MAC (Macintosh) models are revolutionizing the "
        "creative industry. Most professionals prefer the\u00a0MAC (Macintosh)\u00a0for "
        "video editing, but they often forget to check the\u00a0MAC (Media Access "
        "Control)\u00a0address of their wireless peripherals. Without a valid\u00a0MAC "
        "(Media Access Control), the Bluetooth mouse cannot pair with the\u00a0MAC "
        "(Macintosh)\u00a0workstation."
    ),
    "Infrastructure Log.docx": (
        "Infrastructure Log:\u00a0The engineer first configured the Mac (Media Access "
        "Control) address to ensure layer 2 connectivity. After verifying the\u00a0Mac "
        "(Media Access Control) in the router settings, he opened his\u00a0MAC "
        "(Macintosh)\u00a0to download the latest firmware. He noticed that the\u00a0MAC "
        "(Macintosh)\u00a0OS interface was much faster than his previous LT."
    ),
    "Security Protocol.docx": (
        "Security Protocol:\u00a0Each message is appended with a\u00a0MAC (Message "
        "Authentication Code)\u00a0to prevent tampering. While the\u00a0MAC (Message "
        "Authentication Code)\u00a0ensures data integrity, the SA discovered a "
        "vulnerability in the\u00a0MAC (Media Access Control)\u00a0address filtering of the "
        "firewall. The\u00a0MAC (Media Access Control)\u00a0was spoofed by an attacker using "
        "a modified network card."
    ),
}


def _write_documents(root: Path, texts: dict[str, str]) -> None:
    """写入四份合成 DOCX 闭集。

    输入参数：
        root：已存在的隔离目录；texts：固定文件名到单段正文的映射。
    输出返回值：
        无；每个路径生成可由 production parser 解析的 DOCX。
    """

    for filename, text in texts.items():
        document = Document()
        document.add_paragraph(text)
        document.save(root / filename)


def _write_split_documents(root: Path, texts: dict[str, str]) -> None:
    """将同一可见正文机械拆成多个 run 写入 DOCX。

    输入参数：
        root：已存在的隔离目录；texts：固定文件名到单段
        canonical 正文的映射。
    输出返回值：
        无；每 37 个字符一个 run，但容器和可见文字语义不变。
    """

    for filename, text in texts.items():
        document = Document()
        paragraph = document.add_paragraph()
        for offset in range(0, len(text), 37):
            paragraph.add_run(text[offset : offset + 37])
        document.save(root / filename)


def _write_boundary_split_documents(root: Path, texts: dict[str, str]) -> None:
    """将每份正文在首个 ASCII 空格处拆成两个 run。

    输入参数：
        root：已存在的隔离目录；texts：固定文件名到正文的映射。
    输出返回值：
        无；第二个 ``w:t`` 以空格开头，python-docx 必须为其
        写入 ``xml:space=preserve``。
    """

    for filename, text in texts.items():
        split_at = text.index(" ")
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run(text[:split_at])
        paragraph.add_run(text[split_at:])
        document.save(root / filename)


def _baseline(source_root: Path):
    """从四份合成正式路径构造 typed baseline。

    输入参数：
        source_root：含四份 source DOCX 的私有根。
    输出返回值：
        evaluator-only Word-012 baseline，不持有原文或文件句柄。
    """

    files = []
    for filename in _SOURCE_TEXT:
        payload = (source_root / filename).read_bytes()
        files.append(
            WordTextInputFile(
                path=filename,
                size=len(payload),
                sha256=hashlib.sha256(payload).hexdigest(),
                is_docx=True,
            )
        )
    return capture_word_abbreviation_baseline(
        task_id=_TASK_ID,
        protocol_id=OPERATION_PROTOCOL_ID,
        manifest_sha256=_MANIFEST_SHA256,
        source_root=source_root,
        files=tuple(files),
    )


def _task() -> dict[str, object]:
    """读取题面未改的 canonical Word-012 任务。

    输入参数：无。
    输出返回值：完整 trusted task JSON 映射。
    """

    path = _REPO_ROOT / "benchmark/tasks" / f"{_TASK_ID}.json"
    return json.loads(path.read_text(encoding="utf-8"))


def _rewrite_document_xml(
    path: Path,
    mutate: Callable[[ET.Element], None],
) -> None:
    """在测试临时 DOCX 内就地应用一个 OOXML 对抗变换。

    输入参数：
        path：临时 DOCX 路径；mutate：只接收 document.xml 根的
        合成变换函数。
    输出返回值：
        无；除 document.xml 外的 ZIP member 及其元数据原样保留。
    """

    with zipfile.ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    rewritten: list[tuple[zipfile.ZipInfo, bytes]] = []
    for info, payload in entries:
        if info.filename == "word/document.xml":
            root = ET.fromstring(payload)
            mutate(root)
            payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        rewritten.append((info, payload))
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w") as archive:
        for info, payload in rewritten:
            archive.writestr(info, payload)
    path.write_bytes(output.getvalue())


def test_word012_exact_contextual_expansions_pass_with_fixed_denominator(
    tmp_path: Path,
) -> None:
    """验证四份文档按逐处语境映射展开后满分。

    输入参数：
        tmp_path：pytest 提供的 pre/post 隔离根。
    输出返回值：
        无；麻醉、Macintosh、Media Access Control 与 Message
        Authentication Code 必须按各自语境匹配，固定分母为 4。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    _write_documents(source_root, _SOURCE_TEXT)
    _write_documents(result_root, _EXPECTED_TEXT)

    evaluation = evaluate_operation_artifacts(
        result_root,
        _task(),
        input_abbreviation_baseline=_baseline(source_root),
    )

    assert evaluation.passed is True
    assert evaluation.score == 1.0
    assert evaluation.artifact_count == 4
    assert evaluation.evaluated_rule_count == 1
    assert evaluation.rule_results[0].evaluated_artifact_count == 4


def test_word012_equivalent_run_splits_remain_valid(
    tmp_path: Path,
) -> None:
    """验证合法 Writer 仅重分 run 不会被误判为正文漂移。

    输入参数：
        tmp_path：pre/post 隔离根。
    输出返回值：
        无；post 的 canonical 文字可跨多个 run，四文档仍应 PASS/1。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    _write_documents(source_root, _SOURCE_TEXT)
    _write_split_documents(result_root, _EXPECTED_TEXT)

    evaluation = evaluate_operation_artifacts(
        result_root,
        _task(),
        input_abbreviation_baseline=_baseline(source_root),
    )

    assert evaluation.passed is True
    assert evaluation.score == 1.0
    assert evaluation.rule_results[0].evaluated_artifact_count == 4


def test_word012_task_instruction_is_unchanged_and_contains_no_private_mapping() -> (
    None
):
    """验证 task JSON 题面原样保留且不携带 host-only 语义值。

    输入参数：无；读取 canonical task。
    输出返回值：
        无；英文 instruction 必须逐字符匹配原题，task 对象不得
        容纳 task-specific 释义映射。
    """

    task = _task()
    assert task["instruction"] == (
        "Different documents in the folder use the same abbreviation to represent "
        "different concepts. Please expand these abbreviations where this situation "
        "occurs."
    )
    serialized = json.dumps(task, ensure_ascii=False)
    for private_mapping_value in (
        "Minimum Alveolar Concentration",
        "Macintosh",
        "Media Access Control",
        "Message Authentication Code",
    ):
        assert private_mapping_value not in serialized


@pytest.mark.parametrize("carrier", ("hidden", "revision", "textbox"))
def test_word012_visible_text_carrier_drift_cannot_bypass_fidelity(
    tmp_path: Path,
    carrier: str,
) -> None:
    """验证隐藏、修订和文本框容器漂移不能伪装语义正确。

    输入参数：
        tmp_path：pre/post 隔离根；carrier：要注入的可见性或
        OOXML 语义容器漂移。
    输出返回值：
        无；尽管抽取纯文字不变，typed 结构门禁必须 FAIL/0。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    _write_documents(source_root, _SOURCE_TEXT)
    _write_documents(result_root, _EXPECTED_TEXT)
    target = result_root / "Clinical Procedure.docx"
    if carrier == "hidden":
        document = Document(target)
        document.paragraphs[0].runs[0].font.hidden = True
        document.save(target)
    else:
        wrapper_tag = "del" if carrier == "revision" else "txbxContent"

        def _wrap_run(root: ET.Element) -> None:
            """将原文 run 放入修订或文本框语义容器。

            输入参数：
                root：当前临时 DOCX 的 ``word/document.xml`` 根。
            输出返回值：
                无；就地将首个正文 run 用当前对抗容器包裹。
            """

            paragraph = next(root.iter(f"{_W}p"))
            run = next(child for child in paragraph if child.tag == f"{_W}r")
            index = list(paragraph).index(run)
            paragraph.remove(run)
            wrapper = ET.Element(f"{_W}{wrapper_tag}")
            wrapper.append(run)
            paragraph.insert(index, wrapper)

        _rewrite_document_xml(target, _wrap_run)

    evaluation = evaluate_operation_artifacts(
        result_root,
        _task(),
        input_abbreviation_baseline=_baseline(source_root),
    )

    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.reason_codes == ("ABBREVIATION_SEMANTICS_MISMATCH",)
    assert evaluation.rule_results[0].evaluated_artifact_count == 4


@pytest.mark.parametrize("mutation", ("unknown-wrapper", "malformed-package"))
def test_word012_unreliable_post_is_private_error_not_semantic_failure(
    tmp_path: Path,
    mutation: str,
) -> None:
    """验证未知可见载体或损坏容器不被折算为普通 FAIL。

    输入参数：
        tmp_path：pre/post 隔离根；mutation：未登记语义 wrapper
        或无效 ZIP 容器。
    输出返回值：
        无；证据不可靠时必须抛脱敏 ``OperationEvaluationError``，
        不得产生可比较语义零分。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    _write_documents(source_root, _SOURCE_TEXT)
    _write_documents(result_root, _EXPECTED_TEXT)
    target = result_root / "Clinical Procedure.docx"
    if mutation == "malformed-package":
        target.write_bytes(b"private malformed docx sentinel")
    else:

        def _inject_unknown_wrapper(root: ET.Element) -> None:
            """用未登记 namespace 包裹保持文字不变的正文 run。

            输入参数：
                root：当前临时 DOCX 的 document.xml 根。
            输出返回值：
                无；就地注入一个子树含已知文字载体的未知 wrapper。
            """

            paragraph = next(root.iter(f"{_W}p"))
            run = next(child for child in paragraph if child.tag == f"{_W}r")
            index = list(paragraph).index(run)
            paragraph.remove(run)
            wrapper = ET.Element("{urn:private-unknown}wrapper")
            wrapper.append(run)
            paragraph.insert(index, wrapper)

        _rewrite_document_xml(target, _inject_unknown_wrapper)

    with pytest.raises(OperationEvaluationError) as captured:
        evaluate_operation_artifacts(
            result_root,
            _task(),
            input_abbreviation_baseline=_baseline(source_root),
        )

    assert captured.value.code in {
        "WORD_ABBREVIATION_SEMANTICS_INVALID",
        "ARCHIVE_INVALID",
    }
    assert "Clinical Procedure" not in str(captured.value)
    assert "private" not in str(captured.value).casefold()


@pytest.mark.parametrize(
    "mutation",
    (
        "unchanged",
        "empty-parentheses",
        "arbitrary-parentheses",
        "unified-meaning",
        "delete-abbreviation",
        "expand-nontargets",
        "delete-unrelated-text",
        "reorder-body",
    ),
)
def test_word012_semantic_or_body_drift_fails_closed(
    tmp_path: Path,
    mutation: str,
) -> None:
    """验证伪释义、全统一、非目标扩展与正文漂移均零分。

    输入参数：
        tmp_path：pre/post 隔离根；mutation：一种可比较的对抗改写。
    输出返回值：
        无；任一漂移都必须以固定原因、固定四文档分母 FAIL/0。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    _write_documents(source_root, _SOURCE_TEXT)
    texts = dict(_EXPECTED_TEXT)
    if mutation == "unchanged":
        texts = dict(_SOURCE_TEXT)
    elif mutation == "empty-parentheses":
        texts["Security Protocol.docx"] = texts["Security Protocol.docx"].replace(
            "MAC (Message Authentication Code)\u00a0ensures",
            "MAC ()\u00a0ensures",
        )
    elif mutation == "arbitrary-parentheses":
        texts["Security Protocol.docx"] = texts["Security Protocol.docx"].replace(
            "MAC (Message Authentication Code)\u00a0ensures",
            "MAC (banana)\u00a0ensures",
        )
    elif mutation == "unified-meaning":
        texts = {
            name: text.replace("Minimum Alveolar Concentration", "Macintosh")
            .replace("Media Access Control", "Macintosh")
            .replace("Message Authentication Code", "Macintosh")
            for name, text in texts.items()
        }
    elif mutation == "delete-abbreviation":
        texts["Hardware Review.docx"] = texts["Hardware Review.docx"].replace(
            "MAC (Macintosh)\u00a0workstation",
            "Macintosh workstation",
        )
    elif mutation == "expand-nontargets":
        texts["Infrastructure Log.docx"] = texts["Infrastructure Log.docx"].replace(
            "previous LT.",
            "previous LT (Laptop).",
        )
        texts["Security Protocol.docx"] = texts["Security Protocol.docx"].replace(
            "the SA discovered",
            "the SA (Security Analyst) discovered",
        )
    elif mutation == "delete-unrelated-text":
        texts["Hardware Review.docx"] = texts["Hardware Review.docx"].replace(
            "creative industry.",
            "industry.",
        )
    elif mutation == "reorder-body":
        texts["Clinical Procedure.docx"] = texts["Clinical Procedure.docx"].replace(
            "The\u00a0MAC (Macintosh)\u00a0software crashed once, but the\u00a0MAC "
            "(Minimum Alveolar Concentration)\u00a0value remained visible",
            "The\u00a0MAC (Minimum Alveolar Concentration)\u00a0value remained visible, "
            "but the\u00a0MAC (Macintosh)\u00a0software crashed once",
        )
    _write_documents(result_root, texts)

    evaluation = evaluate_operation_artifacts(
        result_root,
        _task(),
        input_abbreviation_baseline=_baseline(source_root),
    )

    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.reason_codes == ("ABBREVIATION_SEMANTICS_MISMATCH",)
    assert evaluation.rule_results[0].evaluated_artifact_count == 4


@pytest.mark.parametrize("mutation", ("missing", "extra"))
def test_word012_document_closed_set_uses_fixed_denominator(
    tmp_path: Path,
    mutation: str,
) -> None:
    """验证缺文档或多文件不能改变 Word-012 固定分母。

    输入参数：
        tmp_path：隔离输入输出根；mutation：删除或添加文件。
    输出返回值：
        无；输出闭集不等时以 artifact contract FAIL/0，分母仍为 4。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    _write_documents(source_root, _SOURCE_TEXT)
    _write_documents(result_root, _EXPECTED_TEXT)
    if mutation == "missing":
        (result_root / "Clinical Procedure.docx").unlink()
    else:
        (result_root / "extra.txt").write_text("extra", encoding="utf-8")

    evaluation = evaluate_operation_artifacts(
        result_root,
        _task(),
        input_abbreviation_baseline=_baseline(source_root),
    )

    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.reason_codes == ("ARTIFACT_CONTRACT_MISMATCH",)
    assert evaluation.rule_results[0].evaluated_artifact_count == 4


def test_word012_baseline_identity_is_required_before_post_io(
    tmp_path: Path,
) -> None:
    """验证缺失或错 manifest baseline 在任何 post 路径读取前 ERROR。

    输入参数：
        tmp_path：用于构造有效 pre DTO 的隔离根。
    输出返回值：
        无；不存在的 post 根不得抢先产生 artifact root 错误。
    """

    source_root = tmp_path / "pre"
    source_root.mkdir()
    _write_documents(source_root, _SOURCE_TEXT)
    baseline = _baseline(source_root)
    missing_root = tmp_path / "not-created"

    with pytest.raises(
        OperationEvaluationError,
        match="WORD_ABBREVIATION_BASELINE_REQUIRED",
    ):
        evaluate_operation_artifacts(missing_root, _task())
    with pytest.raises(
        OperationEvaluationError,
        match="WORD_ABBREVIATION_SEMANTICS_INVALID",
    ):
        evaluate_operation_artifacts(
            missing_root,
            _task(),
            input_abbreviation_baseline=replace(
                baseline,
                manifest_sha256="0" * 64,
            ),
        )


@pytest.mark.parametrize(
    "mutation",
    ("task", "protocol", "path", "count", "digest-shape"),
)
def test_word012_all_forged_baseline_identity_fields_are_private_errors(
    tmp_path: Path,
    mutation: str,
) -> None:
    """验证任务、协议、路径、数量与私有形状伪造均 ERROR。

    输入参数：
        tmp_path：构造有效 pre DTO 的隔离根；mutation：要漂移的
        baseline 身份维度。
    输出返回值：
        无；全部变体必须在不读 post 的情况下映射为同一
        脱敏 ``OperationEvaluationError``。
    """

    source_root = tmp_path / "pre"
    source_root.mkdir()
    _write_documents(source_root, _SOURCE_TEXT)
    baseline = _baseline(source_root)
    if mutation == "task":
        forged = replace(baseline, task_id="private-forged-task")
    elif mutation == "protocol":
        forged = replace(baseline, protocol_id="private-forged-protocol")
    elif mutation == "path":
        first = replace(baseline.expected_documents[0], path="private.docx")
        forged = replace(
            baseline,
            expected_documents=(first, *baseline.expected_documents[1:]),
        )
    elif mutation == "count":
        forged = replace(
            baseline,
            expected_documents=baseline.expected_documents[:-1],
        )
    else:
        first = replace(baseline.expected_documents[0], digest=None)
        forged = replace(
            baseline,
            expected_documents=(first, *baseline.expected_documents[1:]),
        )

    with pytest.raises(
        OperationEvaluationError,
        match="WORD_ABBREVIATION_SEMANTICS_INVALID",
    ) as captured:
        evaluate_operation_artifacts(
            tmp_path / "post-does-not-exist",
            _task(),
            input_abbreviation_baseline=forged,
        )

    for private_value in (
        "private-forged-task",
        "private-forged-protocol",
        "private.docx",
    ):
        assert private_value not in str(captured.value)


@pytest.mark.parametrize(
    "field",
    (
        "digest",
        "token_count",
        "part_count",
        "relationship_digest",
        "image_relationships",
        "media_identities",
    ),
)
def test_word012_valid_shaped_snapshot_drift_is_rejected_before_post_io(
    tmp_path: Path,
    field: str,
) -> None:
    """验证合法形状不能代替 prepare 时封存的复合快照。

    输入参数：
        tmp_path：用于构造有效 pre DTO 的隔离根；field：要替换为
        另一个 schema-valid 值的私有快照字段。
    输出返回值：
        无；六类漂移都必须在不存在的 post 路径被读取前统一
        映射为 ``WORD_ABBREVIATION_SEMANTICS_INVALID``。
    """

    source_root = tmp_path / "pre"
    source_root.mkdir()
    _write_documents(source_root, _SOURCE_TEXT)
    baseline = _baseline(source_root)
    first = baseline.expected_documents[0]
    mutations = {
        "digest": {"digest": "0" * 64},
        "token_count": {"token_count": first.token_count + 1},
        "part_count": {"part_count": first.part_count + 1},
        "relationship_digest": {"relationship_digest": "0" * 64},
        "image_relationships": {
            "image_relationships": (
                (
                    "word/document.xml",
                    (
                        "http://schemas.openxmlformats.org/officeDocument/"
                        "2006/relationships/image"
                    ),
                    "word/media/private.png",
                ),
            )
        },
        "media_identities": {
            "media_identities": (("word/media/private.png", "0" * 64),)
        },
    }
    forged_first = replace(first, **mutations[field])
    forged = replace(
        baseline,
        expected_documents=(forged_first, *baseline.expected_documents[1:]),
    )

    with pytest.raises(
        OperationEvaluationError,
        match="WORD_ABBREVIATION_SEMANTICS_INVALID",
    ):
        evaluate_operation_artifacts(
            tmp_path / "post-does-not-exist",
            _task(),
            input_abbreviation_baseline=forged,
        )


@pytest.mark.parametrize("direction_semantics", ("bdo", "dir", "rtl"))
def test_word012_direction_semantics_drift_cannot_preserve_score(
    tmp_path: Path,
    direction_semantics: str,
) -> None:
    """验证可改变阅读顺序的 Word 方向语义必须进入快照。

    输入参数：
        tmp_path：pre/post 隔离根；direction_semantics：``w:bdo``、
        ``w:dir`` 容器或 ``w:rtl`` run 属性。
    输出返回值：
        无；即使字符序列不变，任一方向漂移也必须 FAIL/0。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    _write_documents(source_root, _SOURCE_TEXT)
    _write_documents(result_root, _EXPECTED_TEXT)

    def _inject_direction_semantics(root: ET.Element) -> None:
        """在首个正文 run 上注入方向容器或属性。

        输入参数：root：当前临时 DOCX 的 document.xml 根。
        输出返回值：无；就地修改首个承载正文的 run。
        """

        paragraph = next(root.iter(f"{_W}p"))
        run = next(child for child in paragraph if child.tag == f"{_W}r")
        if direction_semantics in {"bdo", "dir"}:
            index = list(paragraph).index(run)
            paragraph.remove(run)
            wrapper = ET.Element(
                f"{_W}{direction_semantics}",
                {f"{_W}val": "rtl"},
            )
            wrapper.append(run)
            paragraph.insert(index, wrapper)
            return
        run_properties = run.find(f"{_W}rPr")
        if run_properties is None:
            run_properties = ET.Element(f"{_W}rPr")
            run.insert(0, run_properties)
        ET.SubElement(run_properties, f"{_W}rtl")

    _rewrite_document_xml(
        result_root / "Clinical Procedure.docx",
        _inject_direction_semantics,
    )

    evaluation = evaluate_operation_artifacts(
        result_root,
        _task(),
        input_abbreviation_baseline=_baseline(source_root),
    )

    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.reason_codes == ("ABBREVIATION_SEMANTICS_MISMATCH",)


def test_word012_boundary_space_requires_preserve_semantics(tmp_path: Path) -> None:
    """验证边界空格的 ``xml:space`` 不得被当成无关属性。

    输入参数：tmp_path：pre/post 隔离根。
    输出返回值：
        无；合法的等价 run 拆分先 PASS，仅删除以空格开头的
        ``w:t`` 上 ``xml:space=preserve`` 后必须 FAIL/0。
    """

    source_root = tmp_path / "pre"
    result_root = tmp_path / "post"
    source_root.mkdir()
    result_root.mkdir()
    _write_documents(source_root, _SOURCE_TEXT)
    _write_boundary_split_documents(result_root, _EXPECTED_TEXT)
    baseline = _baseline(source_root)

    valid = evaluate_operation_artifacts(
        result_root,
        _task(),
        input_abbreviation_baseline=baseline,
    )
    assert valid.passed is True

    xml_space = "{http://www.w3.org/XML/1998/namespace}space"

    def _remove_boundary_space_contract(root: ET.Element) -> None:
        """删除首个边界空格文本节点的 preserve 语义。

        输入参数：root：当前 document.xml 根。
        输出返回值：无；就地删除且断言测试前置存在。
        """

        target = next(
            element
            for element in root.iter(f"{_W}t")
            if (element.text or "").startswith(" ")
        )
        assert target.get(xml_space) == "preserve"
        del target.attrib[xml_space]

    _rewrite_document_xml(
        result_root / "Clinical Procedure.docx",
        _remove_boundary_space_contract,
    )

    evaluation = evaluate_operation_artifacts(
        result_root,
        _task(),
        input_abbreviation_baseline=baseline,
    )
    assert evaluation.passed is False
    assert evaluation.score == 0.0
    assert evaluation.reason_codes == ("ABBREVIATION_SEMANTICS_MISMATCH",)
